import io
import json
import os
import sys
import tempfile
import unittest
import zipfile
from datetime import datetime
from pathlib import Path
from unittest.mock import patch
from urllib.parse import unquote

from docx import Document
from openpyxl import Workbook, load_workbook
from PIL import Image
from werkzeug.security import generate_password_hash


TEST_DIR = tempfile.TemporaryDirectory()
DB_PATH = Path(TEST_DIR.name) / "svai-test.db"
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
os.environ.update({
    "DATABASE_URL": f"sqlite:///{DB_PATH.as_posix()}",
    "SECRET_KEY": "test-secret",
    "ENCRYPTION_KEY": "render-style-secret-that-is-not-fernet-formatted",
    "ADMIN_EMAIL": "admin@example.com",
    "ADMIN_PASSWORD": "ChangeMe123!",
    "ENABLE_EMAIL_SCHEDULER": "false",
    "SESSION_COOKIE_SECURE": "false",
    "OPENAI_API_KEY": "",
    "OPENAI_DOCUMENT_EXTRACTION": "false",
})

from ai_service_openai import (  # noqa: E402
    build_case_profile, deterministic_email_candidate,
    enrich_email_details_from_attachments, extract_property_asset,
    regex_email_extract,
)
from report_service import fill_docx_template, fill_excel_template  # noqa: E402
from server import (  # noqa: E402
    BillingTemplate, EmailAccount, FileAsset, SiteEngineer, WhatsAppGroup, User, ValuationCase, app, apply_email_details,
    apply_followup_to_existing_case, db, encrypt_password, is_followup_email,
    existing_case_for_duplicate_assignment, normalized_application_number,
    safe_json, valuation_defaults_from_profile, billing_fee_for_km,
    billing_column_map, generate_billing_workbook, merge_cross_mailbox_duplicate_cases,
    email_fetch_folders, mis_import_rows,
    concise_mis_address, mailbox_source, normalize_whatsapp_group_link,
    fetch_full_message, fetch_mis_message, imap_safe_assignment_folders,
)


class SvaiSmokeTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        app.config.update(TESTING=True)
        with app.app_context():
            db.drop_all()
            db.create_all()

    def setUp(self):
        self.client = app.test_client()
        self.client.get("/login")

    def csrf(self):
        with self.client.session_transaction() as session:
            return session["_csrf_token"]

    def login(self):
        return self.client.post("/login", data={
            "_csrf_token": self.csrf(),
            "email": "admin@example.com",
            "password": "ChangeMe123!",
        })

    def test_health_and_authentication(self):
        self.assertEqual(self.client.get("/health").status_code, 200)
        self.assertEqual(self.client.get("/").status_code, 302)
        response = self.login()
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.headers["Location"], "/")

    def test_configured_admin_password_repairs_stale_persistent_hash(self):
        address = "render-admin@example.com"
        configured_password = "ConfiguredPassword123!"
        with patch.dict(os.environ, {
            "ADMIN_EMAIL": address,
            "ADMIN_PASSWORD": configured_password,
        }):
            self.client.get("/login")
            with app.app_context():
                user = User.query.filter_by(email=address).first()
                self.assertIsNotNone(user)
                user.password_hash = generate_password_hash("StalePassword123!")
                db.session.commit()
            response = self.client.post("/login", data={
                "_csrf_token": self.csrf(),
                "email": address,
                "password": configured_password,
            })
            self.assertEqual(response.status_code, 302)
            self.assertEqual(response.headers["Location"], "/")

    def test_primary_pages_render(self):
        self.login()
        for path in ["/", "/email-accounts", "/site-engineers", "/templates", "/billing", "/settings"]:
            response = self.client.get(path)
            self.assertEqual(response.status_code, 200, path)
            self.assertIn(b"Dashboard / MIS", response.data)
        with app.app_context():
            case = ValuationCase.query.order_by(ValuationCase.id).first()
            case_id = case.id
        response = self.client.get(f"/cases/{case_id}")
        self.assertEqual(response.status_code, 200)
        self.assertIn(b"Dashboard / MIS", response.data)
        self.assertIn(b"Process All Files", response.data)

    def test_dashboard_announces_one_minute_auto_refresh(self):
        self.login()
        response = self.client.get("/")
        self.assertEqual(response.status_code, 200)
        self.assertIn(b"Auto fetch + refresh: every 1 minute", response.data)
        self.assertIn(b"monthly catch-up: hourly", response.data)
        self.assertIn(b"60000", response.data)

    def test_billing_fills_existing_invoice_table_with_km_slab_amount(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet["A1"] = "Original bank invoice header"
        for column, header in enumerate((
            "S.No.", "Lead ID No.", "Product Name", "Customer Name",
            "Property Address", "Distance (K.M)", "Fee",
        ), 1):
            sheet.cell(5, column).value = header
        sheet["G8"] = "Total"
        source = io.BytesIO()
        workbook.save(source)
        result, unmatched = generate_billing_workbook(source.getvalue(), [{
            "application_number": "APP-1", "customer_name": "Asha",
            "case_type": "LAP", "property_address": "Gwalior", "distance": 35,
        }], [(0, 30, 1500), (30.01, 50, 1800)])
        filled = load_workbook(io.BytesIO(result)).active
        self.assertEqual(filled["A6"].value, 1)
        self.assertEqual(filled["B6"].value, "APP-1")
        self.assertEqual(filled["G6"].value, 1800)
        self.assertEqual(unmatched, [])
        self.assertEqual(billing_fee_for_km(None, [(0, 30, 1500)]), None)

    def test_billing_page_generates_from_live_mis_and_saved_bank_format(self):
        workbook = Workbook()
        sheet = workbook.active
        for column, header in enumerate((
            "S.No.", "Lead ID No.", "Product Name", "Customer Name",
            "Property Address", "Distance (K.M)", "Fee",
        ), 1):
            sheet.cell(2, column).value = header
        sheet["G5"] = "Total"
        source = io.BytesIO()
        workbook.save(source)
        with app.app_context():
            template = BillingTemplate(
                bank_name="Test Bank", filename="test-bank.xlsx",
                content=source.getvalue(), mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            case = ValuationCase(
                application_number="BILL-1", customer_name="Billing Customer",
                bank_name="Test Bank", case_type="LAP", property_address="Gwalior",
                extracted_json=json.dumps({"case_profile": {"distance_from_branch": "25"}}),
            )
            db.session.add_all([template, case])
            db.session.commit()
            template_id = template.id
        self.login()
        response = self.client.post("/billing", data={
            "_csrf_token": self.csrf(), "bank_name": "Test Bank",
            "billing_template_id": str(template_id), "source": "live",
            "from": "2026-01-01", "to": "2027-01-01",
            "slab_min[]": "0", "slab_max[]": "30", "slab_amount[]": "1500",
        })
        self.assertEqual(response.status_code, 200)
        result = load_workbook(io.BytesIO(response.data)).active
        self.assertEqual(result["B3"].value, "BILL-1")
        self.assertEqual(result["G3"].value, 1500)

    def test_mis_km_save_is_available_to_billing(self):
        with app.app_context():
            case = ValuationCase(application_number="KM-1", customer_name="KM Case")
            db.session.add(case)
            db.session.commit()
            case_id = case.id
        self.login()
        response = self.client.post(f"/cases/{case_id}/km", data={
            "_csrf_token": self.csrf(), "distance_from_branch": "42.5",
        })
        self.assertEqual(response.status_code, 302)
        with app.app_context():
            case = db.session.get(ValuationCase, case_id)
            self.assertEqual(
                safe_json(case.extracted_json)["case_profile"]["distance_from_branch"], 42.5
            )

    def test_forgot_password_resets_through_saved_linked_mailbox(self):
        address = "reset-user@example.com"
        with app.app_context():
            user = User(
                email=address,
                password_hash=generate_password_hash("OldPassword123!"),
                name="Reset User",
            )
            account = EmailAccount(
                email="saved-mail@gmail.com",
                encrypted_password=encrypt_password("abcdefghijklmnop"),
                provider="gmail",
                active=True,
            )
            db.session.add_all([user, account])
            db.session.commit()

        sent = {}

        def capture_code(account, recipient, code):
            sent.update({"recipient": recipient, "code": code})

        with patch("server.send_password_reset_code", side_effect=capture_code):
            response = self.client.post("/forgot-password", data={
                "_csrf_token": self.csrf(),
                "email": address,
            })
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.headers["Location"], "/reset-password")
        self.assertEqual(sent["recipient"], address)
        self.assertRegex(sent["code"], r"^\d{6}$")

        response = self.client.post("/reset-password", data={
            "_csrf_token": self.csrf(),
            "code": sent["code"],
            "new_password": "NewPassword123!",
            "confirm_password": "NewPassword123!",
        })
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.headers["Location"], "/login")
        self.client.get("/login")
        response = self.client.post("/login", data={
            "_csrf_token": self.csrf(),
            "email": address,
            "password": "NewPassword123!",
        })
        self.assertEqual(response.status_code, 302)

    def test_forgot_password_falls_back_to_another_linked_mailbox(self):
        address = "reset-admin@yahoo.com"
        with app.app_context():
            db.session.add(User(
                email=address,
                password_hash=generate_password_hash("OldPassword123!"),
                name="Reset Admin",
            ))
            db.session.add_all([
                EmailAccount(
                    email=address,
                    encrypted_password=encrypt_password("abcdefghijklmnop"),
                    provider="yahoo",
                    active=True,
                ),
                EmailAccount(
                    email="backup-sender@gmail.com",
                    encrypted_password=encrypt_password("ponmlkjihgfedcba"),
                    provider="gmail",
                    active=True,
                ),
            ])
            db.session.commit()

        senders = []

        def fail_then_send(account, recipient, code):
            senders.append(account.email)
            if account.email == address:
                raise OSError("stale Yahoo app password")

        with patch("server.send_password_reset_code", side_effect=fail_then_send):
            response = self.client.post("/forgot-password", data={
                "_csrf_token": self.csrf(),
                "email": address,
            })
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.headers["Location"], "/reset-password")
        self.assertEqual(senders[0], address)
        self.assertGreaterEqual(len(senders), 2)
        self.assertNotEqual(senders[1], address)

    def test_forgot_password_uses_admin_recovery_when_smtp_is_unavailable(self):
        address = "recovery-admin@example.com"
        recovery_code = "SVAI-Recovery-4827"
        with app.app_context():
            user = User(
                email=address,
                password_hash=generate_password_hash("OldPassword123!"),
                name="Recovery Admin",
            )
            db.session.add(user)
            db.session.commit()

        with patch.dict(os.environ, {
            "ADMIN_EMAIL": address,
            "ADMIN_RECOVERY_CODE": recovery_code,
        }), patch(
            "server.send_password_reset_code",
            side_effect=OSError("SMTP unavailable"),
        ):
            response = self.client.post("/forgot-password", data={
                "_csrf_token": self.csrf(),
                "email": address,
            })
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.headers["Location"], "/reset-password")

        response = self.client.post("/reset-password", data={
            "_csrf_token": self.csrf(),
            "code": recovery_code,
            "new_password": "RecoveredPassword123!",
            "confirm_password": "RecoveredPassword123!",
        })
        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.headers["Location"], "/login")

    def test_free_local_document_extraction_and_profile(self):
        document_text = (
            "Technical Valuation Report\n"
            "Application No: LAPVDS100026755\n"
            "Customer Name: MAHENDRA KIRAR\n"
            "Property Address as per Documents: Plot 12, Village Vidisha, "
            "Tehsil Vidisha, District Vidisha, Madhya Pradesh 464001\n"
            "Khasra No: 151/2\n"
            "Land Area as per Documents: 1162 sq ft\n"
            "Built-up Area as per Documents: 867 sq ft\n"
            "Property Address as per Site: THIS MUST NOT ENTER SITE DATA\n"
            "Land Area as per Site: 9999 sq ft\n"
            "Government Land Rate: 27\n"
            "Government Construction Rate: 300\n"
            "North Boundary: Road\nSouth Boundary: Plot 13\n"
        )
        document = extract_property_asset(
            "technical_report.pdf", b"", document_text, "property_document"
        )
        visit = extract_property_asset(
            "engineer_visit.pdf",
            b"",
            "Property Address as per Site: Plot 12, Vidisha\n"
            "Actual Khasra No: 151/3\n"
            "Land Area as per Site: 1042 sq ft\n"
            "Built-up Area as per Site: 867 sq ft\n"
            "Property Address as per Documents: MUST NOT ENTER DOC DATA\n"
            "Land Area as per Documents: 8888 sq ft\n"
            "Road Width: 10 ft\n",
            "visit_data",
        )
        profile = build_case_profile(
            {"customer_name": "MAHENDRA KIRAR"},
            [document],
            [visit],
            {},
        )
        self.assertEqual(document["application_number"], "LAPVDS100026755")
        self.assertEqual(document["applicant_name"], "MAHENDRA KIRAR")
        self.assertIn("1162", document["land_area_as_per_docs"])
        self.assertEqual(document["property_address_as_per_site"], "")
        self.assertEqual(document["land_area_as_per_site"], "")
        self.assertIn("151/2", document["survey_khasra_plot_no_as_per_docs"])
        self.assertIn("1042", visit["land_area_as_per_site"])
        self.assertEqual(visit["property_address_as_per_docs"], "")
        self.assertEqual(visit["land_area_as_per_docs"], "")
        self.assertIn("151/3", visit["survey_khasra_plot_no_as_per_site"])
        self.assertEqual(document["govt_land_rate"], "27")
        self.assertEqual(document["govt_construction_rate"], "300")
        self.assertEqual(profile["processing_mode"], "Free Local")
        self.assertIn("Plot 12", profile["property_address_as_per_docs"])
        self.assertIn("151/2", profile["survey_khasra_plot_no_as_per_docs"])
        self.assertIn("151/3", profile["survey_khasra_plot_no_as_per_site"])

    def test_flattened_technical_report_recovers_docs_site_boundaries_and_areas(self):
        text = """
        PROPERTY ADDRESS
        Plot no.- 50, Part of Survey no. 62/1, Vill- Katsara, Dist- Vidisha MP 464001
        Part of Survey no. 62/1, Vill- Katsara, Dist- Vidisha MP 464001
        Total Permissible Built up area of Land in sqft. 867.00
        Total Built up area of Land in sqft. 867.00
        As Per Document Road H/O Prem narayan Ahirwar H/O Rajaram H/o Narmadaprashad
        Actual Road H/O Prem narayan Ahirwar 6' Gali Then H/O Rajaram H/o Narmadaprashad
        The plot area as per docs is 1162.00 sqft. And actual area at site 1042.00 sqft
        Plot demarcated at site Yes
        Occupied Status Self
        Nature of Construction Load Bearing
        No. Of Floors (Permissible & Actual) 2
        14
        46
        100%
        16-06-2026
        """
        docs = extract_property_asset(
            "TECHNICAL_REPORT.PDF", b"", text, "property_document"
        )
        site = extract_property_asset(
            "TECHNICAL_REPORT.PDF", b"", text, "visit_data"
        )
        self.assertEqual(docs["land_area_as_per_docs"], "1162.00 Sqft.")
        self.assertEqual(site["land_area_as_per_site"], "1042.00 Sqft.")
        self.assertEqual(docs["east_boundary_as_per_docs"], "Road")
        self.assertEqual(docs["north_boundary_as_per_docs"], "H/O Rajaram")
        self.assertEqual(site["north_boundary_as_per_site"], "6' Gali Then H/O Rajaram")
        self.assertEqual(site["south_boundary_as_per_site"], "H/o Narmadaprashad")
        self.assertEqual(docs["builtup_area_as_per_docs"], "867.00 Sqft.")
        self.assertEqual(site["builtup_area_as_per_site"], "867.00 Sqft.")
        self.assertEqual(site["plot_demarcated"], "Yes")
        self.assertEqual(site["occupancy"], "Self Occupied")
        self.assertEqual(site["property_age_years"], "14")
        self.assertEqual(site["residual_age_years"], "46")
        self.assertEqual(site["visit_date"], "16-06-2026")

    def test_existing_mis_row_is_stable_when_same_case_is_fetched_again(self):
        with app.app_context():
            account = EmailAccount.query.filter_by(
                email="stable-mis@gmail.com"
            ).first()
            if not account:
                account = EmailAccount(
                    email="stable-mis@gmail.com",
                    encrypted_password=encrypt_password("abcdefghijklmnop"),
                    provider="gmail",
                    bank_name="Test Bank",
                    active=True,
                )
                db.session.add(account)
            case = ValuationCase(
                application_number="STABLE-001",
                customer_name="Correct Customer",
                property_address="Correct manually reviewed address",
                bank_name="Correct Bank",
                branch_name="Correct Branch",
                case_type="Fresh",
                status="Completed",
                extracted_json='{"case_profile":{"reviewed":true}}',
            )
            db.session.add(case)
            db.session.commit()
            apply_email_details(
                case,
                {
                    "application_number": "WRONG-NEW-ID",
                    "customer_name": "Wrong Name",
                    "contact_number": "9876543210",
                    "property_address": "Wrong signature address",
                    "bank_name": "Wrong Bank",
                    "branch_name": "Wrong Branch",
                    "case_type": "Subsequent",
                },
                account,
                "Re: Please share report STABLE-001",
                None,
                "<stable-followup@example.com>",
            )
            db.session.commit()
            self.assertEqual(case.application_number, "STABLE-001")
            self.assertEqual(case.customer_name, "Correct Customer")
            self.assertEqual(case.property_address, "Correct manually reviewed address")
            self.assertEqual(case.branch_name, "Correct Branch")
            self.assertEqual(case.status, "Completed")
            self.assertEqual(case.contact_number, "9876543210")
            self.assertTrue(safe_json(case.extracted_json)["case_profile"]["reviewed"])

    def test_csrf_rejects_missing_token(self):
        self.assertEqual(self.client.post("/login", data={}).status_code, 400)

    def test_saved_engineer_initiates_offline_whatsapp_visit_without_case_link(self):
        self.login()
        response = self.client.post("/site-engineers", data={
            "_csrf_token": self.csrf(), "name": "Nikhil Engineer",
            "mobile_number": "9876543210", "area": "Vidisha",
        })
        self.assertEqual(response.status_code, 302)
        with app.app_context():
            engineer = SiteEngineer.query.filter_by(name="Nikhil Engineer").first()
            case = ValuationCase(
                application_number="VISIT-101", customer_name="Visit Customer",
                contact_number="9999999999", bank_name="Test Bank",
                case_type="Fresh", property_address=(
                    "Plot 12, Village Hasuya, Tehsil Vidisha, Dist- Vidisha"
                ), branch_name="Vidisha",
            )
            db.session.add(case)
            db.session.commit()
            engineer_id, case_id = engineer.id, case.id
        dashboard = self.client.get("/")
        self.assertIn(b'target="_blank"', dashboard.data)
        response = self.client.post(f"/cases/{case_id}/initiate-visit", data={
            "_csrf_token": self.csrf(), "recipient": f"engineer:{engineer_id}",
        })
        self.assertEqual(response.status_code, 302)
        location = unquote(response.headers["Location"])
        self.assertIn("https://wa.me/919876543210?text=", location)
        self.assertIn("Visit Customer", location)
        self.assertIn("Village Hasuya, Tehsil Vidisha, Dist- Vidisha", location)
        self.assertNotIn("photos offline phone camera", location)
        self.assertNotIn("svai-valuation-app.onrender.com", location)
        with app.app_context():
            self.assertEqual(db.session.get(ValuationCase, case_id).visit_by, "Nikhil Engineer")

        empty = self.client.post(f"/cases/{case_id}/initiate-visit", data={
            "_csrf_token": self.csrf(), "recipient": "",
        })
        self.assertEqual(empty.status_code, 302)

        saved_group = self.client.post("/whatsapp-groups", data={
            "_csrf_token": self.csrf(), "name": "Ashta Site Team", "area": "Ashta",
            "invite_url": "https://chat.whatsapp.com/AbCdEf1234567890?s=sw&p=i&ilr=0",
        })
        self.assertEqual(saved_group.status_code, 302)
        with app.app_context():
            group_id = WhatsAppGroup.query.filter_by(name="Ashta Site Team").first().id
        group_page = self.client.post(f"/cases/{case_id}/initiate-visit", data={
            "_csrf_token": self.csrf(), "recipient": f"group:{group_id}",
        })
        self.assertEqual(group_page.status_code, 200)
        self.assertIn(b"Open WhatsApp", group_page.data)
        self.assertIn(b"No Copy/Paste", group_page.data)
        self.assertIn(b"Visit Customer", group_page.data)
        self.assertNotIn(b"photos offline phone camera", group_page.data)
        self.assertEqual(
            normalize_whatsapp_group_link(
                "https://chat.whatsapp.com/CZiTfIiZlEYCkBK8y7pyTa?s=sw&p=i&ilr=0"
            ),
            "https://chat.whatsapp.com/CZiTfIiZlEYCkBK8y7pyTa",
        )

    def test_case_upload_valuation_and_report(self):
        self.login()
        response = self.client.post("/cases/new", data={
            "_csrf_token": self.csrf(),
            "application_number": "SMOKE-001",
            "customer_name": "Smoke Test",
            "bank_name": "Test Bank",
            "property_address": "Vidisha, MP",
        })
        self.assertEqual(response.status_code, 302)
        case_id = int(response.headers["Location"].rstrip("/").split("/")[-1])

        visit_archive = io.BytesIO()
        with zipfile.ZipFile(visit_archive, "w") as bundle:
            bundle.writestr("photos/front.jpg", b"test-image")
            bundle.writestr("../blocked.jpg", b"blocked")
        visit_archive.seek(0)
        response = self.client.post(
            f"/cases/{case_id}/upload/visit",
            data={
                "_csrf_token": self.csrf(),
                "files": (visit_archive, "visit.zip"),
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 302)

        response = self.client.post(
            f"/cases/{case_id}/upload/documents",
            data={
                "_csrf_token": self.csrf(),
                "files": (io.BytesIO(b"scanned-title-document"), "registry.jpg"),
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 302)

        response = self.client.post(f"/cases/{case_id}/valuation", data={
            "_csrf_token": self.csrf(),
            "land_area": "1000",
            "land_rate": "500",
            "builtup_area": "800",
            "construction_rate": "1200",
            "depreciation_percent": "10",
            "conservative_percent": "80",
            "distress_percent": "70",
        })
        self.assertEqual(response.status_code, 302)

        with app.app_context():
            template = FileAsset.query.filter_by(
                asset_type="template", filename="Laxmi India.xlsx"
            ).first()
            self.assertIsNotNone(template)
            template_id = template.id

        report = self.client.post(f"/cases/{case_id}/report", data={
            "_csrf_token": self.csrf(),
            "template_id": str(template_id),
        })
        self.assertEqual(report.status_code, 200)
        self.assertIn(
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            report.content_type,
        )
        self.assertGreater(len(report.data), 1000)
        generated = load_workbook(io.BytesIO(report.data), data_only=False)
        self.assertIn("MOTA RAM", generated.sheetnames)

        with app.app_context():
            self.assertIsNotNone(ValuationCase.query.get(case_id))
            self.assertEqual(FileAsset.query.filter_by(case_id=case_id, asset_type="photo").count(), 1)
            self.assertEqual(FileAsset.query.filter_by(case_id=case_id, asset_type="document").count(), 1)
            self.assertEqual(FileAsset.query.filter_by(case_id=case_id, asset_type="report").count(), 1)

    def test_visit_form_pages_are_separate_and_source_review_is_saved(self):
        self.login()
        response = self.client.post("/cases/new", data={
            "_csrf_token": self.csrf(),
            "application_number": "VISIT-FORM-001",
            "customer_name": "Visit Form Test",
        })
        case_id = int(response.headers["Location"].rstrip("/").split("/")[-1])
        response = self.client.post(
            f"/cases/{case_id}/upload/visit-form",
            data={
                "_csrf_token": self.csrf(),
                "files": [
                    (io.BytesIO(b"handwritten-page-one"), "visit-page-1.jpg"),
                    (io.BytesIO(b"mp-kisan-page"), "mp-kisan-khasra.png"),
                    (io.BytesIO(b"handwritten-page-three"), "visit-page-3.jpeg"),
                ],
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 302)
        with app.app_context():
            assets = FileAsset.query.filter_by(
                case_id=case_id, asset_type="visit_data"
            ).order_by(FileAsset.id).all()
            self.assertEqual(len(assets), 3)
            self.assertEqual(
                [asset.category for asset in assets],
                ["Visit Form Page 1", "Visit Form Page 2", "Visit Form Page 3"],
            )
            self.assertEqual(
                FileAsset.query.filter_by(case_id=case_id, asset_type="photo").count(),
                0,
            )

        response = self.client.post(
            f"/cases/{case_id}/source-review",
            data={
                "_csrf_token": self.csrf(),
                "property_address_as_per_docs": "Registry Address",
                "survey_khasra_plot_no_as_per_docs": "Khasra 151/2",
                "land_area_as_per_docs": "1162 sq ft",
                "north_boundary_as_per_docs": "Document Road",
                "property_address_as_per_site": "Engineer Actual Address",
                "survey_khasra_plot_no_as_per_site": "Khasra 151/3",
                "land_area_as_per_site": "1042 sq ft",
                "north_boundary_as_per_site": "Actual 12 ft Road",
                "road_width": "12 ft",
            },
        )
        self.assertEqual(response.status_code, 302)
        with app.app_context():
            case = db.session.get(ValuationCase, case_id)
            profile = safe_json(case.extracted_json)["case_profile"]
            self.assertEqual(profile["property_address_as_per_docs"], "Registry Address")
            self.assertEqual(profile["property_address_as_per_site"], "Engineer Actual Address")
            self.assertEqual(profile["survey_khasra_plot_no_as_per_docs"], "Khasra 151/2")
            self.assertEqual(profile["survey_khasra_plot_no_as_per_site"], "Khasra 151/3")
            self.assertEqual(profile["survey_khasra_plot_no"], "Khasra 151/2")
            self.assertTrue(profile["source_reviewed"])
            self.assertEqual(case.property_address, "Engineer Actual Address")
            self.assertEqual(case.status, "Source Data Reviewed")

    def test_email_fetch_scans_assignment_folders(self):
        gmail = EmailAccount(email="valuer@gmail.com", encrypted_password="x", provider="gmail")
        yahoo = EmailAccount(email="valuer@yahoo.com", encrypted_password="x", provider="yahoo")
        self.assertIn('"[Gmail]/All Mail"', email_fetch_folders(gmail))
        self.assertIn("INBOX", email_fetch_folders(gmail))
        self.assertEqual(email_fetch_folders(yahoo), ["INBOX", "Archive"])

    def test_existing_mis_import_adds_and_merges_without_duplicate_application(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "ALL BANK"
        sheet.append([
            "SR NO", "Date", "CUSTOMER NAME", "APPLICATION NO", "CONTACT NUMBER",
            "BANK", "CASE TYPE", "STATUS", "ADDRESS", "VISIT BY", "BRANCH", "K.M",
        ])
        sheet.append([
            1, datetime(2026, 8, 7), "Import Customer", "IMPORT-2026-001",
            "9876543210", "Test Bank", "Fresh", "Visit Pending",
            "Plot 1, Vidisha", "Engineer", "Vidisha", 12,
        ])
        stream = io.BytesIO()
        workbook.save(stream)
        stream.seek(0)
        parsed = mis_import_rows(stream)
        self.assertEqual(parsed[0]["application_number"], "IMPORT-2026-001")
        self.assertEqual(parsed[0]["distance"], 12)

        self.login()
        first = self.client.post("/mis/import", data={
            "_csrf_token": self.csrf(),
            "mis_file": (io.BytesIO(stream.getvalue()), "MIS.xlsx"),
        }, content_type="multipart/form-data")
        self.assertEqual(first.status_code, 302)
        second = self.client.post("/mis/import", data={
            "_csrf_token": self.csrf(),
            "mis_file": (io.BytesIO(stream.getvalue()), "MIS.xlsx"),
        }, content_type="multipart/form-data")
        self.assertEqual(second.status_code, 302)
        with app.app_context():
            matches = ValuationCase.query.filter_by(application_number="IMPORT-2026-001").all()
            self.assertEqual(len(matches), 1)
            self.assertEqual(matches[0].customer_name, "Import Customer")

    def test_email_fetch_includes_custom_incoming_folder_but_not_sent_or_junk(self):
        class MailboxList:
            def list(self):
                return "OK", [
                    b'(\\HasNoChildren) "/" "INBOX"',
                    b'(\\HasNoChildren) "/" "LIFC Assignments"',
                    b'(\\Sent) "/" "[Gmail]/Sent Mail"',
                    b'(\\Junk) "/" "Spam"',
                ]

        account = EmailAccount(email="valuer@gmail.com", encrypted_password="x", provider="gmail")
        folders = imap_safe_assignment_folders(account, MailboxList())
        self.assertIn('"LIFC Assignments"', folders)
        self.assertNotIn('"[Gmail]/Sent Mail"', folders)
        self.assertNotIn("Spam", folders)

    def test_mis_fetch_uses_bounded_message_body_without_attachments(self):
        class Mailbox:
            def __init__(self):
                self.query = ""

            def fetch(self, msg_id, query):
                self.query = query
                return "OK", [
                    (b"header", b"Subject: Technical assignment APP-101\r\nFrom: bank@example.com\r\n"),
                    (b"body", b"\r\nCustomer Name: Test Customer"),
                    b")",
                ]

        mailbox = Mailbox()
        raw = fetch_mis_message(mailbox, b"1")
        self.assertIn(b"Technical assignment", raw)
        self.assertIn(b"Test Customer", raw)
        self.assertIn("BODY.PEEK[TEXT]<0.262144>", mailbox.query)
        self.assertNotIn("RFC822", mailbox.query)

    def test_full_email_is_available_only_for_missing_address_fallback(self):
        class Mailbox:
            def fetch(self, msg_id, query):
                self.query = query
                return "OK", [(b"full", b"Subject: Valuation\r\n\r\nBody")]

        mailbox = Mailbox()
        raw = fetch_full_message(mailbox, b"2")
        self.assertIn(b"Valuation", raw)
        self.assertEqual(mailbox.query, "(RFC822)")

    def test_bounded_fetch_falls_back_for_yahoo_style_imap_error(self):
        class YahooMailbox:
            def __init__(self):
                self.queries = []

            def fetch(self, msg_id, query):
                self.queries.append(query)
                if "BODY.PEEK" in query:
                    raise __import__("imaplib").IMAP4.error("FETCH Bad sequence")
                return "OK", [(b"full", b"Subject: Yahoo valuation\r\n\r\nBody")]

        mailbox = YahooMailbox()
        raw = fetch_mis_message(mailbox, b"3")
        self.assertIn(b"Yahoo valuation", raw)
        self.assertEqual(mailbox.queries[-1], "(RFC822)")

    def test_email_prefilter_and_regex_extraction(self):
        subject = "Fresh Technical Valuation - APP NO: LAP-2026-0091"
        body = (
            "Customer Name: Ramesh Kumar\n"
            "Property Address: Plot 14, Vidisha, Madhya Pradesh\n"
            "Contact: 9876543210"
        )
        self.assertTrue(deterministic_email_candidate(subject, body))
        extracted = regex_email_extract(subject, body, "credit@samplebank.com")
        self.assertEqual(extracted["application_number"], "LAP-2026-0091")
        self.assertEqual(extracted["customer_name"], "Ramesh Kumar")
        self.assertTrue(extracted["is_valuation"])
        self.assertFalse(deterministic_email_candidate(
            "Your account statement", "Monthly statement is attached."
        ))
        self.assertTrue(is_followup_email(
            "Re: Fresh Technical Valuation - APP NO: LAP-2026-0091",
            "Dear Team,\nPlease share the technical report.\n"
            "On Friday someone wrote:\nCustomer Name: Ramesh Kumar",
        ))

    def test_lifc_applicant_source_and_concise_mis_address(self):
        extracted = regex_email_extract(
            "LIFC - TECHNICAL Case Assignment | LAPSJP100029929 | Order No - 58429 | SHUJALPUR (MP) Branch",
            "Applicant Name: Ajay Malviya\nVendor Code: TECH073",
            "notifications@lifl.in",
        )
        self.assertEqual(extracted["customer_name"], "Ajay Malviya")
        gmail_case = ValuationCase(source_email="valuer@gmail.com")
        yahoo_case = ValuationCase(source_email="valuer@yahoo.com")
        self.assertEqual(mailbox_source(gmail_case), "Gmail")
        self.assertEqual(mailbox_source(yahoo_case), "Yahoo")
        self.assertEqual(
            concise_mis_address(
                "Plot 12, Ward 4, Village Hasuya, Tehsil Vidisha, Dist- Vidisha"
            ),
            "Village Hasuya, Tehsil Vidisha, Dist- Vidisha",
        )
        self.assertEqual(
            concise_mis_address("House 9, Parvati Puram Colony Vidisha"),
            "Parvati Puram Colony Vidisha",
        )

    def test_public_mail_sender_keeps_strong_new_assignment_only(self):
        self.assertTrue(deterministic_email_candidate(
            "Fresh Technical Valuation - Application No LAPGUN100030646",
            "Applicant: Surendra Singh Yadav\nProperty Address: Village Guna",
            "bank.employee@gmail.com",
        ))
        self.assertFalse(deterministic_email_candidate(
            "Technical discussion",
            "Please review whenever convenient.",
            "somebody@yahoo.com",
        ))
        self.assertFalse(is_followup_email(
            "Subsequent visit required - LAP-2026-0091",
            "Please arrange subsequent visit due to construction stage complete.",
        ))
        self.assertEqual(
            normalized_application_number(" LAP-2026/0091 "),
            "LAP20260091",
        )

        muthoot_subject = "Task TSR - Audit Initiation -GWA-PRO-003294GWALIOR (Sanjay Yadav)"
        self.assertTrue(deterministic_email_candidate(muthoot_subject, "Audit task initiated."))
        muthoot = regex_email_extract(
            muthoot_subject, "Audit task initiated.", "homefinconnect@muthoothomefin.com"
        )
        self.assertEqual(muthoot["application_number"], "GWA-PRO-003294")
        self.assertEqual(muthoot["customer_name"], "Sanjay Yadav")
        self.assertEqual(muthoot["bank_name"], "Muthoot Homefin")

        lifc = regex_email_extract(
            "LIFC - TECHNICAL Case Assignment | LAPAST100029996 | Order No - 57902 | ASHTA (MP) Branch",
            "Applicant Name: Jitendra\nApplication Number: LAPAST100029996",
            "notifications@lifl.in",
        )
        self.assertTrue(lifc["is_valuation"])
        self.assertEqual(lifc["bank_name"], "Laxmi India Finance")

    def test_duplicate_fresh_assignment_merges_but_subsequent_stays_separate(self):
        with app.app_context():
            account = EmailAccount(
                email="duplicates@example.com",
                encrypted_password=encrypt_password("abcdefghijklmnop"),
                provider="gmail",
                active=True,
            )
            fresh = ValuationCase(
                application_number="GWA-PRO-003283",
                customer_name="Ginny Sabharwal",
                source_email=account.email,
                case_type="Fresh",
            )
            db.session.add_all([account, fresh])
            db.session.commit()
            self.assertEqual(
                existing_case_for_duplicate_assignment({
                    "application_number": "GWA PRO 003283", "case_type": "Fresh",
                }, account).id,
                fresh.id,
            )
            self.assertIsNone(existing_case_for_duplicate_assignment({
                "application_number": "GWA-PRO-003283", "case_type": "Subsequent",
            }, account))

    def test_same_special_assignment_from_two_mailboxes_merges(self):
        received = __import__("datetime").datetime(2026, 8, 8, 10, 30)
        with app.app_context():
            gmail = ValuationCase(
                application_number="PART-101", case_type="Part / Tranche",
                source_email="one@gmail.com", email_subject="Part valuation PART-101",
                email_received_at=received,
            )
            yahoo = EmailAccount(
                email="two@yahoo.com", encrypted_password="x", provider="yahoo",
            )
            db.session.add_all([gmail, yahoo])
            db.session.commit()
            match = existing_case_for_duplicate_assignment(
                {"application_number": "PART 101", "case_type": "Part / Tranche"},
                yahoo, "Part valuation PART-101", received,
            )
            self.assertEqual(match.id, gmail.id)
            same_mailbox = EmailAccount(
                email="one@gmail.com", encrypted_password="x", provider="gmail",
            )
            db.session.add(same_mailbox)
            db.session.commit()
            same_match = existing_case_for_duplicate_assignment(
                {"application_number": "PART-101", "case_type": "Part / Tranche"},
                same_mailbox, "Part valuation PART-101", received,
            )
            self.assertEqual(same_match.id, gmail.id)

    def test_case_delete_removes_case_and_related_assets(self):
        with app.app_context():
            case = ValuationCase(application_number="DELETE-101", customer_name="Remove Me")
            db.session.add(case)
            db.session.commit()
            case_id = case.id
            db.session.add(FileAsset(
                case_id=case_id, asset_type="document", filename="remove.pdf", content=b"x",
            ))
            db.session.commit()
        self.login()
        response = self.client.post(
            f"/cases/{case_id}/delete",
            data={"_csrf_token": self.csrf()},
        )
        self.assertEqual(response.status_code, 302)
        with app.app_context():
            self.assertIsNone(db.session.get(ValuationCase, case_id))
            self.assertEqual(FileAsset.query.filter_by(case_id=case_id).count(), 0)

    def test_email_document_cleanup_preserves_mis_and_manual_case_files(self):
        with app.app_context():
            email_case = ValuationCase(
                application_number="EMAIL-DOC-1", source_email="valuer@gmail.com",
            )
            manual_case = ValuationCase(application_number="MANUAL-DOC-1")
            db.session.add_all([email_case, manual_case])
            db.session.commit()
            db.session.add_all([
                FileAsset(
                    case_id=email_case.id, asset_type="document",
                    filename="email.pdf", content=b"email-document",
                ),
                FileAsset(
                    case_id=email_case.id, asset_type="report",
                    filename="report.pdf", content=b"report",
                ),
                FileAsset(
                    case_id=manual_case.id, asset_type="document",
                    filename="manual.pdf", content=b"manual-document",
                ),
            ])
            db.session.commit()
            email_case_id = email_case.id
            manual_case_id = manual_case.id
        self.login()
        response = self.client.post(
            "/settings/cleanup-email-documents",
            data={"_csrf_token": self.csrf()},
        )
        self.assertEqual(response.status_code, 302)
        with app.app_context():
            self.assertIsNotNone(db.session.get(ValuationCase, email_case_id))
            self.assertEqual(FileAsset.query.filter_by(
                case_id=email_case_id, asset_type="document"
            ).count(), 0)
            self.assertEqual(FileAsset.query.filter_by(
                case_id=email_case_id, asset_type="report"
            ).count(), 1)
            self.assertEqual(FileAsset.query.filter_by(
                case_id=manual_case_id, asset_type="document"
            ).count(), 1)

    def test_cross_mailbox_same_application_merges_to_one_active_mis_case(self):
        with app.app_context():
            gmail_case = ValuationCase(
                application_number="SAME-CASE-101", customer_name="Customer One",
                bank_name="Test Bank", case_type="Fresh", source_email="gmail@example.com",
                source_message_id="<gmail-copy@example.com>",
            )
            yahoo_case = ValuationCase(
                application_number="SAME CASE 101", contact_number="9999999999",
                bank_name="Test Bank", case_type="Fresh", source_email="yahoo@example.com",
                source_message_id="<yahoo-copy@example.com>",
            )
            subsequent = ValuationCase(
                application_number="SAME-CASE-101", bank_name="Test Bank",
                case_type="Subsequent", source_email="yahoo@example.com",
            )
            received = __import__("datetime").datetime(2026, 8, 7, 20, 0)
            part_one = ValuationCase(
                application_number="PART-REPEAT-101", bank_name="Test Bank",
                case_type="Part / Tranche", email_subject="Part valuation request",
                email_received_at=received,
            )
            part_two = ValuationCase(
                application_number="PART REPEAT 101", bank_name="Test Bank",
                case_type="Part / Tranche", email_subject="FW: Tranche release",
                email_received_at=received.replace(minute=15),
            )
            db.session.add_all([gmail_case, yahoo_case, subsequent, part_one, part_two])
            db.session.commit()
            self.assertEqual(merge_cross_mailbox_duplicate_cases(), 2)
            self.assertFalse(gmail_case.archived)
            self.assertTrue(yahoo_case.archived)
            self.assertEqual(gmail_case.contact_number, "9999999999")
            self.assertFalse(subsequent.archived)
            self.assertFalse(part_one.archived)
            self.assertTrue(part_two.archived)

    def test_followup_report_request_updates_existing_case_without_new_mis_row(self):
        with app.app_context():
            case = ValuationCase(
                application_number="FOLLOWUP-2026-001",
                customer_name="Existing Customer",
                case_type="Fresh",
                source_email="valuer@example.com",
                source_message_id="<initiation@example.com>",
                email_subject="Technical initiation FOLLOWUP-2026-001",
            )
            db.session.add(case)
            db.session.commit()
            count_before = ValuationCase.query.count()
            changed = apply_followup_to_existing_case(
                case,
                {
                    "application_number": "FOLLOWUP-2026-001",
                    "contact_number": "9876543210",
                },
                [],
                "Re: Technical initiation FOLLOWUP-2026-001",
                None,
                "<followup@example.com>",
            )
            self.assertTrue(changed)
            self.assertEqual(ValuationCase.query.count(), count_before)
            self.assertEqual(case.contact_number, "9876543210")
            self.assertEqual(
                len(safe_json(case.extracted_json)["followup_emails"]), 1
            )
            self.assertFalse(apply_followup_to_existing_case(
                case, {}, [], "Repeated fetch", None, "<followup@example.com>"
            ))

    def test_real_bank_subject_patterns_and_signature_are_parsed_safely(self):
        bajaj = regex_email_extract(
            "Re: Technical INITIATION // H425HLD1885354 // AMAN AGRAWAL // "
            "BHOPAL // RESALE",
            "Applicant Contact: 7000700463",
            "credit@bajajhousing.co.in",
        )
        self.assertEqual(bajaj["application_number"], "H425HLD1885354")
        self.assertEqual(bajaj["customer_name"], "AMAN AGRAWAL")
        self.assertEqual(bajaj["branch_name"], "Bhopal")
        self.assertEqual(bajaj["case_type"], "Resale")
        self.assertEqual(bajaj["bank_name"], "Bajaj Housing Finance")

        ummeed = regex_email_extract(
            "Task Technical Initiation -00180296 JHANSI (PREETI KEVAT)",
            "",
            "technical@ummeedhfc.com",
        )
        self.assertEqual(ummeed["application_number"], "00180296")
        self.assertEqual(ummeed["customer_name"], "PREETI KEVAT")
        self.assertEqual(ummeed["branch_name"], "Jhansi")
        self.assertNotEqual(ummeed.get("application_number"), "application")

        lifc = regex_email_extract(
            "RE: REQUEST FOR POSITIVE VALUATION REPORT CASE /MAHENDRA KIRAR/"
            "LAPVDS100026755/VIDISHA BRANCH",
            "Contact: 9752725420\n\nRegards,\nLaxmi India Finance Ltd.\n"
            "Address: Kartra Arcade, Raisen Road Bhopal\n"
            "Email: ravi.trivedi@lifc.in\nWebsite: www.lifc.co.in\n"
            "Toll Free Number: 1800-121-7747",
            "ravi.trivedi@lifc.in",
        )
        self.assertEqual(lifc["application_number"], "LAPVDS100026755")
        self.assertEqual(lifc["customer_name"], "MAHENDRA KIRAR")
        self.assertEqual(lifc["branch_name"], "Vidisha")
        self.assertEqual(lifc.get("property_address", ""), "")

        lifc_assignment = regex_email_extract(
            "LIFC - TECHNICAL Case Assignment | LAPAST100029996 | Order No - 57902 | ASHTA (MP) Branch",
            "Applicant Name: Jitendra\nVendor Code: TECH073\n"
            "Vendor Dashboard: https://login.synofin.tech/?app_id=vendor_portal",
            "notifications@lifl.in",
        )
        self.assertTrue(lifc_assignment["is_valuation"])
        self.assertEqual(lifc_assignment["bank_name"], "Laxmi India Finance")
        self.assertEqual(lifc_assignment["application_number"], "LAPAST100029996")
        self.assertEqual(lifc_assignment["customer_name"], "Jitendra")
        self.assertEqual(lifc_assignment["branch_name"], "Ashta (Mp)")

        public_mail_trail = regex_email_extract(
            "Technical case assignment - TENOR-20YRS",
            "Applicant: SATYAPAL SINGH JADON\nCase type: Purchase + Construction",
            "somebody@gmail.com",
        )
        self.assertFalse(public_mail_trail["is_valuation"])
        self.assertEqual(public_mail_trail["bank_name"], "")

        forwarded_structured_assignment = regex_email_extract(
            "Technical case assignment - AYE-000003928097-LOS",
            "Application No: AYE-000003928097-LOS\n"
            "Applicant: JAGDISH\nCase Type: LAP\n"
            "Property Address: Village Mendori, Bhopal\n"
            "Branch Name: Karond",
            "bank.team@gmail.com",
        )
        self.assertTrue(forwarded_structured_assignment["is_valuation"])
        self.assertEqual(
            forwarded_structured_assignment["application_number"],
            "AYE-000003928097-LOS",
        )

    def test_refetch_recovers_a_previously_rejected_assignment(self):
        with app.app_context():
            account = EmailAccount.query.first()
            case = ValuationCase(
                application_number="AYE-000003928097-LOS",
                customer_name="JAGDISH",
                status="Ignored - Not Valuation Email",
                archived=True,
                source_email=account.email,
                source_message_id="<rejected-assignment@example.test>",
                email_subject="Technical case assignment",
            )
            db.session.add(case)
            db.session.commit()
            apply_email_details(
                case,
                {
                    "is_valuation": True,
                    "application_number": "AYE-000003928097-LOS",
                    "customer_name": "JAGDISH",
                    "case_type": "LAP",
                    "property_address": "Village Mendori, Bhopal",
                },
                account,
                case.email_subject,
                datetime(2026, 8, 10, 6, 5),
                case.source_message_id,
            )
            db.session.commit()
            self.assertFalse(case.archived)
            self.assertIn(case.status, {"New - Email", "Email Parsed - Review"})

    def test_attachment_fills_missing_real_mis_fields(self):
        details = {
            "is_valuation": True,
            "application_number": "",
            "customer_name": "",
            "contact_number": "",
            "branch_name": "",
            "case_type": "",
            "property_address": "",
        }
        text = (
            "TECHNICAL SCRUTINY REPORT FOR Mr. ANAND SAINI\n"
            "Branch Name Sagar Lead Id No HLSA0004F0E1 Report Date 26-06-2023\n"
            "Case Type Plot Purchase House Delivery Agency Self-Construction\n"
            "Contact person name And Number Jayanti Saini, 9179275115\n"
            "Address as per Provided Documents Mouja-Sagarkhas, Brandavan ward, "
            "old khasra 491/1, Tehsil and Distt. Sagar 470002\n"
            "Address as per Plan Not provided\n"
        )
        enriched = enrich_email_details_from_attachments(
            details, [("technical_report.pdf", text)]
        )
        self.assertEqual(enriched["application_number"], "HLSA0004F0E1")
        self.assertEqual(enriched["customer_name"], "ANAND SAINI")
        self.assertEqual(enriched["contact_number"], "9179275115")
        self.assertEqual(enriched["branch_name"], "Sagar")
        self.assertEqual(enriched["case_type"], "Plot Purchase")
        self.assertIn("Sagarkhas", enriched["property_address"])
        self.assertFalse(deterministic_email_candidate(
            "Mukesh estimate floor plan and key plan",
            "Attached construction estimate and drawing.",
            "sender@ummeedhfc.com",
        ))

    def test_additional_real_initiation_and_portal_patterns(self):
        fusion = regex_email_extract(
            "Valuation required -//NEELESH // Application Id - 465486// "
            "Vidisha Branch",
            "Application Id\n465486\nBranch\nVidisha\nApplicant\nNEELESH\n"
            "Contact Number of Applicant\n9685063376\n"
            "Address of property to be mortgaged with pin code\n"
            "Narwar, Raisen, Madhya Pradesh 464551",
            "arvind.verma@fusionfin.com",
        )
        self.assertEqual(fusion["application_number"], "465486")
        self.assertEqual(fusion["customer_name"], "NEELESH")
        self.assertEqual(fusion["branch_name"], "Vidisha")
        self.assertEqual(fusion["case_type"], "Fresh")
        self.assertEqual(fusion["bank_name"], "Fusion Finance")

        subsequent = regex_email_extract(
            "Subsequent visit required in case of Indralal Vishwakarma",
            "Please arrange subsequent visit in captioned case.",
            "rajat.deshmukh@grihashakti.com",
        )
        self.assertTrue(subsequent["is_valuation"])
        self.assertEqual(subsequent["customer_name"], "Indralal Vishwakarma")
        self.assertEqual(subsequent["case_type"], "Subsequent")

        revisit = regex_email_extract(
            "INITIATE TECHNICAL FOR THE CASE NAME Mr. ROOPNARAYAN SHARMA "
            "Lan No: -HVDS26000116402",
            "Kindly arrange re-visit due to construction stage complete.",
            "pratham.gour@jmfl.com",
        )
        self.assertEqual(revisit["customer_name"], "ROOPNARAYAN SHARMA")
        self.assertEqual(revisit["application_number"], "HVDS26000116402")
        self.assertEqual(revisit["case_type"], "Revisit")

        portal = regex_email_extract(
            "LIFC - TECHNICAL Case Assignment | LAPSRJ100027125 | "
            "Order No - 56750 | Sironj(MP) Branch",
            "Application Number: LAPSRJ100027125\n"
            "Applicant Name: Jagdish Sen\nVendor Dashboard: https://example.test/portal",
            "notifications@lifc.in",
        )
        self.assertEqual(portal["application_number"], "LAPSRJ100027125")
        self.assertEqual(portal["customer_name"], "Jagdish Sen")
        self.assertEqual(portal["branch_name"], "Sironj(Mp)")
        self.assertTrue(portal["portal_case"])
        lifc_guna = regex_email_extract(
            "LIFC - TECHNICAL Case Assignment | CLGUN100028420 | "
            "Order No - 56727 | GUNA(MP) Branch",
            "Applicant Name: Test Applicant",
            "notifications@lifc.in",
        )
        self.assertEqual(lifc_guna["application_number"], "CLGUN100028420")
        business = regex_email_extract(
            "Technical Request - CRYSTAL CROWN HERBAL - "
            "LUGL871520260706111432",
            "",
            "technical@examplebank.com",
        )
        self.assertEqual(
            business["application_number"], "LUGL871520260706111432"
        )
        self.assertEqual(business["customer_name"], "CRYSTAL CROWN HERBAL")

    def test_application_number_quick_start_and_draft_report(self):
        self.login()
        response = self.client.post("/cases/new", data={
            "_csrf_token": self.csrf(),
            "application_number": "QUICK-2026-001",
        })
        self.assertEqual(response.status_code, 302)
        case_id = int(response.headers["Location"].rstrip("/").split("/")[-1])
        response = self.client.post(
            f"/cases/{case_id}/upload/all",
            data={
                "_csrf_token": self.csrf(),
                "files": (io.BytesIO(b"site-photo"), "front.jpg"),
                "template": (
                    io.BytesIO(
                        (
                            Path(__file__).resolve().parents[1]
                            / "seed_templates" / "SBFC.xlsx"
                        ).read_bytes()
                    ),
                    "SBFC.xlsx",
                ),
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 302)
        report = self.client.post(f"/cases/{case_id}/report", data={
            "_csrf_token": self.csrf(),
            "template_id": "",
        })
        self.assertEqual(report.status_code, 200)
        with app.app_context():
            case = db.session.get(ValuationCase, case_id)
            self.assertEqual(case.status, "Draft Report Generated")
            self.assertEqual(
                FileAsset.query.filter_by(case_id=case_id, asset_type="photo").count(),
                1,
            )

    def test_unknown_uploaded_excel_keeps_structure_and_fills_in_place(self):
        source = Workbook()
        sheet = source.active
        sheet.title = "BANK ORIGINAL"
        sheet.merge_cells("A1:D1")
        sheet["A1"] = "ORIGINAL BANK VALUATION FORMAT"
        sheet["B3"] = "Customer Name"
        sheet["B5"] = "Front Elevation Photo"
        sheet["D10"] = "=1+1"
        sheet.column_dimensions["B"].width = 28
        sheet.row_dimensions[6].height = 120
        source_stream = io.BytesIO()
        source.save(source_stream)

        image_stream = io.BytesIO()
        Image.new("RGB", (200, 120), "#b7d7ee").save(image_stream, format="JPEG")
        output = fill_excel_template(
            source_stream.getvalue(),
            {"customer_name": "Exact Format Customer"},
            [{
                "filename": "front.jpg",
                "category": "Front Elevation",
                "content": image_stream.getvalue(),
            }],
            "Other Bank Original.xlsx",
            "Other Bank",
        )
        filled = load_workbook(io.BytesIO(output), data_only=False)
        result = filled["BANK ORIGINAL"]
        self.assertEqual(filled.sheetnames, ["BANK ORIGINAL"])
        self.assertEqual([str(item) for item in result.merged_cells.ranges], ["A1:D1"])
        self.assertEqual(result.column_dimensions["B"].width, 28)
        self.assertEqual(result.row_dimensions[6].height, 120)
        self.assertEqual(result["C3"].value, "Exact Format Customer")
        self.assertEqual(result["D10"].value, "=1+1")
        self.assertEqual(len(result._images), 1)

    def test_filtered_mis_uses_uploaded_format_columns(self):
        self.login()
        with app.app_context():
            inside = ValuationCase(
                application_number="MIS-IN-001",
                customer_name="MIS Included",
                bank_name="DCB",
                case_type="Fresh",
                created_at=__import__("datetime").datetime(2025, 1, 15, 10, 30),
            )
            outside = ValuationCase(
                application_number="MIS-OUT-001",
                customer_name="MIS Excluded",
                bank_name="SBFC",
                case_type="Subsequent",
                created_at=__import__("datetime").datetime(2025, 2, 1, 10, 30),
            )
            db.session.add_all([inside, outside])
            db.session.commit()

        response = self.client.get("/mis/export?from=2025-01-01&to=2025-01-31")
        self.assertEqual(response.status_code, 200)
        workbook = load_workbook(io.BytesIO(response.data), data_only=False)
        sheet = workbook["ALL BANK"]
        expected = [
            "SR NO", "Date", "Time", "CUSTOMER NAME", "APPLICATION NO", "CONTACT NUMBER",
            "CASE TYPE", "BANK", "STATUS", "ADDRESS", "VISIT BY", "BRANCH",
            "Pending", "K.M",
        ]
        self.assertEqual([sheet.cell(1, col).value for col in range(1, 15)], expected)
        application_numbers = [
            sheet.cell(row, 5).value for row in range(2, sheet.max_row + 1)
        ]
        self.assertIn("MIS-IN-001", application_numbers)
        self.assertNotIn("MIS-OUT-001", application_numbers)

    def test_bank_formats_keep_document_and_site_facts_separate(self):
        root = Path(__file__).resolve().parents[1]
        image_buffer = io.BytesIO()
        Image.new("RGB", (160, 100), "#d9e8f5").save(image_buffer, format="JPEG")
        photos = [{
            "filename": "front.jpg",
            "category": "Front Elevation",
            "content": image_buffer.getvalue(),
        }]
        profile = {
            "application_number": "MAP-001",
            "customer_name": "Mapped Customer",
            "property_address_as_per_docs": "Document Address",
            "property_address_as_per_site": "Actual Site Address",
            "survey_khasra_plot_no_as_per_docs": "Document Khasra 151/2",
            "survey_khasra_plot_no_as_per_site": "Actual Khasra 151/3",
            "north_boundary_as_per_docs": "Document North",
            "north_boundary_as_per_site": "Actual North",
            "land_area_as_per_docs": "1,500 sq ft",
            "land_area_as_per_site": "1,488 sq ft",
            "builtup_area_as_per_site": "920 sq ft",
            "land_rate": 850,
            "construction_rate": 1400,
        }
        sbfc_source = (root / "seed_templates" / "SBFC.xlsx").read_bytes()
        sbfc_output = fill_excel_template(
            sbfc_source, profile, [], "SBFC.xlsx", "SBFC Finance"
        )
        sbfc = load_workbook(io.BytesIO(sbfc_output), data_only=False)["Table 1"]
        self.assertEqual(sbfc["B11"].value, "Actual Site Address")
        self.assertEqual(sbfc["B12"].value, "Document Address")
        self.assertEqual(sbfc["F13"].value, "Document Khasra 151/2")
        self.assertEqual(sbfc["D15"].value, "Document Khasra 151/2")
        self.assertEqual(sbfc["B36"].value, "Document North")
        self.assertEqual(sbfc["B37"].value, "Actual North")
        self.assertEqual(sbfc["E68"].value, "=C68*D68")
        self.assertEqual(sbfc["D70"].value, "=E68+E69")

        laxmi_source = (root / "seed_templates" / "Laxmi India.xlsx").read_bytes()
        laxmi_output = fill_excel_template(
            laxmi_source, profile, photos, "Laxmi India.xlsx", "Laxmi India Finance"
        )
        laxmi = load_workbook(io.BytesIO(laxmi_output), data_only=False)["MOTA RAM"]
        self.assertEqual(laxmi["C15"].value, "Document Address")
        self.assertEqual(laxmi["C16"].value, "Actual Site Address")
        self.assertEqual(laxmi["G38"].value, 1488)
        self.assertEqual(laxmi["D102"].value, "=G42")
        self.assertEqual(laxmi["I102"].value, "=D102*E102")
        self.assertEqual(len(laxmi._images), 1)
        self.assertEqual(laxmi._images[0].anchor._from.row + 1, 161)
        self.assertEqual(laxmi._images[0].anchor._from.col + 1, 1)

        no_site_profile = dict(profile)
        no_site_profile["property_address_as_per_site"] = ""
        no_site_output = fill_excel_template(
            laxmi_source,
            no_site_profile,
            [],
            "Laxmi India.xlsx",
            "Laxmi India Finance",
        )
        no_site = load_workbook(
            io.BytesIO(no_site_output), data_only=False
        )["MOTA RAM"]
        self.assertIn(no_site["C16"].value, ("", None))

        ummeed_source = (root / "seed_templates" / "Ummeed.docx").read_bytes()
        ummeed_output = fill_docx_template(ummeed_source, profile, photos)
        ummeed = Document(io.BytesIO(ummeed_output))
        self.assertEqual(len(ummeed.tables), 6)
        self.assertIn("MAP-001", ummeed.tables[0].rows[1].cells[-1].text)
        self.assertEqual(len(ummeed.inline_shapes), 1)

    def test_valuer_standing_defaults_use_document_and_site_authority(self):
        defaults = valuation_defaults_from_profile({
            "land_area_as_per_docs": "1,250 sq ft",
            "land_area_as_per_site": "1,180 sq ft",
            "builtup_area_as_per_site": "780 sq ft",
            "govt_land_rate": "250",
            "construction_stage": "Without Plaster / Brickwork",
            "construction_year": "2020",
        })
        self.assertEqual(defaults["land_area"], 1250)
        self.assertEqual(defaults["builtup_area"], 780)
        self.assertEqual(defaults["land_rate"], 500)
        self.assertEqual(defaults["construction_rate"], 700)
        self.assertEqual(defaults["age_years"], 6)
        self.assertEqual(defaults["depreciation_percent"], 0)
        self.assertEqual(defaults["conservative_percent"], 70)
        self.assertEqual(defaults["distress_percent"], 80)

    def test_laxmi_generic_whatsapp_photos_fill_available_slots(self):
        root = Path(__file__).resolve().parents[1]
        image_buffer = io.BytesIO()
        Image.new("RGB", (160, 100), "#d9e8f5").save(image_buffer, format="JPEG")
        photos = [
            {
                "filename": f"WhatsApp_Image_{index}.jpeg",
                "category": "Other Site Photo",
                "content": image_buffer.getvalue(),
            }
            for index in range(6)
        ]
        output = fill_excel_template(
            (root / "seed_templates" / "Laxmi India.xlsx").read_bytes(),
            {"application_number": "PHOTO-001"},
            photos,
            "Laxmi India.xlsx",
            "Laxmi India Finance",
        )
        sheet = load_workbook(io.BytesIO(output), data_only=False)["MOTA RAM"]
        self.assertEqual(len(sheet._images), 6)

    @classmethod
    def tearDownClass(cls):
        with app.app_context():
            db.session.remove()
            db.drop_all()
            db.engine.dispose()
        TEST_DIR.cleanup()


if __name__ == "__main__":
    unittest.main(verbosity=2)

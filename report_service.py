"""Preserve-and-fill services for bank Excel and Word valuation formats."""

import io
import json
import re
from collections import defaultdict
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook, load_workbook
from openpyxl.drawing.image import Image as ExcelImage
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from PIL import Image as PillowImage, ImageOps

from ai_service_openai import map_template_cells


TOKEN_ALIASES = {
    "APPLICATION_NUMBER": ("application_number", "deal_number"),
    "CUSTOMER_NAME": ("customer_name", "applicant_name"),
    "OWNER_NAME": ("owner_name",),
    "CONTACT_NUMBER": ("contact_number",),
    "PROPERTY_ADDRESS": ("property_address_as_per_site", "property_address", "property_address_as_per_docs"),
    "PROPERTY_ADDRESS_DOCS": ("property_address_as_per_docs", "property_address"),
    "PROPERTY_ADDRESS_SITE": ("property_address_as_per_site",),
    "SURVEY_KHASRA_PLOT_NO": ("survey_khasra_plot_no_as_per_docs", "survey_khasra_plot_no"),
    "SURVEY_KHASRA_PLOT_NO_DOCS": ("survey_khasra_plot_no_as_per_docs",),
    "SURVEY_KHASRA_PLOT_NO_SITE": ("survey_khasra_plot_no_as_per_site",),
    "BANK_NAME": ("bank_name",),
    "BRANCH_NAME": ("branch_name",),
    "CASE_TYPE": ("case_type",),
    "REPORT_DATE": ("report_date",),
    "VISIT_DATE": ("visit_date",),
    "VISIT_BY": ("visit_engineer", "visit_by"),
    "LAND_AREA": ("land_area", "land_area_as_per_docs", "land_area_as_per_site"),
    "LAND_AREA_DOCS": ("land_area_as_per_docs", "land_area"),
    "LAND_AREA_SITE": ("land_area_as_per_site", "land_area"),
    "LAND_RATE": ("land_rate",),
    "LAND_VALUE": ("land_value",),
    "BUILTUP_AREA": ("builtup_area", "builtup_area_as_per_site", "builtup_area_as_per_docs"),
    "BUILTUP_AREA_DOCS": ("builtup_area_as_per_docs",),
    "BUILTUP_AREA_SITE": ("builtup_area_as_per_site", "builtup_area"),
    "CONSTRUCTION_RATE": ("construction_rate",),
    "MARKET_VALUE": ("market_value",),
    "CONSERVATIVE_VALUE": ("conservative_value",),
    "DISTRESS_VALUE": ("distress_value",),
    "GOVT_VALUE": ("govt_value",),
    "REMARKS": ("remarks",),
    "LATITUDE": ("latitude",),
    "LONGITUDE": ("longitude",),
}

LABEL_ALIASES = (
    (("deal number", "case no", "appl no", "application no", "lead purposal number"), "application_number"),
    (("name of the customer", "name of customer", "customer name", "applicant name", "name of the borrower"), "customer_name"),
    (("name of the property owner", "property owner", "owner name", "ownership as per"), "owner_name"),
    (("mobile no", "contact no", "contact number"), "contact_number"),
    (("address as per site", "property address as per site"), "property_address_as_per_site"),
    (("address as per document", "legal address", "property address as per doc"), "property_address_as_per_docs"),
    (("actual khasra", "khasra no as per site", "survey no as per site", "plot no as per site"), "survey_khasra_plot_no_as_per_site"),
    (("khasra no as per document", "survey no as per document", "plot no as per document"), "survey_khasra_plot_no_as_per_docs"),
    (("branch name", "name of nearest branch"), "branch_name"),
    (("type of loan", "product", "case type"), "case_type"),
    (("date of technical visit", "technical visit date", "date of inspection"), "visit_date"),
    (("technical visit done by", "property was inspected by"), "visit_engineer"),
    (("latitude",), "latitude"),
    (("longitude", "longtitude"), "longitude"),
    (("plot no", "survey no", "khasra no"), "survey_khasra_plot_no_as_per_docs"),
    (("plot area", "land area as per sale deed", "land area as per document"), "land_area_as_per_docs"),
    (("actual builtup", "actual built up", "built up area"), "builtup_area_as_per_site"),
    (("age of the property",), "property_age_years"),
    (("residual age", "estimated future life"), "residual_age_years"),
    (("type of construction", "structure type"), "structure_type"),
    (("actual use", "property usage"), "property_usage_as_per_site"),
    (("occupancy status", "occupancy"), "occupancy"),
    (("person met at site",), "person_met"),
    (("width of the approach road", "front side road width", "road width"), "road_width"),
    (("present condition", "quality of construction"), "construction_quality"),
    (("marketability",), "marketability"),
    (("land rate", "rate adopted"), "land_rate"),
    (("construction rate", "rate/sft"), "construction_rate"),
)


# Exact data-entry cells for the supplied bank formats. These maps deliberately
# avoid headings, declarations and formula cells. Every mapped field is replaced
# even when the case has no value, so sample applicant/property data from the
# reference workbook cannot leak into a newly generated report.
KNOWN_CELL_MAPPINGS = {
    "laxmi": {
        "D8": ("visit_date", "report_date"),
        "D9": ("application_number",),
        "D10": ("customer_name", "applicant_name"),
        "D11": ("owner_name",),
        "D12": ("contact_number",),
        "I9": ("case_type",),
        "I13": ("occupancy",),
        "I14": ("branch_name",),
        "C15": ("property_address_as_per_docs", "property_address"),
        "C16": ("property_address_as_per_site",),
        "G38": ("land_area_as_per_site", "land_area"),
        "G42": ("builtup_area_as_per_site", "builtup_area"),
        "G44": ("property_usage_as_per_site",),
        "G47": ("property_age_years", "age_years"),
        "G48": ("residual_age_years",),
        "G49": ("occupancy",),
        "G51": ("plot_demarcated",),
        "F32": ("road_type",),
        "F33": ("landmark",),
        "F35": ("amenities",),
        "D55": ("north_boundary_as_per_docs",),
        "F55": ("east_boundary_as_per_docs",),
        "H55": ("west_boundary_as_per_docs",),
        "J55": ("south_boundary_as_per_docs",),
        "D56": ("north_boundary_as_per_site",),
        "F56": ("east_boundary_as_per_site",),
        "H56": ("west_boundary_as_per_site",),
        "J56": ("south_boundary_as_per_site",),
        "D59": ("north_boundary_as_per_docs",),
        "F59": ("east_boundary_as_per_docs",),
        "H59": ("west_boundary_as_per_docs",),
        "J59": ("south_boundary_as_per_docs",),
        "D60": ("north_boundary_as_per_site",),
        "F60": ("east_boundary_as_per_site",),
        "H60": ("west_boundary_as_per_site",),
        "J60": ("south_boundary_as_per_site",),
        "G62": ("marketability",),
        "F67": ("construction_quality",),
        "G73": ("structure_type",),
        "G74": ("number_of_floors",),
        "G77": ("room_configuration",),
        "G79": ("plan_details",),
        "G80": ("plan_details",),
        "G81": ("construction_permission",),
        "G82": ("demolition_risk",),
        "E89": ("govt_land_rate",),
        "E91": ("govt_construction_rate",),
        "E100": ("land_rate",),
        "D102": ("builtup_area_as_per_site", "builtup_area"),
        "E102": ("construction_rate",),
    },
    "dcb": {
        "N3": ("report_date",),
        "N5": ("visit_date",),
        "N6": ("bank_name",),
        "N7": ("locality_type", "village"),
        "N8": ("application_number",),
        "N9": ("case_type",),
        "N10": ("project_name",),
        "N11": ("customer_name", "applicant_name"),
        "N12": ("co_applicant_name",),
        "N13": ("property_address_as_per_site", "property_address", "property_address_as_per_docs"),
        "N14": ("pincode",),
        "N15": ("owner_name",),
        "N16": ("document_type",),
        "N17": ("property_type",),
        "N18": ("property_type",),
        "N19": ("road_type",),
        "N20": ("survey_khasra_plot_no_as_per_docs", "survey_khasra_plot_no"),
        "N21": ("locality_type",),
        "N24": ("land_tenure",),
        "N25": ("approving_authority",),
        "N26": ("plan_details",),
        "W28": ("north_boundary_as_per_docs",),
        "W29": ("south_boundary_as_per_docs",),
        "W30": ("east_boundary_as_per_docs",),
        "W31": ("west_boundary_as_per_docs",),
        "W32": ("north_boundary_as_per_site",),
        "W33": ("south_boundary_as_per_site",),
        "W34": ("east_boundary_as_per_site",),
        "W35": ("west_boundary_as_per_site",),
        "N37": ("demolition_risk",),
        "N38": ("property_identified_through",),
        "N39": ("property_identified_through",),
        "N41": ("plot_demarcated",),
        "N42": ("coordinates",),
        "M44": ("construction_year",),
        "M45": ("property_age_years", "age_years"),
        "M46": ("residual_age_years",),
        "M47": ("construction_quality",),
        "M49": ("structure_type",),
        "M50": ("property_usage_as_per_docs",),
        "M51": ("property_usage_as_per_site",),
        "M52": ("number_of_floors",),
        "M58": ("amenities",),
        "M60": ("number_of_floors",),
        "M62": ("floor_wise_usage", "room_configuration"),
        "M70": ("doors_windows",),
        "M71": ("flooring",),
        "M77": ("demolition_risk",),
        "M80": ("occupancy",),
        "M85": ("land_area_as_per_site", "land_area"),
        "M86": ("property_age_years", "age_years"),
        "M87": ("residual_age_years",),
        "M88": ("construction_quality",),
        "M90": ("land_rate",),
        "J101": ("builtup_area_as_per_site", "builtup_area"),
        "J103": ("land_area_as_per_docs", "land_area"),
        "C108": ("land_area_as_per_docs", "land_area"),
        "I108": ("land_rate",),
        "C110": ("builtup_area_as_per_site", "builtup_area"),
        "I110": ("construction_rate",),
        "E148": ("land_area_as_per_docs", "land_area"),
        "K148": ("land_rate",),
    },
    "sbfc": {
        "B8": ("branch_name",),
        "D8": ("customer_name", "applicant_name"),
        "F8": ("application_number",),
        "B9": ("case_type",),
        "D9": ("contact_number",),
        "B11": ("property_address_as_per_site",),
        "F11": ("address_match_status",),
        "B12": ("property_address_as_per_docs", "property_address"),
        "D13": ("house_number",),
        "F13": ("survey_khasra_plot_no_as_per_docs", "survey_khasra_plot_no"),
        "F14": ("street_name",),
        "B15": ("floor_number",),
        "D15": ("survey_khasra_plot_no_as_per_docs", "survey_khasra_plot_no"),
        "B16": ("village",),
        "D16": ("ward_number",),
        "F16": ("landmark",),
        "B17": ("tehsil",),
        "D17": ("district",),
        "F17": ("state",),
        "B18": ("branch_name",),
        "D18": ("pincode",),
        "F18": ("property_type",),
        "B19": ("plot_demarcated",),
        "D19": ("distance_from_branch",),
        "F19": ("structure_type",),
        "C26": ("document_type",),
        "B36": ("north_boundary_as_per_docs",),
        "C36": ("south_boundary_as_per_docs",),
        "D36": ("east_boundary_as_per_docs",),
        "E36": ("west_boundary_as_per_docs",),
        "B37": ("north_boundary_as_per_site",),
        "C37": ("south_boundary_as_per_site",),
        "D37": ("east_boundary_as_per_site",),
        "E37": ("west_boundary_as_per_site",),
        "B39": ("approving_authority",),
        "D39": ("property_age_years", "age_years"),
        "F39": ("residual_age_years",),
        "D40": ("occupancy",),
        "F40": ("occupant_name",),
        "D41": ("person_met",),
        "F41": ("road_width",),
        "B42": ("owner_name",),
        "D42": ("relation_with_person_met",),
        "F42": ("dwelling_unit_type", "property_usage_as_per_site"),
        "B43": ("number_of_tenants",),
        "D43": ("property_identification_status",),
        "F43": ("property_completion_status",),
        "B44": ("person_met_contact",),
        "D44": ("negative_area",),
        "F44": ("community_dominated_area",),
        "C55": ("builtup_area_as_per_site", "builtup_area"),
        "C56": ("builtup_area_as_per_docs",),
        "C57": ("builtup_area_as_per_site", "builtup_area"),
        "B59": ("room_configuration", "floor_wise_usage"),
        "B62": ("land_area_as_per_docs", "land_area"),
        "C62": ("land_area_as_per_site", "land_area"),
        "D62": ("land_area_as_per_site", "land_area"),
        "F62": ("site_dimension_remarks",),
        "C64": ("builtup_area_as_per_site", "builtup_area"),
        "D64": ("builtup_area_as_per_site", "builtup_area"),
        "C68": ("land_area_as_per_site", "land_area"),
        "D68": ("land_rate",),
        "C69": ("builtup_area_as_per_site", "builtup_area"),
        "D69": ("construction_rate",),
        "D72": ("valuation_method",),
        "D73": ("completion_percentage",),
        "C79": ("marketability",),
        "C80": ("surrounding_development_status",),
        "F80": ("occupancy_level",),
        "B82": ("remarks",),
        "B83": ("_blank",),
        "B84": ("_blank",),
        "B85": ("_blank",),
        "B86": ("_blank",),
        "B87": ("_blank",),
        "B88": ("_blank",),
        "B89": ("_blank",),
        "C90": ("latitude",),
        "F90": ("longitude",),
        "B92": ("visit_engineer", "visit_by"),
        "E92": ("visit_engineer", "visit_by"),
        "B93": ("visit_date",),
        "E93": ("report_date",),
    },
}

KNOWN_NUMERIC_CELLS = {
    "laxmi": {
        "G38", "G42", "G47", "G48", "E89", "E91", "E100", "D102", "E102",
    },
    "dcb": {"C108", "I108", "C110", "I110", "E148", "K148"},
    "sbfc": {
        "B62", "C62", "D62", "C64", "D64", "C68", "D68", "C69", "D69",
    },
}

KNOWN_FIXED_FORMULAS = {
    "laxmi": {
        "D89": "=G38/10.76",
        "D91": "=G42/10.76",
        "D100": "=G38/9",
        "D102": "=G42",
    },
}

KNOWN_PRESERVE_FORMULA_IF_EMPTY = {
    "laxmi": {"G48"},
}

KNOWN_CLEAR_CELLS = {
    "laxmi": {"G22", "F25", "F32", "F33", "F34"},
    "dcb": {
        "N22", "W22", "AA22", "M44", "M48", "M53", "M54", "M55", "M56",
        "M57", "M58", "M59", "M61", "M63", "M64", "M65", "M66", "M68",
        "M69", "M70", "M71", "M72", "M73", "M74", "M75", "M76", "M77",
        "M78", "M79", "M80", "M81",
    },
    "sbfc": {
        "C47", "D47", "C48", "D48", "C49", "D49", "C50", "D50", "C51",
        "D51", "C52", "B35", "C35", "D35", "E35", "F11", "D13", "F14",
        "B15", "E44", "F62", "C63", "D63", "C79", "C80", "F80",
    },
}


def _first(profile, *keys):
    for key in keys:
        value = profile.get(key)
        if value not in (None, "", [], {}):
            return value
    return ""


def normalize_profile(profile):
    profile = dict(profile or {})
    profile.setdefault("report_date", date.today().strftime("%d-%m-%Y"))
    profile.setdefault("property_address", _first(
        profile, "property_address_as_per_site", "property_address_as_per_docs"
    ))
    profile.setdefault("customer_name", _first(profile, "applicant_name"))
    profile.setdefault("land_area", _first(profile, "land_area_as_per_docs", "land_area_as_per_site"))
    profile.setdefault("builtup_area", _first(profile, "builtup_area_as_per_site", "builtup_area_as_per_docs"))
    return profile


def token_mapping(profile):
    profile = normalize_profile(profile)
    output = {}
    for token, keys in TOKEN_ALIASES.items():
        output[token] = _first(profile, *keys)
    return output


def _display(value):
    if isinstance(value, (datetime, date)):
        return value.strftime("%d-%m-%Y")
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return value


def _numeric_value(value):
    if isinstance(value, (int, float)):
        return value
    match = re.search(r"-?\d[\d,]*(?:\.\d+)?", str(value or ""))
    if not match:
        return ""
    number = float(match.group(0).replace(",", ""))
    return int(number) if number.is_integer() else number


def _replace_tokens(text, mapping):
    if not isinstance(text, str):
        return text
    result = text
    for token, value in mapping.items():
        result = result.replace("{{" + token + "}}", str(_display(value) if value is not None else ""))
    return result


def _normalized(text):
    return re.sub(r"[^a-z0-9]+", " ", str(text).lower()).strip()


def _key_for_label(text):
    normalized = _normalized(text)
    for phrases, key in LABEL_ALIASES:
        if any(phrase in normalized for phrase in phrases):
            return key
    return None


def _merged_range_for_cell(ws, row, column):
    for merged in ws.merged_cells.ranges:
        if merged.min_row <= row <= merged.max_row and merged.min_col <= column <= merged.max_col:
            return merged
    return None


def _value_target(ws, cell):
    label_range = _merged_range_for_cell(ws, cell.row, cell.column)
    start_col = (label_range.max_col + 1) if label_range else (cell.column + 1)
    if start_col > ws.max_column:
        return None
    candidates = []
    for merged in ws.merged_cells.ranges:
        if merged.min_row <= cell.row <= merged.max_row and merged.min_col >= start_col:
            candidates.append(merged)
    if candidates:
        merged = min(candidates, key=lambda item: item.min_col)
        return ws.cell(merged.min_row, merged.min_col)
    for column in range(start_col, min(ws.max_column, start_col + 12) + 1):
        candidate = ws.cell(cell.row, column)
        if candidate.value not in (None, ""):
            return candidate
    return ws.cell(cell.row, start_col)


def _profile_value(profile, key, label=""):
    if key == "property_address":
        return _first(profile, "property_address_as_per_site", "property_address", "property_address_as_per_docs")
    if key == "customer_name":
        return _first(profile, "customer_name", "applicant_name")
    if key == "market_value":
        return _first(profile, "market_value", "total_market_value")
    if key == "land_area_as_per_docs" and "actual" in _normalized(label):
        return _first(profile, "land_area_as_per_site", "land_area")
    return _first(profile, key)


def fill_excel_labels(workbook, profile):
    profile = normalize_profile(profile)
    for ws in workbook.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str) or not cell.value.strip():
                    continue
                key = _key_for_label(cell.value)
                if not key:
                    continue
                value = _profile_value(profile, key, cell.value)
                if value in ("", None, [], {}):
                    continue
                target = _value_target(ws, cell)
                if target is None:
                    continue
                if isinstance(target.value, str) and target.value.startswith("="):
                    continue
                target.value = _display(value)


def _template_cells(workbook, limit=650):
    output = []
    for ws in workbook.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value in (None, ""):
                    continue
                output.append({
                    "sheet": ws.title,
                    "cell": cell.coordinate,
                    "value": str(cell.value)[:240],
                    "formula": isinstance(cell.value, str) and cell.value.startswith("="),
                })
                if len(output) >= limit:
                    return output
    return output


def apply_ai_assignments(workbook, profile):
    allowed = {
        (item["sheet"], item["cell"]): item
        for item in _template_cells(workbook)
        if not item["formula"] and not _key_for_label(item["value"])
    }
    assignments = map_template_cells(list(allowed.values()), profile)
    for item in assignments:
        sheet_name = item.get("sheet") or workbook.worksheets[0].title
        coordinate = str(item.get("cell", "")).upper()
        if (sheet_name, coordinate) not in allowed:
            continue
        cell = workbook[sheet_name][coordinate]
        if isinstance(cell.value, str) and cell.value.startswith("="):
            continue
        value = item.get("value")
        if value not in (None, ""):
            cell.value = value


def fill_known_excel_cells(workbook, profile, layout_key):
    assignments = KNOWN_CELL_MAPPINGS.get(layout_key, {})
    if not assignments:
        return
    sheet_name = EXCEL_PHOTO_LAYOUTS[layout_key]["sheet"]
    ws = workbook[sheet_name] if sheet_name in workbook.sheetnames else workbook.worksheets[0]
    for coordinate in KNOWN_CLEAR_CELLS.get(layout_key, set()):
        target = ws[coordinate]
        merged = _merged_range_for_cell(ws, target.row, target.column)
        if merged:
            target = ws.cell(merged.min_row, merged.min_col)
        if not (isinstance(target.value, str) and target.value.startswith("=")):
            target.value = ""
    for coordinate, keys in assignments.items():
        target = ws[coordinate]
        merged = _merged_range_for_cell(ws, target.row, target.column)
        if merged:
            target = ws.cell(merged.min_row, merged.min_col)
        value = _first(profile, *keys)
        if (
            value in ("", None, [], {})
            and coordinate in KNOWN_PRESERVE_FORMULA_IF_EMPTY.get(layout_key, set())
            and isinstance(target.value, str)
            and target.value.startswith("=")
        ):
            continue
        if coordinate == "N42" and not value:
            latitude = _first(profile, "latitude")
            longitude = _first(profile, "longitude")
            value = ", ".join(str(part) for part in (latitude, longitude) if part not in ("", None))
        if coordinate in KNOWN_NUMERIC_CELLS.get(layout_key, set()):
            value = _numeric_value(value)
        target.value = _display(value) if value not in ("", None, [], {}) else ""
    for coordinate, formula in KNOWN_FIXED_FORMULAS.get(layout_key, {}).items():
        ws[coordinate] = formula


NUMBERED_PROPERTY_PHOTO_CATEGORIES = {
    1: "Front Side View",
    2: "Front Elevation",
    3: "Distant Property View",
    4: "Front Side View",
    5: "Internal Room",
    6: "Internal Room",
    7: "Internal Room",
    8: "Kitchen",
    9: "Other Site Photo",
    10: "Electricity Meter",
}


def _effective_photo_category(asset):
    category = asset.get("category") or "Other Site Photo"
    filename = Path(asset.get("filename") or "").stem.casefold()
    if any(token in filename for token in ("google_map", "google map")):
        return "Google Map"
    if any(token in filename for token in ("mp_kisan", "mp kishan", "mp_kishan")):
        return "MP Kisan"
    if category == "Other Site Photo":
        match = re.search(r"property[_ -]*photos?[_ -]*(\d+)$", filename)
        if match:
            return NUMBERED_PROPERTY_PHOTO_CATEGORIES.get(int(match.group(1)), category)
    return category


def _photo_groups(photo_assets):
    groups = defaultdict(list)
    for asset in photo_assets or []:
        groups[_effective_photo_category(asset)].append(asset)
    return groups


def _ordered_photos(photo_assets):
    groups = _photo_groups(photo_assets)
    order = [
        "Front Elevation", "Front Side View", "Front Side View", "Approach Road",
        "Distant Property View", "Property Selfie", "Kitchen",
        "Internal Room", "Internal Room", "Internal Room", "Internal Room",
        "Electricity Meter", "Electricity Bill", "Site Sketch", "Location Map",
        "Other Site Photo",
    ]
    output = []
    used = set()
    for category in order:
        for asset in groups.get(category, []):
            identity = id(asset)
            if identity not in used:
                output.append(asset)
                used.add(identity)
                break
    for assets in groups.values():
        for asset in assets:
            if id(asset) not in used:
                output.append(asset)
                used.add(id(asset))
    return output


def _excel_image(asset, max_width, max_height):
    stream = io.BytesIO(asset["content"])
    with PillowImage.open(stream) as image:
        image = ImageOps.exif_transpose(image)
        category = _effective_photo_category(asset)
        protected = category in {"Property Document", "Site Sketch", "Location Map", "MP Kisan"}
        if protected:
            width, height = image.size
            prepared = image.convert("RGB")
            ratio = min(max_width / max(width, 1), max_height / max(height, 1))
            output_width = max(40, int(width * ratio))
            output_height = max(40, int(height * ratio))
        else:
            output_width, output_height = int(max_width), int(max_height)
            prepared = ImageOps.fit(
                image.convert("RGB"),
                (output_width, output_height),
                method=PillowImage.Resampling.LANCZOS,
                centering=(0.5, 0.5),
            )
        image_stream = io.BytesIO()
        prepared.save(image_stream, format="PNG", optimize=True)
        image_stream.seek(0)
    excel_image = ExcelImage(image_stream)
    excel_image.width = output_width
    excel_image.height = output_height
    return excel_image


EXCEL_PHOTO_LAYOUTS = {
    "dcb": {
        "sheet": "Table 1",
        "remove_after_row": 1,
        "slots": [
            ("B172", 420, 300), ("N172", 420, 300), ("B174", 420, 300),
            ("N174", 420, 300), ("B176", 420, 300), ("N176", 420, 300),
            ("B178", 420, 300), ("N178", 420, 300), ("B180", 260, 280),
            ("H180", 260, 280), ("W180", 260, 280), ("B182", 260, 280),
            ("H182", 260, 280), ("W182", 260, 280), ("B184", 650, 330),
        ],
    },
    "sbfc": {
        "sheet": "Table 1",
        "remove_after_row": 1,
        "slots": [
            ("A118", 240, 230, "Front Elevation"),
            ("C118", 240, 230, "Front Side View"),
            ("E118", 240, 230, "Internal Room"),
            ("A129", 240, 230, "Approach Road"),
            ("C129", 240, 230, "Distant Property View"),
            ("E129", 240, 230, "Other Site Photo"),
            ("A140", 240, 230, "Approach Road"),
            ("C140", 240, 230, "Distant Property View"),
            ("E140", 240, 230, "Property Selfie"),
            ("A149", 240, 210, "Internal Room"),
            ("C149", 240, 210, "Front Elevation"),
            ("E149", 240, 210, "Other Site Photo"),
            ("A96", 500, 290, ("Google Map", "Location Map")),
            ("E96", 220, 290, "MP Kisan"),
        ],
    },
    "laxmi": {
        "sheet": "MOTA RAM",
        "remove_after_row": 1,
        "slots": [
            ("A139", 360, 300, "Property Document"),
            ("G139", 360, 300, "Property Document"),
            ("A161", 360, 300, "Front Elevation"),
            ("G161", 360, 300, "Approach Road"),
            ("A182", 360, 300, "Internal Room"),
            ("G182", 360, 300, "Kitchen"),
            ("A207", 360, 260, "Property Selfie"),
            ("G207", 360, 260, ("Electricity Meter", "Electricity Bill")),
            ("A210", 720, 360, "Site Sketch"),
            ("A236", 720, 360, "Location Map"),
        ],
    },
}


def _layout_key(template_name, bank_name=""):
    text = f"{template_name} {bank_name}".lower()
    if "dcb" in text:
        return "dcb"
    if "sbfc" in text:
        return "sbfc"
    if "laxmi" in text:
        return "laxmi"
    return ""


def insert_excel_photos(workbook, photo_assets, template_name="", bank_name=""):
    layout = EXCEL_PHOTO_LAYOUTS.get(_layout_key(template_name, bank_name))
    if not layout:
        return
    ws = workbook[layout["sheet"]] if layout["sheet"] in workbook.sheetnames else workbook.worksheets[0]
    # openpyxl anchors are zero-based. Preserve a bank logo/header anchored in
    # the first row and replace the old sample property photographs.
    cutoff = layout["remove_after_row"]
    ws._images = [
        image for image in ws._images
        if not hasattr(image.anchor, "_from") or image.anchor._from.row < cutoff
    ]
    photos = _ordered_photos(photo_assets)
    slots = layout["slots"]
    if slots and len(slots[0]) == 4:
        groups = _photo_groups(photos)
        selected = []
        used = set()
        for anchor, width, height, requested in slots:
            categories = requested if isinstance(requested, tuple) else (requested,)
            asset = next(
                (
                    candidate
                    for category in categories
                    for candidate in groups.get(category, [])
                    if id(candidate) not in used
                ),
                None,
            )
            if asset is None and requested not in (
                "Property Document", "Site Sketch", "Location Map", "Google Map", "MP Kisan",
            ):
                asset = next(
                    (
                        candidate for candidate in photos
                        if (
                            id(candidate) not in used
                            and (candidate.get("category") or "") not in {
                                "Property Document", "Site Sketch", "Location Map", "Google Map", "MP Kisan",
                            }
                        )
                    ),
                    None,
                )
            if asset is None:
                continue
            used.add(id(asset))
            selected.append((asset, (anchor, width, height)))
    else:
        selected = list(zip(photos, slots))
    for asset, (anchor, width, height) in selected:
        try:
            image = _excel_image(asset, width, height)
            image.anchor = anchor
            ws.add_image(image)
        except Exception:
            continue


PHOTO_LABELS = (
    (("front elevation", "front view", "property front"), "Front Elevation"),
    (("side view",), "Front Side View"),
    (("approach road", "road view"), "Approach Road"),
    (("distant view", "long view"), "Distant Property View"),
    (("selfie",), "Property Selfie"),
    (("kitchen",), "Kitchen"),
    (("interior", "internal view", "room photo"), "Internal Room"),
    (("electricity meter",), "Electricity Meter"),
    (("electricity bill",), "Electricity Bill"),
    (("site sketch",), "Site Sketch"),
    (("location map",), "Location Map"),
)


def _photo_category_for_label(value):
    normalized = _normalized(value)
    for labels, category in PHOTO_LABELS:
        if any(label in normalized for label in labels):
            return category
    if "photo" in normalized or "photograph" in normalized:
        return "Other Site Photo"
    return ""


def insert_labeled_excel_photos(workbook, photo_assets):
    photos = _ordered_photos(photo_assets or [])
    if not photos:
        return
    by_category = _photo_groups(photos)
    used = set()
    fallback_index = 0
    for ws in workbook.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str):
                    continue
                category = _photo_category_for_label(cell.value)
                if not category:
                    continue
                asset = next(
                    (
                        item for item in by_category.get(category, [])
                        if id(item) not in used
                    ),
                    None,
                )
                while asset is None and fallback_index < len(photos):
                    candidate = photos[fallback_index]
                    fallback_index += 1
                    if id(candidate) not in used:
                        asset = candidate
                if asset is None:
                    return
                merged = _merged_range_for_cell(ws, cell.row, cell.column)
                anchor_row = (merged.max_row + 1) if merged else (cell.row + 1)
                anchor_col = merged.min_col if merged else cell.column
                try:
                    image = _excel_image(asset, 360, 260)
                    image.anchor = f"{get_column_letter(anchor_col)}{anchor_row}"
                    ws.add_image(image)
                    used.add(id(asset))
                except Exception:
                    continue


def _excel_structure(workbook):
    return tuple(
        (
            ws.title,
            ws.calculate_dimension(),
            tuple(str(item) for item in ws.merged_cells.ranges),
            tuple(
                (index, dimension.height, dimension.hidden)
                for index, dimension in ws.row_dimensions.items()
            ),
            tuple(
                (index, dimension.width, dimension.hidden)
                for index, dimension in ws.column_dimensions.items()
            ),
            str(ws.freeze_panes or ""),
            str(ws.print_area or ""),
        )
        for ws in workbook.worksheets
    )


def fill_excel_template(content, profile, photo_assets=None, template_name="", bank_name=""):
    keep_vba = Path(template_name).suffix.lower() == ".xlsm"
    workbook = load_workbook(io.BytesIO(content), keep_vba=keep_vba)
    original_structure = _excel_structure(workbook)
    profile = normalize_profile(profile)
    mapping = token_mapping(profile)
    for ws in workbook.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    cell.value = _replace_tokens(cell.value, mapping)
    layout_key = _layout_key(template_name, bank_name)
    if layout_key:
        fill_known_excel_cells(workbook, profile, layout_key)
        insert_excel_photos(workbook, photo_assets or [], template_name, bank_name)
    else:
        fill_excel_labels(workbook, profile)
        insert_labeled_excel_photos(workbook, photo_assets or [])
    if _excel_structure(workbook) != original_structure:
        raise ValueError("Uploaded bank workbook structure changed; report was not generated.")
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


UMMEED_ROW_KEYS = {
    1: "application_number", 2: "case_type", 3: "customer_name", 4: "owner_name",
    5: "document_type", 6: "property_address_as_per_site",
    7: "property_address_as_per_docs", 8: "contact_number", 9: "landmark",
    10: "visit_date", 11: "property_usage_as_per_docs", 12: "occupancy",
    14: "marketability", 15: "road_width", 17: "ward_number",
    18: "locality_type", 20: "property_type", 21: "distance_from_branch",
    22: "site_access", 23: "approving_authority", 26: "road_type",
    28: "number_of_floors", 29: "floor_wise_usage", 30: "property_age_years",
    37: "property_identified_through", 38: "plot_demarcated",
    44: "structure_type", 47: "construction_quality",
    49: "room_configuration", 51: "flooring", 54: "doors_windows",
    57: "plan_details", 58: "construction_permission", 59: "demolition_risk",
    75: "land_area_as_per_docs", 76: "land_value", 78: "builtup_area_as_per_site",
    80: "net_building_value", 81: "amenities_value", 82: "depreciation_amount",
    83: "market_value", 84: "distress_value",
}


def _set_docx_cell(cell, value):
    value = str(_display(value))
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _docx_picture(cell, asset, width=2.9, height=2.15):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = 1
    with PillowImage.open(io.BytesIO(asset["content"])) as image:
        image = ImageOps.exif_transpose(image).convert("RGB")
        if (asset.get("category") or "") not in {
            "Property Document", "Site Sketch", "Location Map", "MP Kisan"
        }:
            image = ImageOps.fit(
                image,
                (int(width * 300), int(height * 300)),
                method=PillowImage.Resampling.LANCZOS,
                centering=(0.5, 0.5),
            )
        stream = io.BytesIO()
        image.save(stream, format="PNG", optimize=True)
        stream.seek(0)
        paragraph.add_run().add_picture(
            stream, width=Inches(width),
            **({} if (asset.get("category") or "") in {
                "Property Document", "Site Sketch", "Location Map", "MP Kisan"
            } else {"height": Inches(height)})
        )


def _looks_like_ummeed(document, template_name="", bank_name=""):
    if "ummeed" in f"{template_name} {bank_name}".casefold():
        return True
    if len(document.tables) != 6 or not document.tables:
        return False
    if len(document.tables[0].rows) >= 85:
        return True
    text = " ".join(
        cell.text for row in document.tables[0].rows[:12] for cell in row.cells
    ).casefold()
    return "application" in text and "property address" in text


def _fill_docx_labels(document, profile):
    for table in document.tables:
        for row in table.rows:
            for index, cell in enumerate(row.cells[:-1]):
                key = _key_for_label(cell.text)
                if not key:
                    continue
                value = _profile_value(profile, key, cell.text)
                if value not in ("", None, [], {}):
                    _set_docx_cell(row.cells[index + 1], value)


def _insert_labeled_docx_photos(document, photo_assets):
    photos = _ordered_photos(photo_assets or [])
    used = set()
    for table in document.tables:
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                category = _photo_category_for_label(cell.text)
                if not category:
                    continue
                asset = next(
                    (
                        item for item in photos
                        if id(item) not in used
                        and (
                            item.get("category") == category
                            or category == "Other Site Photo"
                        )
                    ),
                    None,
                )
                if asset is None:
                    asset = next((item for item in photos if id(item) not in used), None)
                if asset is None:
                    return
                target = row.cells[index + 1] if index + 1 < len(row.cells) else cell
                try:
                    _docx_picture(target, asset)
                    used.add(id(asset))
                except Exception:
                    continue


def fill_docx_template(
    content, profile, photo_assets=None, template_name="", bank_name=""
):
    profile = normalize_profile(profile)
    document = Document(io.BytesIO(content))
    mapping = token_mapping(profile)
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.text = _replace_tokens(run.text, mapping)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = _replace_tokens(run.text, mapping)

    if _looks_like_ummeed(document, template_name, bank_name):
        summary = document.tables[0]
        for row_index, key in UMMEED_ROW_KEYS.items():
            if row_index >= len(summary.rows):
                continue
            value = _profile_value(profile, key)
            if value in ("", None, [], {}):
                continue
            row = summary.rows[row_index]
            _set_docx_cell(row.cells[-1], value)

    else:
        _fill_docx_labels(document, profile)

    if _looks_like_ummeed(document, template_name, bank_name) and len(document.tables) > 3:
        photos = _ordered_photos(photo_assets or [])
        photo_table = document.tables[3]
        cells = [
            photo_table.rows[2].cells[0], photo_table.rows[2].cells[2],
            photo_table.rows[2].cells[3], photo_table.rows[4].cells[0],
            photo_table.rows[4].cells[2], photo_table.rows[4].cells[3],
            photo_table.rows[7].cells[0], photo_table.rows[7].cells[2],
            photo_table.rows[7].cells[3],
        ]
        for asset, cell in zip(photos, cells):
            try:
                _docx_picture(cell, asset)
            except Exception:
                continue
    else:
        _insert_labeled_docx_photos(document, photo_assets or [])

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def generic_report(profile):
    profile = normalize_profile(profile)
    workbook = Workbook()
    ws = workbook.active
    ws.title = "Valuation Report"
    ws.append(["SVAI - SAKSHAM ASSOCIATE PROPERTY VALUATION REPORT"])
    ws.merge_cells("A1:D1")
    ws["A1"].font = Font(bold=True, size=15, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor="17365D")
    rows = [
        ("Application Number", _first(profile, "application_number")),
        ("Customer Name", _first(profile, "customer_name", "applicant_name")),
        ("Bank / Branch", f"{_first(profile, 'bank_name')} / {_first(profile, 'branch_name')}"),
        ("Property Address - Documents", _first(profile, "property_address_as_per_docs")),
        ("Property Address - Actual Site", _first(profile, "property_address_as_per_site")),
        ("Owner as per Documents", _first(profile, "owner_name")),
        ("Land Area - Documents", _first(profile, "land_area_as_per_docs", "land_area")),
        ("Land Area - Actual Site", _first(profile, "land_area_as_per_site")),
        ("Built-up Area - Documents", _first(profile, "builtup_area_as_per_docs")),
        ("Built-up Area - Actual Site", _first(profile, "builtup_area_as_per_site", "builtup_area")),
        ("North Boundary - Documents", _first(profile, "north_boundary_as_per_docs")),
        ("North Boundary - Actual", _first(profile, "north_boundary_as_per_site")),
        ("South Boundary - Documents", _first(profile, "south_boundary_as_per_docs")),
        ("South Boundary - Actual", _first(profile, "south_boundary_as_per_site")),
        ("East Boundary - Documents", _first(profile, "east_boundary_as_per_docs")),
        ("East Boundary - Actual", _first(profile, "east_boundary_as_per_site")),
        ("West Boundary - Documents", _first(profile, "west_boundary_as_per_docs")),
        ("West Boundary - Actual", _first(profile, "west_boundary_as_per_site")),
        ("Market Value", _first(profile, "market_value")),
        ("Conservative Value", _first(profile, "conservative_value")),
        ("Distress / QSV", _first(profile, "distress_value")),
        ("Remarks", _first(profile, "remarks")),
    ]
    for label, value in rows:
        ws.append([label, _display(value)])
    for row in range(2, len(rows) + 2):
        ws.cell(row, 1).font = Font(bold=True)
        ws.cell(row, 1).fill = PatternFill("solid", fgColor="D9EAF7")
    ws.column_dimensions["A"].width = 35
    ws.column_dimensions["B"].width = 95
    ws.column_dimensions["C"].width = 3
    ws.column_dimensions["D"].width = 3
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()

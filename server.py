import base64
import gc
import io
import json
import os
import re
import secrets
import imaplib
import smtplib
import tempfile
import time
import threading
import email as email_lib
import zipfile
from datetime import date, datetime, timedelta
from email.message import EmailMessage
from functools import wraps
from html import unescape
from pathlib import Path
from typing import Optional
from urllib.parse import quote, urlsplit
from zoneinfo import ZoneInfo

from cryptography.fernet import Fernet
from dotenv import dotenv_values, load_dotenv
from docx import Document
from flask import (
    Flask, abort, flash, jsonify, redirect, render_template, request,
    send_file, session, url_for
)
from flask_sqlalchemy import SQLAlchemy
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from pypdf import PdfReader
from sqlalchemy import inspect, text
from sqlalchemy.orm import defer
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename

BASE_DIR = Path(__file__).resolve().parent
ENV_PATH = BASE_DIR / ".env"
load_dotenv(ENV_PATH)
# Never inherit an unrelated machine/process API key. SVAI uses only the key
# explicitly saved in its own local .env file.
local_env = dotenv_values(ENV_PATH)
if "OPENAI_API_KEY" in local_env:
    os.environ["OPENAI_API_KEY"] = local_env.get("OPENAI_API_KEY") or ""

from ai_service_openai import (
    OPENAI_MODEL, ai_enabled, build_case_profile, configure_openai,
    classify_property_photo, deterministic_email_candidate,
    document_ai_enabled,
    enrich_email_details_from_attachments, extract_property_asset,
    extract_valuation_email,
)
from report_service import fill_docx_template, fill_excel_template

try:
    from apscheduler.schedulers.background import BackgroundScheduler
except Exception:
    BackgroundScheduler = None


app = Flask(__name__)
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", "change-this-secret-in-render")
database_url = os.getenv("DATABASE_URL", f"sqlite:///{BASE_DIR / 'svai.db'}")
if database_url.startswith("postgres://"):
    database_url = database_url.replace(
        "postgres://", "postgresql+psycopg://", 1
    )
elif database_url.startswith("postgresql://"):
    database_url = database_url.replace(
        "postgresql://", "postgresql+psycopg://", 1
    )
app.config["SQLALCHEMY_DATABASE_URI"] = database_url
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
if database_url.startswith("postgresql+"):
    app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
        "pool_pre_ping": True,
        "pool_recycle": 280,
    }
app.config["MAX_CONTENT_LENGTH"] = int(os.getenv("MAX_UPLOAD_MB", "20")) * 1024 * 1024
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_SECURE"] = os.getenv("SESSION_COOKIE_SECURE", "false").lower() == "true"
db = SQLAlchemy(app)
_setup_lock = threading.Lock()
_runtime_setup_done = None
_email_fetch_lock = threading.Lock()

DOCUMENT_EXTENSIONS = {
    ".pdf", ".jpg", ".jpeg", ".png", ".webp", ".docx", ".xlsx"
}
PHOTO_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}
TEMPLATE_EXTENSIONS = {".xlsx", ".xlsm", ".docx"}

_offline_ocr_engine = None


def offline_ocr_text(filename: str, content: bytes) -> str:
    """Read scanned PDFs/images locally; no network or paid API is used."""
    global _offline_ocr_engine
    ext = Path(filename).suffix.lower()
    if ext not in PHOTO_EXTENSIONS | {".pdf"} or not content:
        return ""
    try:
        import cv2
        import fitz
        import numpy as np
        from rapidocr_onnxruntime import RapidOCR

        if _offline_ocr_engine is None:
            _offline_ocr_engine = RapidOCR()
        images = []
        if ext == ".pdf":
            document = fitz.open(stream=content, filetype="pdf")
            max_pages = max(1, int(os.getenv("LOCAL_OCR_MAX_PAGES", "6")))
            for page in document[:max_pages]:
                pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                image = cv2.imdecode(
                    np.frombuffer(pixmap.tobytes("png"), dtype=np.uint8),
                    cv2.IMREAD_COLOR,
                )
                if image is not None:
                    images.append(image)
            document.close()
        else:
            image = cv2.imdecode(
                np.frombuffer(content, dtype=np.uint8), cv2.IMREAD_COLOR
            )
            if image is not None:
                images.append(image)
        pages = []
        for image in images:
            result, _ = _offline_ocr_engine(image)
            if result:
                pages.append("\n".join(
                    str(line[1]).strip() for line in result
                    if len(line) > 1 and str(line[1]).strip()
                ))
        return "\n\n".join(pages)[:50000]
    except Exception:
        app.logger.exception("Offline OCR failed for %s", filename)
        return ""
SEED_TEMPLATES_DIR = BASE_DIR / "seed_templates"
APP_TIMEZONE = ZoneInfo(os.getenv("APP_TIMEZONE", "Asia/Kolkata"))

PHOTO_CATEGORIES = {
    "front": "Front Elevation",
    "elevation": "Front Elevation",
    "approach": "Approach Road",
    "road": "Approach Road",
    "kitchen": "Kitchen",
    "hall": "Hall / Drawing Room",
    "bed": "Bedroom",
    "room": "Internal Room",
    "meter": "Electricity Meter",
    "selfie": "Property Selfie",
    "side": "Side View",
    "back": "Rear View",
    "toilet": "Toilet / Bathroom",
    "bath": "Toilet / Bathroom",
    "terrace": "Terrace",
    "map": "Location / Map",
    "sketch": "Site Sketch",
}


class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(180), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    name = db.Column(db.String(180), nullable=False, default="SVAI User")
    role = db.Column(db.String(30), nullable=False, default="admin")
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class EmailAccount(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(180), unique=True, nullable=False)
    encrypted_password = db.Column(db.Text, nullable=False)
    provider = db.Column(db.String(30), nullable=False, default="auto")
    imap_host = db.Column(db.String(180))
    bank_name = db.Column(db.String(180))
    active = db.Column(db.Boolean, default=True)
    last_fetch_at = db.Column(db.DateTime)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class SiteEngineer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(180), nullable=False, index=True)
    mobile_number = db.Column(db.String(20), nullable=False)
    area = db.Column(db.String(180), default="")
    active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class WhatsAppGroup(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(180), nullable=False, index=True)
    area = db.Column(db.String(180), default="")
    invite_url = db.Column(db.Text, nullable=False)
    active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class ValuationCase(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    application_number = db.Column(db.String(180), index=True)
    customer_name = db.Column(db.String(220), index=True)
    contact_number = db.Column(db.String(40))
    property_address = db.Column(db.Text)
    bank_name = db.Column(db.String(180), index=True)
    branch_name = db.Column(db.String(180))
    case_type = db.Column(db.String(100), default="LAP")
    source_email = db.Column(db.String(180))
    source_message_id = db.Column(db.String(255), unique=True)
    email_subject = db.Column(db.Text)
    email_received_at = db.Column(db.DateTime)
    visit_by = db.Column(db.String(180))
    status = db.Column(db.String(80), default="New")
    archived = db.Column(db.Boolean, default=False)
    extracted_json = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class FileAsset(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    case_id = db.Column(db.Integer, db.ForeignKey("valuation_case.id"), nullable=True, index=True)
    asset_type = db.Column(db.String(40), nullable=False)  # document/photo/template/report
    category = db.Column(db.String(120))
    filename = db.Column(db.String(255), nullable=False)
    mime_type = db.Column(db.String(120))
    content = db.Column(db.LargeBinary, nullable=False)
    extracted_text = db.Column(db.Text)
    extraction_json = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class BillingTemplate(db.Model):
    """A bank's original invoice workbook plus its reusable billing identity."""
    id = db.Column(db.Integer, primary_key=True)
    bank_name = db.Column(db.String(180), nullable=False, index=True)
    branch_name = db.Column(db.String(180))
    filename = db.Column(db.String(255), nullable=False)
    mime_type = db.Column(db.String(120))
    content = db.Column(db.LargeBinary, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class BillingRatePlan(db.Model):
    """Reusable KM slabs; one current plan per bank and optional branch."""
    id = db.Column(db.Integer, primary_key=True)
    bank_name = db.Column(db.String(180), nullable=False, index=True)
    branch_name = db.Column(db.String(180), default="")
    slabs_json = db.Column(db.Text, nullable=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class StaffPaymentProfile(db.Model):
    """Reusable salary/visit and petrol rules for one staff member or office item."""
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(180), unique=True, nullable=False, index=True)
    category = db.Column(db.String(30), nullable=False, default="Staff")
    payment_mode = db.Column(db.String(30), nullable=False, default="Per Visit")
    fixed_salary = db.Column(db.Float, nullable=False, default=0)
    visit_rate = db.Column(db.Float, nullable=False, default=0)
    petrol_rate = db.Column(db.Float, nullable=False, default=0)
    active = db.Column(db.Boolean, nullable=False, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class StaffMonthlyPayment(db.Model):
    """One billing-ready payment row per staff/profile and calendar month."""
    id = db.Column(db.Integer, primary_key=True)
    profile_id = db.Column(
        db.Integer, db.ForeignKey("staff_payment_profile.id"), nullable=False, index=True
    )
    month = db.Column(db.String(7), nullable=False, index=True)
    visits = db.Column(db.Integer, nullable=False, default=0)
    km = db.Column(db.Float, nullable=False, default=0)
    base_amount = db.Column(db.Float, nullable=False, default=0)
    conveyance_amount = db.Column(db.Float, nullable=False, default=0)
    advance_amount = db.Column(db.Float, nullable=False, default=0)
    adjustment_amount = db.Column(db.Float, nullable=False, default=0)
    final_amount = db.Column(db.Float, nullable=False, default=0)
    status = db.Column(db.String(30), nullable=False, default="Pending")
    notes = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    profile = db.relationship("StaffPaymentProfile")
    __table_args__ = (
        db.UniqueConstraint("profile_id", "month", name="uq_staff_payment_month"),
    )


class Valuation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    case_id = db.Column(db.Integer, db.ForeignKey("valuation_case.id"), unique=True, nullable=False)
    land_area = db.Column(db.Float, default=0)
    land_rate = db.Column(db.Float, default=0)
    builtup_area = db.Column(db.Float, default=0)
    construction_rate = db.Column(db.Float, default=0)
    age_years = db.Column(db.Float, default=0)
    depreciation_percent = db.Column(db.Float, default=0)
    govt_land_rate = db.Column(db.Float, default=0)
    govt_construction_rate = db.Column(db.Float, default=0)
    conservative_percent = db.Column(db.Float, default=100)
    distress_percent = db.Column(db.Float, default=80)
    land_value = db.Column(db.Float, default=0)
    gross_building_value = db.Column(db.Float, default=0)
    depreciation_amount = db.Column(db.Float, default=0)
    net_building_value = db.Column(db.Float, default=0)
    market_value = db.Column(db.Float, default=0)
    conservative_value = db.Column(db.Float, default=0)
    distress_value = db.Column(db.Float, default=0)
    govt_value = db.Column(db.Float, default=0)
    remarks = db.Column(db.Text)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


def encryption_key() -> bytes:
    configured = os.getenv("ENCRYPTION_KEY", "").strip()
    if configured:
        try:
            Fernet(configured.encode())
            return configured.encode()
        except (TypeError, ValueError):
            # Render's generated secrets are not guaranteed to already be
            # Fernet-formatted. Derive a stable valid key from any secret.
            raw = configured.encode()
            return base64.urlsafe_b64encode(__import__("hashlib").sha256(raw).digest())
    raw = str(app.config["SECRET_KEY"]).encode()
    return base64.urlsafe_b64encode(__import__("hashlib").sha256(raw).digest())


FERNET = Fernet(encryption_key())


def encrypt_password(value: str) -> str:
    return FERNET.encrypt(value.encode()).decode()


def decrypt_password(value: str) -> str:
    return FERNET.decrypt(value.encode()).decode()


def login_required(fn):
    @wraps(fn)
    def wrapped(*args, **kwargs):
        if not session.get("user_id"):
            return redirect(url_for("login"))
        return fn(*args, **kwargs)
    return wrapped


def _backup_json_value(value):
    """Encode database values losslessly for a portable private backup."""
    if value is None or isinstance(value, (bool, int, float, str)):
        return value
    if isinstance(value, bytes):
        return {"$type": "bytes", "base64": base64.b64encode(value).decode("ascii")}
    if isinstance(value, (datetime, date)):
        return {"$type": type(value).__name__, "value": value.isoformat()}
    return {"$type": "string", "value": str(value)}


@app.route("/admin/portable-database-backup")
@login_required
def portable_database_backup():
    """Download a private, host-neutral backup without exposing DB credentials."""
    user = db.session.get(User, session.get("user_id"))
    if not user or user.role != "admin":
        abort(403)

    stamp = datetime.now(APP_TIMEZONE).strftime("%Y%m%d-%H%M%S")
    handle, temp_name = tempfile.mkstemp(prefix="svai-private-backup-", suffix=".zip")
    os.close(handle)
    manifest = {
        "format": "SVAI portable database backup v1",
        "created_at": datetime.now(APP_TIMEZONE).isoformat(),
        "tables": [],
    }
    try:
        schema = inspect(db.engine)
        preparer = db.engine.dialect.identifier_preparer
        with zipfile.ZipFile(temp_name, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as bundle:
            with db.engine.connect() as connection:
                transaction = connection.begin()
                try:
                    if db.engine.dialect.name == "postgresql":
                        connection.execute(text("SET TRANSACTION ISOLATION LEVEL REPEATABLE READ, READ ONLY"))
                    for table_name in sorted(schema.get_table_names()):
                        columns = schema.get_columns(table_name)
                        quoted_table = preparer.quote(table_name)
                        result = connection.execution_options(stream_results=True).execute(
                            text(f"SELECT * FROM {quoted_table}")
                        ).mappings()
                        count = 0
                        with bundle.open(f"tables/{table_name}.jsonl", "w") as output:
                            while True:
                                rows = result.fetchmany(50)
                                if not rows:
                                    break
                                for row in rows:
                                    payload = {key: _backup_json_value(value) for key, value in row.items()}
                                    line = json.dumps(payload, ensure_ascii=False, separators=(",", ":")) + "\n"
                                    output.write(line.encode("utf-8"))
                                    count += 1
                        manifest["tables"].append({
                            "name": table_name,
                            "row_count": count,
                            "columns": [
                                {"name": column["name"], "type": str(column["type"]), "nullable": column["nullable"]}
                                for column in columns
                            ],
                        })
                finally:
                    transaction.rollback()
            bundle.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
            bundle.writestr(
                "RESTORE_README.txt",
                "PRIVATE SVAI PRODUCTION DATA. Do not upload to GitHub or share publicly.\n"
                "Create the schema from the matching SVAI source commit, then restore the JSONL tables.\n",
            )
        response = send_file(
            temp_name,
            as_attachment=True,
            download_name=f"svai-private-production-backup-{stamp}.zip",
            mimetype="application/zip",
        )
        response.call_on_close(lambda: os.path.exists(temp_name) and os.remove(temp_name))
        return response
    except Exception:
        if os.path.exists(temp_name):
            os.remove(temp_name)
        raise


def api_login_required(fn):
    @wraps(fn)
    def wrapped(*args, **kwargs):
        if not session.get("user_id"):
            return jsonify({"success": False, "message": "Login required"}), 401
        return fn(*args, **kwargs)
    return wrapped


def safe_json(text: Optional[str], default=None):
    if not text:
        return default if default is not None else {}
    try:
        return json.loads(text)
    except Exception:
        return default if default is not None else {}


def update_env_file(updates):
    """Update selected .env values while preserving unrelated configuration."""
    env_path = BASE_DIR / ".env"
    existing = env_path.read_text(encoding="utf-8").splitlines() if env_path.exists() else []
    remaining = dict(updates)
    output = []
    for line in existing:
        match = re.match(r"^\s*([A-Za-z_][A-Za-z0-9_]*)\s*=", line)
        key = match.group(1) if match else None
        if key in remaining:
            value = str(remaining.pop(key)).replace("\\", "\\\\").replace('"', '\\"')
            output.append(f'{key}="{value}"')
        else:
            output.append(line)
    for key, raw_value in remaining.items():
        value = str(raw_value).replace("\\", "\\\\").replace('"', '\\"')
        output.append(f'{key}="{value}"')
    env_path.write_text("\n".join(output) + "\n", encoding="utf-8")


def detect_imap(email_address: str, provider: str = "auto", custom_host: str = ""):
    if custom_host:
        return custom_host
    domain = email_address.lower().split("@")[-1]
    provider = (provider or "auto").lower()
    if "gmail" in provider or domain in {"gmail.com", "googlemail.com"}:
        return "imap.gmail.com"
    if "yahoo" in provider or "yahoo" in domain:
        return "imap.mail.yahoo.com"
    raise ValueError("Only Gmail and Yahoo IMAP accounts are supported.")


def detect_smtp(email_address: str, provider: str = "auto"):
    domain = email_address.lower().split("@")[-1]
    provider = (provider or "auto").lower()
    if "gmail" in provider or domain in {"gmail.com", "googlemail.com"}:
        return "smtp.gmail.com"
    if "yahoo" in provider or "yahoo" in domain:
        return "smtp.mail.yahoo.com"
    raise ValueError("Only Gmail and Yahoo SMTP accounts are supported.")


def send_password_reset_code(account, recipient, code):
    message = EmailMessage()
    message["Subject"] = "SVAI password reset code"
    message["From"] = account.email
    message["To"] = recipient
    message.set_content(
        "Your SVAI password reset code is:\n\n"
        f"{code}\n\n"
        "This code expires in 10 minutes. If you did not request it, ignore "
        "this email."
    )
    host = detect_smtp(account.email, account.provider)
    password = decrypt_password(account.encrypted_password)
    try:
        with smtplib.SMTP_SSL(host, 465, timeout=12) as smtp:
            smtp.login(account.email, password)
            smtp.send_message(message)
            return
    except (OSError, smtplib.SMTPException):
        # Some cloud hosts block outbound SMTPS/465 while allowing the
        # standard STARTTLS submission port. Gmail and Yahoo support both.
        with smtplib.SMTP(host, 587, timeout=20) as smtp:
            smtp.ehlo()
            smtp.starttls()
            smtp.ehlo()
            smtp.login(account.email, password)
            smtp.send_message(message)


def decode_header_value(value: str) -> str:
    if not value:
        return ""
    parts = []
    for piece, charset in email_lib.header.decode_header(value):
        if isinstance(piece, bytes):
            parts.append(piece.decode(charset or "utf-8", errors="ignore"))
        else:
            parts.append(str(piece))
    return "".join(parts)


def email_body(message) -> str:
    chunks = []
    if message.is_multipart():
        for part in message.walk():
            content_type = part.get_content_type()
            disposition = str(part.get("Content-Disposition", ""))
            if content_type in ("text/plain", "text/html") and "attachment" not in disposition:
                try:
                    payload = part.get_payload(decode=True) or b""
                    text = payload.decode(part.get_content_charset() or "utf-8", errors="ignore")
                    if content_type == "text/html":
                        text = re.sub(
                            r"(?is)<(?:script|style)\b[^>]*>.*?</(?:script|style)>",
                            " ",
                            text,
                        )
                        text = re.sub(
                            r"(?i)</?(?:p|div|br|tr|li|table|h[1-6])\b[^>]*>",
                            "\n",
                            text,
                        )
                        text = unescape(re.sub(r"<[^>]+>", " ", text))
                    chunks.append(text)
                except Exception:
                    continue
    else:
        payload = message.get_payload(decode=True) or b""
        chunks.append(payload.decode(message.get_content_charset() or "utf-8", errors="ignore"))
    normalized = []
    for line in "\n".join(chunks).replace("\r", "\n").splitlines():
        line = re.sub(r"[ \t\f\v]+", " ", line).strip()
        if line:
            normalized.append(line)
        elif normalized and normalized[-1] != "":
            normalized.append("")
    return "\n".join(normalized).strip()[:25000]


def latest_email_body(body: str) -> str:
    """Keep the current message separate from quoted older email trails.

    Old assignments in a reply/forward must not create a new MIS case.  The
    newest part is still read for genuine current instructions and facts.
    """
    return re.split(
        r"(?im)^\s*(?:-{2,}\s*original message\s*-{2,}|from\s*:|"
        r"on .{0,160} wrote\s*:)",
        str(body or ""),
        maxsplit=1,
    )[0].strip()


def ai_extract_email(subject: str, body: str, sender: str):
    return extract_valuation_email(subject, body, sender)


def is_valuation_email(subject: str, body: str, sender: str = "") -> bool:
    return deterministic_email_candidate(subject, body, sender)


def parse_iso_date(value, default):
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except (TypeError, ValueError):
        return default


def current_month_range():
    today = datetime.now(APP_TIMEZONE).date()
    return today.replace(day=1), today


def normalized_header(value):
    return re.sub(r"[^a-z0-9]+", "", str(value or "").lower())


def bank_template_key(value):
    key = normalized_header(value)
    for noise in ("housingfinance", "homefinance", "finance", "limited", "ltd"):
        key = key.replace(noise, "")
    return key


def matching_master_template(bank_name):
    bank_key = bank_template_key(bank_name)
    if not bank_key:
        return None
    candidates = FileAsset.query.filter_by(asset_type="template").order_by(
        FileAsset.created_at.desc()
    ).all()
    for candidate in candidates:
        keys = (bank_template_key(candidate.category), bank_template_key(candidate.filename))
        if any(key and (key == bank_key or key in bank_key or bank_key in key) for key in keys):
            return candidate
    return None


def numeric_km(value):
    match = re.search(r"\d+(?:\.\d+)?", str(value or ""))
    return float(match.group()) if match else None


def billing_fee_for_km(km, slabs):
    if km is None:
        return None
    for minimum, maximum, amount in slabs:
        if km >= minimum and (maximum is None or km <= maximum):
            return amount
    return None


def parse_billing_slabs(form):
    slabs = []
    for minimum, maximum, amount in zip(
        form.getlist("slab_min[]"), form.getlist("slab_max[]"), form.getlist("slab_amount[]")
    ):
        if not any((minimum, maximum, amount)):
            continue
        try:
            parsed_minimum = float(minimum or 0)
            parsed_maximum = float(maximum) if str(maximum).strip() else None
            parsed_amount = float(amount)
        except ValueError:
            raise ValueError("Har KM slab me valid minimum aur amount enter karein.")
        if parsed_minimum < 0 or parsed_amount < 0 or (parsed_maximum is not None and parsed_maximum < parsed_minimum):
            raise ValueError("KM slab range valid nahi hai.")
        slabs.append((parsed_minimum, parsed_maximum, parsed_amount))
    if not slabs:
        raise ValueError("Kam se kam ek KM slab amount zaroor add karein.")
    return sorted(slabs, key=lambda row: row[0])


def saved_billing_slabs(bank_name, branch_name=""):
    plan = BillingRatePlan.query.filter(
        db.func.lower(BillingRatePlan.bank_name) == (bank_name or "").lower(),
        db.func.lower(BillingRatePlan.branch_name) == (branch_name or "").lower(),
    ).order_by(BillingRatePlan.updated_at.desc()).first()
    if not plan:
        return []
    slabs = safe_json(plan.slabs_json, [])
    return [tuple(item) for item in slabs if isinstance(item, list) and len(item) == 3]


def save_billing_slabs(bank_name, branch_name, slabs):
    plan = BillingRatePlan.query.filter(
        db.func.lower(BillingRatePlan.bank_name) == (bank_name or "").lower(),
        db.func.lower(BillingRatePlan.branch_name) == (branch_name or "").lower(),
    ).first()
    if plan is None:
        plan = BillingRatePlan(bank_name=bank_name, branch_name=branch_name or "")
    plan.slabs_json = json.dumps(slabs)
    db.session.add(plan)


def nonnegative_number(value, label, allow_negative=False):
    try:
        parsed = float(str(value or "0").strip() or 0)
    except (TypeError, ValueError):
        raise ValueError(f"{label} valid number hona chahiye.")
    if not allow_negative and parsed < 0:
        raise ValueError(f"{label} negative nahi ho sakta.")
    return parsed


def normalized_payment_month(value=None):
    month = str(value or datetime.now(APP_TIMEZONE).strftime("%Y-%m")).strip()
    try:
        return datetime.strptime(month, "%Y-%m").strftime("%Y-%m")
    except ValueError:
        raise ValueError("Month YYYY-MM format me select karein.")


def staff_payment_amounts(profile, visits=0, km=0, base_override="", advance=0, adjustment=0):
    parsed_visits = nonnegative_number(visits, "Visits")
    if not parsed_visits.is_integer():
        raise ValueError("Visits whole number hona chahiye.")
    visits = int(parsed_visits)
    km = nonnegative_number(km, "KM")
    advance = nonnegative_number(advance, "Advance")
    adjustment = nonnegative_number(adjustment, "Adjustment", allow_negative=True)
    override_text = str(base_override or "").strip()
    if override_text:
        base_amount = nonnegative_number(override_text, "Base amount")
    elif profile.payment_mode == "Salary":
        base_amount = profile.fixed_salary or 0
    elif profile.payment_mode == "Per Visit":
        base_amount = visits * (profile.visit_rate or 0)
    else:
        base_amount = 0
    conveyance = km * (profile.petrol_rate or 0)
    final_amount = base_amount + conveyance + adjustment - advance
    return {
        "visits": visits,
        "km": km,
        "base_amount": round(base_amount, 2),
        "conveyance_amount": round(conveyance, 2),
        "advance_amount": round(advance, 2),
        "adjustment_amount": round(adjustment, 2),
        "final_amount": round(final_amount, 2),
    }


def staff_names_from_workbook(upload):
    try:
        workbook = load_workbook(io.BytesIO(upload.read()), data_only=True, read_only=True)
    except Exception as exc:
        raise ValueError(f"Staff Excel read nahi hua: {exc}")
    sheet = workbook.active
    header_row = None
    name_column = None
    for row_number, row in enumerate(sheet.iter_rows(values_only=True), 1):
        for column, value in enumerate(row, 1):
            if normalized_header(value) in {"employeename", "staffname", "engineername", "name"}:
                header_row, name_column = row_number, column
                break
        if header_row:
            break
    if not header_row:
        raise ValueError("Excel me Employee Name heading nahi mili.")
    names = []
    for row in sheet.iter_rows(min_row=header_row + 1, values_only=True):
        raw = row[name_column - 1] if name_column <= len(row) else ""
        name = re.sub(r"\s+", " ", str(raw or "")).strip()
        if not name:
            continue
        if normalized_header(name).startswith("total"):
            break
        if normalized_header(name) in {"sbi", "sbicc"}:
            continue
        if name not in names:
            names.append(name)
    return names


def staff_payments_workbook(month, entries):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Conveyance MIS"
    headers = [
        "Month", "Name", "Category", "Payment Basis", "Visits", "Visit Rate",
        "Salary / Base", "KM", "Petrol Rate", "Conveyance", "Adjustment",
        "Advance", "Final Payable", "Status", "Notes",
    ]
    sheet.append([f"SVAI Staff Payment & Conveyance MIS - {month}"])
    sheet.append(headers)
    for entry in entries:
        profile = entry.profile
        sheet.append([
            month, profile.name, profile.category, profile.payment_mode,
            entry.visits, profile.visit_rate, entry.base_amount, entry.km,
            profile.petrol_rate, entry.conveyance_amount, entry.adjustment_amount,
            entry.advance_amount, entry.final_amount, entry.status, entry.notes or "",
        ])
    total_row = sheet.max_row + 1
    sheet.cell(total_row, 1).value = "MONTH TOTAL"
    for column in (7, 10, 11, 12, 13):
        letter = get_column_letter(column)
        sheet.cell(total_row, column).value = f"=SUM({letter}3:{letter}{total_row - 1})"
    sheet.merge_cells("A1:O1")
    sheet["A1"].font = Font(bold=True, size=16, color="FFFFFF")
    sheet["A1"].fill = PatternFill("solid", fgColor="163A63")
    for cell in sheet[2]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F6FA8")
    for cell in sheet[total_row]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="DDEBF7")
    widths = [12, 24, 14, 16, 10, 12, 15, 10, 12, 15, 12, 12, 15, 12, 28]
    for index, width in enumerate(widths, 1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    for row in sheet.iter_rows(min_row=3, max_row=total_row, min_col=6, max_col=13):
        for cell in row:
            cell.number_format = '₹#,##0.00'
    sheet.freeze_panes = "A3"
    sheet.auto_filter.ref = f"A2:O{max(2, total_row - 1)}"
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def billing_case_rows(cases):
    rows = []
    for case in cases:
        profile = safe_json(case.extracted_json)
        profile = profile.get("case_profile") or profile.get("email") or profile
        rows.append({
            "application_number": case.application_number or "",
            "customer_name": case.customer_name or "",
            "property_address": case.property_address or "",
            "case_type": case.case_type or "",
            "bank_name": case.bank_name or "",
            "branch_name": case.branch_name or "",
            "distance": numeric_km(profile.get("distance_from_branch", profile.get("km", ""))),
        })
    return rows


def billing_upload_rows(upload):
    """Read a standard MIS export. It is input only; the original invoice stays untouched."""
    try:
        workbook = load_workbook(io.BytesIO(upload.read()), data_only=True, read_only=True)
    except Exception as exc:
        raise ValueError(f"MIS Excel read nahi hua: {exc}")
    sheet = workbook.active
    header_row = None
    columns = {}
    aliases = {
        "application_number": {"applicationno", "applicationnumber", "leadidno", "idno", "leadproposalno", "leadpurposalno"},
        "customer_name": {"customername", "applicantname", "name"},
        "property_address": {"address", "propertyaddress"},
        "case_type": {"casetype", "product", "productname"},
        "bank_name": {"bank", "bankname"},
        "branch_name": {"branch", "branchname"},
        "distance": {"km", "distance", "distence", "distancefrombranchinkm", "distencefrombranchinkm", "distancekm"},
    }
    for row_number, row in enumerate(sheet.iter_rows(values_only=True), 1):
        found = {normalized_header(value): index for index, value in enumerate(row) if value is not None}
        if any(key in found for key in aliases["application_number"]) and any(
            key in found for key in aliases["customer_name"]
        ):
            header_row = row_number
            for field, names in aliases.items():
                for name in names:
                    if name in found:
                        columns[field] = found[name]
                        break
            break
    if not header_row:
        raise ValueError("Uploaded MIS me Application No aur Customer Name heading nahi mili.")
    rows = []
    for row in sheet.iter_rows(min_row=header_row + 1, values_only=True):
        application = row[columns.get("application_number", -1)] if columns.get("application_number") is not None else ""
        customer = row[columns.get("customer_name", -1)] if columns.get("customer_name") is not None else ""
        if not application and not customer:
            continue
        rows.append({
            field: (row[index] if index is not None and index < len(row) else "")
            for field, index in columns.items()
        })
        rows[-1]["distance"] = numeric_km(rows[-1].get("distance"))
    return rows


def mis_import_rows(upload):
    """Read an existing MIS workbook so its cases can be merged into SVAI."""
    try:
        workbook = load_workbook(io.BytesIO(upload.read()), data_only=True, read_only=True)
    except Exception as exc:
        raise ValueError(f"MIS Excel read nahi hua: {exc}")
    sheet = workbook["ALL BANK"] if "ALL BANK" in workbook.sheetnames else workbook.active
    aliases = {
        "date": {"date", "receiveddate"},
        "customer_name": {"customername", "applicantname", "customer"},
        "application_number": {"applicationno", "applicationnumber", "leadidno", "idno", "leadproposalno", "leadpurposalno"},
        "contact_number": {"contactnumber", "mobilenumber", "mobile", "contact"},
        "bank_name": {"bank", "bankname"},
        "case_type": {"casetype", "product", "productname"},
        "status": {"status", "casestatus"},
        "property_address": {"address", "propertyaddress"},
        "visit_by": {"visitby", "engineer", "engineername"},
        "branch_name": {"branch", "branchname"},
        "distance": {"km", "distance", "distancekm"},
    }
    header_row = None
    columns = {}
    for row_number, row in enumerate(sheet.iter_rows(values_only=True), 1):
        found = {normalized_header(value): index for index, value in enumerate(row) if value is not None}
        if any(name in found for name in aliases["customer_name"]) and any(
            name in found for name in aliases["application_number"]
        ):
            header_row = row_number
            for field, names in aliases.items():
                match = next((found[name] for name in names if name in found), None)
                if match is not None:
                    columns[field] = match
            break
    if not header_row:
        raise ValueError("Uploaded MIS me Customer Name aur Application No headings nahi mili.")

    rows = []
    for row in sheet.iter_rows(min_row=header_row + 1, values_only=True):
        item = {
            field: (row[index] if index < len(row) else "")
            for field, index in columns.items()
        }
        application = str(item.get("application_number") or "").strip()
        customer = str(item.get("customer_name") or "").strip()
        if not application and not customer:
            continue
        raw_date = item.get("date")
        if isinstance(raw_date, datetime):
            received_at = raw_date
        elif isinstance(raw_date, date):
            received_at = datetime.combine(raw_date, datetime.min.time())
        elif isinstance(raw_date, (int, float)):
            received_at = datetime(1899, 12, 30) + timedelta(days=float(raw_date))
        else:
            received_at = None
            for date_format in ("%d-%b-%y", "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%Y"):
                try:
                    received_at = datetime.strptime(str(raw_date or "").strip(), date_format)
                    break
                except ValueError:
                    continue
        item["received_at"] = received_at
        item["application_number"] = application
        item["customer_name"] = customer
        item["distance"] = numeric_km(item.get("distance"))
        rows.append(item)
    return rows


def billing_column_map(sheet):
    aliases = {
        "serial": {"sno", "srno", "serialno"},
        "application": {"applicationno", "applicationnumber", "leadidno", "idno", "leadproposalno", "leadpurposalno"},
        "customer": {"customername", "applicantname", "customer"},
        "product": {"product", "productname", "casetype"},
        "address": {"propertyaddress", "address"},
        "distance": {"distance", "distence", "distancekm", "distancefrombranchinkm", "distencefrombranchinkm", "km"},
        "fee": {"fee", "fees", "amount", "valuationlegalamount"},
    }
    for row_number in range(1, min(sheet.max_row, 100) + 1):
        found = {normalized_header(sheet.cell(row_number, col).value): col for col in range(1, sheet.max_column + 1)}
        mapping = {}
        for field, names in aliases.items():
            for name in names:
                if name in found:
                    mapping[field] = found[name]
                    break
        if "application" in mapping and "customer" in mapping and ("fee" in mapping or "distance" in mapping):
            return row_number, mapping
    raise ValueError("Invoice template me Application/Customer/Fee wali detail heading auto-detect nahi hui. Bank ka line-item invoice format upload karein.")


def generate_billing_workbook(template_content, rows, slabs):
    """Fill an uploaded bank invoice workbook without rebuilding its layout or formulas."""
    source = io.BytesIO(template_content)
    workbook = load_workbook(source, keep_vba=False)
    sheet = workbook.active
    header_row, columns = billing_column_map(sheet)
    start_row = header_row + 1
    # Never overwrite a total / tax / grand-total section. Clear only existing detail rows.
    end_row = start_row
    while end_row <= sheet.max_row:
        label = " ".join(str(sheet.cell(end_row, col).value or "") for col in range(1, min(sheet.max_column, 4) + 1)).lower()
        if any(word in label for word in ("total", "cgst", "sgst", "grand", "tax")):
            break
        if end_row - start_row > 500:
            break
        for col in range(1, sheet.max_column + 1):
            if sheet.cell(end_row, col).data_type != "f":
                sheet.cell(end_row, col).value = None
        end_row += 1
    for index, row in enumerate(rows, start=1):
        excel_row = start_row + index - 1
        if excel_row >= end_row:
            sheet.insert_rows(excel_row)
            # Copy blank-row formatting from the first detail row.
            for col in range(1, sheet.max_column + 1):
                source_cell = sheet.cell(start_row, col)
                target = sheet.cell(excel_row, col)
                if source_cell.has_style:
                    target._style = source_cell._style
                if source_cell.number_format:
                    target.number_format = source_cell.number_format
        fee = billing_fee_for_km(row.get("distance"), slabs)
        values = {
            "serial": index,
            "application": row.get("application_number", ""),
            "customer": row.get("customer_name", ""),
            "product": row.get("case_type", ""),
            "address": row.get("property_address", ""),
            "distance": row.get("distance", ""),
            "fee": fee if fee is not None else "",
        }
        for field, value in values.items():
            if field in columns:
                sheet.cell(excel_row, columns[field]).value = value
    stream = io.BytesIO()
    workbook.save(stream)
    return stream.getvalue(), [row for row in rows if billing_fee_for_km(row.get("distance"), slabs) is None]


def collect_email_attachments(message):
    attachments = []
    for part in message.walk():
        raw_name = decode_header_value(part.get_filename() or "")
        if not raw_name:
            continue
        filename = secure_filename(raw_name)
        content = part.get_payload(decode=True) or b""
        ext = Path(filename).suffix.lower()
        if not content or ext not in DOCUMENT_EXTENSIONS | {".zip"}:
            continue
        if ext == ".zip":
            try:
                with zipfile.ZipFile(io.BytesIO(content)) as bundle:
                    for member in safe_zip_members(bundle):
                        inner_name = secure_filename(Path(member.filename).name)
                        if Path(inner_name).suffix.lower() not in DOCUMENT_EXTENSIONS:
                            continue
                        attachments.append({
                            "filename": inner_name,
                            "content": bundle.read(member),
                            "mime_type": "application/octet-stream",
                        })
            except (zipfile.BadZipFile, ValueError):
                continue
        else:
            attachments.append({
                "filename": filename,
                "content": content,
                "mime_type": part.get_content_type(),
            })
    return attachments


def normalized_email_subject(subject):
    value = re.sub(
        r"(?i)^(?:(?:re|fw|fwd)(?:\[\d+\])?\s*:\s*)+",
        "",
        re.sub(r"\s+", " ", subject or "").strip(),
    )
    return value.casefold()


def normalized_application_number(value):
    """Stable key used only for matching; the displayed application stays unchanged."""
    key = re.sub(r"[^A-Z0-9]", "", str(value or "").upper())
    # Excel/manual MIS often drops numeric leading zeroes while lender emails
    # retain them. They are the same application and must not become two rows.
    return (key.lstrip("0") or "0") if key.isdigit() else key


def is_followup_email(subject, body=""):
    """Return True for report reminders/status mails, not new visits or assignments."""
    subject_text = re.sub(r"\s+", " ", subject or "").strip()
    if re.search(
        r"(?i)^(?:(?:re|fw|fwd)(?:\[\d+\])?\s*:\s*)*"
        r"(?:recall\s*:|revise(?:d)?\s+(?:technical|valuation)\s+report\b)",
        subject_text,
    ):
        return True
    latest_body = latest_email_body(body)
    latest_body = re.sub(r"\s+", " ", latest_body[:2500]).strip()
    combined = f"{subject_text}\n{latest_body}"
    followup_patterns = (
        r"(?i)\b(?:please|pls|kindly)\s+(?:share|send|forward)\s+"
        r"(?:the\s+)?(?:(?:technical|valuation|final)\s+)?report\b",
        r"(?i)\b(?:technical|valuation)?\s*report\s+(?:is\s+)?"
        r"(?:awaited|pending|overdue)\b",
        r"(?i)\bawaiting\s+(?:the\s+)?(?:(?:technical|valuation)\s+)?report\b",
        r"(?i)\b(?:reminder|follow[\s-]*up)\b.{0,100}\b(?:report|status)\b",
        r"(?i)\b(?:report|case)\s+status\s+(?:required|needed|update)\b",
    )
    if not any(re.search(pattern, combined) for pattern in followup_patterns):
        return False
    new_visit_patterns = (
        r"(?i)\bsubsequent\s+visit\b",
        r"(?i)\bre[\s-]*visit\b",
        r"(?i)\b(?:part|tranche)\s+(?:valuation|visit|technical)\b",
        r"(?i)\bconstruction\s+stage\b.{0,80}\b(?:visit|required|complete)\b",
        r"(?i)\b(?:new\s+)?(?:technical|valuation)\s+(?:case\s+)?assignment\b",
    )
    return not any(re.search(pattern, latest_body) for pattern in new_visit_patterns)


def _message_already_recorded(case, unique_id):
    if not case or not unique_id:
        return False
    if case.source_message_id == unique_id:
        return True
    stored = safe_json(case.extracted_json)
    return any(
        item.get("message_id") == unique_id
        for item in stored.get("followup_emails", [])
        if isinstance(item, dict)
    )


def existing_case_for_message(account, unique_id, subject, received):
    direct = ValuationCase.query.filter_by(source_message_id=unique_id).first()
    if direct:
        return direct
    recorded = ValuationCase.query.filter(
        ValuationCase.source_email == account.email,
        ValuationCase.extracted_json.contains(unique_id),
    ).all()
    for item in recorded:
        if _message_already_recorded(item, unique_id):
            return item
    start = datetime.combine(received.date(), datetime.min.time())
    end = start + timedelta(days=1)
    candidates = ValuationCase.query.filter(
        ValuationCase.source_email == account.email,
        ValuationCase.email_received_at >= start,
        ValuationCase.email_received_at < end,
    ).all()
    normalized = normalized_email_subject(subject)
    return next(
        (
            item for item in candidates
            if normalized_email_subject(item.email_subject) == normalized
        ),
        None,
    )


def existing_case_for_application(application_number, exclude_case=None, subject=""):
    key = normalized_application_number(application_number)
    if not key:
        return None
    candidates = ValuationCase.query.filter_by(archived=False).order_by(
        db.func.coalesce(
            ValuationCase.email_received_at, ValuationCase.created_at
        ).asc(),
        ValuationCase.id.asc(),
    ).all()
    matches = [
        item for item in candidates
        if (not exclude_case or item.id != exclude_case.id)
        and normalized_application_number(item.application_number) == key
    ]
    thread_subject = normalized_email_subject(subject)
    if thread_subject:
        threaded = next(
            (
                item for item in matches
                if normalized_email_subject(item.email_subject) == thread_subject
            ),
            None,
        )
        if threaded:
            return threaded
    return matches[0] if matches else None


def existing_case_for_duplicate_assignment(details, account, subject="", received=None):
    """Merge same assignment arriving in Gmail and Yahoo; keep new work separate."""
    key = normalized_application_number(details.get("application_number", ""))
    if not key:
        return None
    incoming_type = (details.get("case_type") or "").strip().casefold()
    candidates = ValuationCase.query.filter(
        ValuationCase.archived.is_(False),
    ).order_by(
        db.func.coalesce(ValuationCase.email_received_at, ValuationCase.created_at).asc(),
        ValuationCase.id.asc(),
    ).all()
    matches = [
        item for item in candidates
        if normalized_application_number(item.application_number) == key
    ]
    if incoming_type in {"subsequent", "revisit", "part / tranche"}:
        incoming_subject = normalized_email_subject(subject)
        return next(
            (
                item for item in matches
                if (item.case_type or "").strip().casefold() == incoming_type
                and incoming_subject
                and normalized_email_subject(item.email_subject) == incoming_subject
                and (
                    not received or not item.email_received_at
                    or abs((item.email_received_at.date() - received.date()).days) <= 1
                )
            ),
            None,
        )
    return next(
        (
            item for item in matches
            if (item.case_type or "").strip().casefold()
            not in {"subsequent", "revisit", "part / tranche"}
        ),
        None,
    )


def merge_cross_mailbox_duplicate_cases():
    """Archive only exact non-follow-up duplicate assignments and retain their files."""
    groups = {}
    for case in ValuationCase.query.filter_by(archived=False).order_by(
        db.func.coalesce(ValuationCase.email_received_at, ValuationCase.created_at).asc(),
        ValuationCase.id.asc(),
    ).all():
        key = normalized_application_number(case.application_number)
        kind = (case.case_type or "").strip().casefold()
        if not key:
            continue
        if kind in {"subsequent", "revisit", "part / tranche"}:
            received_day = case.email_received_at.date() if case.email_received_at else None
            if not received_day:
                continue
            key = (key, kind, received_day)
        else:
            key = (key, "standard")
        groups.setdefault(key, []).append(case)
    merged = 0
    for cases in groups.values():
        if len(cases) < 2:
            continue
        canonical = cases[0]
        for duplicate in cases[1:]:
            # Do not merge two genuinely different banks under an accidentally reused ID.
            if (
                canonical.bank_name and duplicate.bank_name
                and normalized_header(canonical.bank_name) != normalized_header(duplicate.bank_name)
            ):
                continue
            for field in ("customer_name", "contact_number", "property_address", "branch_name", "bank_name"):
                if not (getattr(canonical, field, "") or "").strip() and (getattr(duplicate, field, "") or "").strip():
                    setattr(canonical, field, getattr(duplicate, field))
            for asset in FileAsset.query.filter_by(case_id=duplicate.id).all():
                asset.case_id = canonical.id
            stored = safe_json(canonical.extracted_json)
            stored.setdefault("followup_emails", []).append({
                "message_id": duplicate.source_message_id or f"duplicate-case-{duplicate.id}",
                "subject": duplicate.email_subject or "Duplicate assignment from another mailbox",
                "received_at": duplicate.email_received_at.isoformat() if duplicate.email_received_at else "",
                "action": "Duplicate Gmail/Yahoo assignment merged into existing MIS case",
            })
            canonical.extracted_json = json.dumps(stored, ensure_ascii=False)
            duplicate.archived = True
            duplicate.status = f"Duplicate merged into Case #{canonical.id}"
            merged += 1
    if merged:
        db.session.commit()
    return merged


def apply_followup_to_existing_case(
    target, details, attachments, subject, received, unique_id,
    action="Merged into existing MIS case; no new row created",
    mark_for_review=False,
):
    if _message_already_recorded(target, unique_id):
        return False
    fill_if_missing = {
        "customer_name": details.get("customer_name", ""),
        "contact_number": details.get("contact_number", ""),
        "property_address": details.get("property_address", ""),
        "bank_name": details.get("bank_name", ""),
        "branch_name": details.get("branch_name", ""),
    }
    for field, value in fill_if_missing.items():
        if value and not (getattr(target, field, "") or "").strip():
            setattr(target, field, str(value).strip())
    stored = safe_json(target.extracted_json)
    stored.setdefault("followup_emails", []).append({
        "message_id": unique_id,
        "subject": subject,
        "received_at": received.isoformat() if received else "",
        "action": action,
    })
    # A reminder/review/correction message belongs to the case audit trail, not
    # the billing MIS chronology. Preserve the original assignment date/status
    # so an old case does not reappear as today's work or inflate invoice rows.
    target.extracted_json = json.dumps(stored, ensure_ascii=False)
    db.session.commit()
    store_email_attachments(target, attachments)
    return True


FOLLOWUP_ONLY_STATUSES = {
    "Correction Pending",
    "System Pending - Action Required",
    "Existing Case - New Mail Review",
}


def clean_non_billing_followups():
    """Repair legacy follow-up mutations and hide standalone action mails."""
    changed = 0
    candidates = ValuationCase.query.filter(
        ValuationCase.archived.is_(False),
        ValuationCase.source_email.isnot(None),
    ).all()
    for case in candidates:
        stored = safe_json(case.extracted_json)
        followups = stored.get("followup_emails", [])
        initial_received = stored.get("initial_email_received_at", "")
        if followups and initial_received:
            try:
                restored = datetime.fromisoformat(initial_received)
            except (TypeError, ValueError):
                restored = None
            if restored and case.email_received_at != restored:
                case.email_received_at = restored
                changed += 1
        if followups and case.status in FOLLOWUP_ONLY_STATUSES:
            complete = bool(case.application_number and case.customer_name and case.case_type)
            case.status = "New - Email" if complete else "Email Parsed - Review"
            changed += 1
        elif not followups and case.status in FOLLOWUP_ONLY_STATUSES:
            # This row was created from a correction/system-action message
            # itself. It is audit mail, not a billable valuation assignment.
            case.archived = True
            case.status = "Non-billing follow-up email"
            changed += 1
    if changed:
        db.session.commit()
    return changed


def email_case_status(details):
    if details.get("correction_mail") or details.get("correction_request_mail"):
        return "Correction Pending"
    if details.get("system_pending_mail"):
        return "System Pending - Action Required"
    if details.get("portal_case"):
        return "Portal Pending"
    required = (
        details.get("application_number"),
        details.get("customer_name"),
        details.get("case_type"),
    )
    return "New - Email" if all(required) else "Email Parsed - Review"


def mailbox_source(case):
    email = (case.source_email or "").strip().casefold()
    if email.endswith("@gmail.com"):
        return "Gmail"
    if email.endswith("@yahoo.com") or email.endswith("@yahoo.in"):
        return "Yahoo"
    return "Manual / Import" if not email else email


def concise_mis_address(value, limit=150):
    text = re.sub(r"\s+", " ", str(value or "")).strip(" ,;-")
    if not text:
        return ""
    colony = re.search(r"(?i)\bcolony\b", text)
    locality = re.search(r"(?i)\b(?:village|vill\.?|gram)\b", text)
    if colony and (not locality or colony.start() < locality.start()):
        comma = text.rfind(",", 0, colony.start())
        text = text[comma + 1:].strip() if comma >= 0 else text
    elif locality:
        text = text[locality.start():]
    end = re.search(
        r"(?i)\b(?:district|distt?\.?)[\s:.-]*[A-Za-z][A-Za-z .'-]{1,45}", text
    )
    if end:
        text = text[:end.end()]
    if len(text) > limit:
        text = text[:limit].rsplit(" ", 1)[0].rstrip(" ,;-") + "…"
    return text


def normalize_whatsapp_group_link(value):
    raw = str(value or "").strip()
    try:
        parsed = urlsplit(raw)
    except ValueError:
        return ""
    token = parsed.path.strip("/")
    if (
        parsed.scheme.casefold() == "https"
        and parsed.netloc.casefold() == "chat.whatsapp.com"
        and re.fullmatch(r"[A-Za-z0-9_-]{10,100}", token)
    ):
        return f"https://chat.whatsapp.com/{token}"
    return ""


def apply_email_details(case, details, account, subject, received, unique_id):
    existing_row = bool(case.id)
    values = {
        "application_number": details.get("application_number", ""),
        "customer_name": details.get("customer_name", ""),
        "contact_number": details.get("contact_number", ""),
        "property_address": details.get("property_address", ""),
        "bank_name": details.get("bank_name", "") or account.bank_name or "",
        "branch_name": details.get("branch_name", ""),
        "case_type": details.get("case_type", ""),
    }
    invalid_existing = {
        "application_number": {"applicant", "application", "app", "case", "lead"},
        "customer_name": {"to be reviewed", "applicant", "customer", "pending"},
    }
    for field, value in values.items():
        current = (getattr(case, field, "") or "").strip()
        invalid_current = current.casefold() in invalid_existing.get(field, set())
        if value and (not existing_row or not current or invalid_current):
            setattr(case, field, str(value).strip())
        elif invalid_current:
            setattr(case, field, "")
    case.source_email = account.email
    case.source_message_id = case.source_message_id or unique_id
    case.email_subject = case.email_subject or subject
    case.email_received_at = case.email_received_at or received
    email_managed_statuses = {
        "", "New", "New - Email", "Email Parsed - Review",
        "Correction Pending", "Portal Pending",
        "System Pending - Action Required",
        "Ignored - Not Valuation Email",
        "Ignored - Follow-up Without Initiation",
    }
    if not existing_row or (case.status or "") in email_managed_statuses:
        case.status = email_case_status(details)
    case.archived = False
    stored = safe_json(case.extracted_json)
    stored["email"] = details
    case.extracted_json = json.dumps(stored, ensure_ascii=False)


def store_email_attachments(case, attachments):
    existing = {
        (asset.filename.casefold(), len(asset.content or b""))
        for asset in FileAsset.query.filter_by(case_id=case.id).all()
    }
    stored = 0
    for attachment in attachments:
        key = (attachment["filename"].casefold(), len(attachment["content"]))
        if key in existing:
            continue
        asset_type, source_kind = quick_asset_type(attachment["filename"])
        store_asset(
            case.id, asset_type, attachment["filename"], attachment["content"],
            attachment["mime_type"], source_kind=source_kind,
            process_ai=False,
        )
        existing.add(key)
        stored += 1
    return stored


def apply_application_correction(details, existing_case, attachments):
    if not (
        details.get("correction_mail")
        and details.get("application_number")
        and details.get("customer_name")
    ):
        return False
    query = ValuationCase.query.filter(
        db.func.lower(ValuationCase.customer_name)
        == details["customer_name"].strip().lower(),
        ValuationCase.archived.is_(False),
    )
    if existing_case:
        query = query.filter(ValuationCase.id != existing_case.id)
    target = query.order_by(
        db.func.coalesce(ValuationCase.email_received_at, ValuationCase.created_at).desc()
    ).first()
    if not target:
        return False
    target.application_number = details["application_number"]
    stored = safe_json(target.extracted_json)
    stored.setdefault("corrections", []).append({
        "application_number": details["application_number"],
        "customer_name": details["customer_name"],
        "reason": "Application number correction received by email",
    })
    target.extracted_json = json.dumps(stored, ensure_ascii=False)
    target.status = "Email Correction Applied - Review"
    store_email_attachments(target, attachments)
    if existing_case:
        existing_case.archived = True
        existing_case.status = f"Correction Applied to Case #{target.id}"
    db.session.commit()
    return True


def email_fetch_folders(account):
    """Folders that can contain a real incoming assignment for a provider."""
    provider = (account.provider or "").casefold()
    address = (account.email or "").casefold()
    if provider == "gmail" or address.endswith("@gmail.com") or address.endswith("@googlemail.com"):
        # Gmail moves messages out of Inbox after a rule/archive, but retains
        # them in All Mail. Inbox remains as a compatibility fallback.
        return ['"[Gmail]/All Mail"', '"[Google Mail]/All Mail"', "INBOX"]
    if provider == "yahoo" or address.endswith("@yahoo.com"):
        # Yahoo users commonly archive case-assignment mails after reading.
        return ["INBOX", "Archive"]
    return ["INBOX"]


def imap_safe_assignment_folders(account, mail):
    """Return configured folders plus user-created incoming labels/folders.

    Bank assignment rules can move a message into a custom Gmail label or a
    Yahoo folder.  Do not scan folders that normally contain mail sent by us,
    drafts, deleted mail or junk, as those must never create MIS cases.
    """
    folders = list(email_fetch_folders(account))
    known = {folder.strip('"').casefold() for folder in folders}
    try:
        status, payload = mail.list()
    except Exception:
        return folders
    if status != "OK":
        return folders

    blocked_flags = (r"\sent", r"\draft", r"\trash", r"\junk", r"\spam")
    for raw_folder in payload or []:
        if not raw_folder:
            continue
        line = raw_folder.decode("utf-8", errors="replace") if isinstance(raw_folder, bytes) else str(raw_folder)
        if any(flag in line.casefold() for flag in blocked_flags):
            continue
        quoted = re.findall(r'"((?:\\.|[^"\\])*)"', line)
        name = (quoted[-1] if quoted else line.rsplit(" ", 1)[-1]).strip()
        if not name or name.casefold() in known:
            continue
        # IMAP requires an argument quoted when the mailbox name contains a
        # space; names received through LIST are otherwise already exact.
        folders.append(f'"{name}"' if re.search(r"\s", name) else name)
        known.add(name.casefold())
    return folders


def fetch_mis_message(mail, msg_id):
    """Fetch enough of a message for MIS parsing without downloading documents.

    The header plus first 256 KB of MIME text comfortably covers normal bank
    assignment bodies while avoiding multi-megabyte PDFs and site-photo sets.
    """
    try:
        status, msg_data = mail.fetch(
            msg_id,
            "(BODY.PEEK[HEADER] BODY.PEEK[TEXT]<0.262144>)",
        )
    except imaplib.IMAP4.error:
        status, msg_data = "BAD", []
    # Yahoo rejects the bounded multi-part FETCH command on some mailboxes.
    # Fall back to the widely supported full-message command; attachment bytes
    # remain temporary and are never stored in SVAI.
    if status != "OK":
        status, msg_data = mail.fetch(msg_id, "(RFC822)")
    if status != "OK":
        return None
    chunks = [item[1] for item in msg_data if isinstance(item, tuple) and item[1]]
    if not chunks:
        return None
    return b"\r\n".join(chunks)


def fetch_full_message(mail, msg_id):
    """Fetch a full message only when a missing MIS address needs documents."""
    status, msg_data = mail.fetch(msg_id, "(RFC822)")
    if status != "OK":
        return None
    return next(
        (item[1] for item in msg_data if isinstance(item, tuple) and item[1]),
        None,
    )


def enrich_missing_address_from_email_document(details, mail, msg_id):
    """Read supported documents temporarily; never save them as FileAssets."""
    if (details.get("property_address") or "").strip():
        return details
    raw = fetch_full_message(mail, msg_id)
    if not raw:
        return details
    message = email_lib.message_from_bytes(raw)
    readable = []
    for item in collect_email_attachments(message):
        extension = Path(item["filename"]).suffix.casefold()
        if extension not in {".pdf", ".docx", ".xlsx", ".xlsm"}:
            continue
        # Scheduled MIS fetch must stay below the production worker's memory
        # limit. Searchable document text is enough for an address fallback;
        # scanned-file OCR remains available in the explicit report workflow.
        text = extract_basic_text(
            item["filename"], item["content"], allow_ocr=False
        )
        if text:
            readable.append((item["filename"], text))
    return enrich_email_details_from_attachments(details, readable)


def archived_case_has_assignment_identity(case):
    """Recover a structured MIS assignment without trusting blank/junk mail.

    Some lender systems send a useful case table but no literal "valuation"
    phrase.  Earlier parser versions could create the structured row and a
    later fetch could archive it.  Require a real application key plus several
    independent case facts before reopening such a row.
    """
    application = normalized_application_number(case.application_number)
    if (
        len(application) < 6
        or application in {
            "APP", "APPLICANT", "APPLICATION", "CASE", "LEAD", "LOAN",
            "PROPOSAL", "REFERENCE",
        }
        or any(token in application for token in ("ISO8859", "TENOR20YRS"))
    ):
        return False
    bank = normalized_header(case.bank_name)
    real_bank = bool(bank and bank not in {"gmail", "yahoo", "googlemail"})
    has_property = bool((case.property_address or "").strip())
    facts = (
        bool((case.customer_name or "").strip()),
        bool((case.contact_number or "").strip()),
        has_property,
        real_bank,
        bool((case.branch_name or "").strip()),
        bool((case.case_type or "").strip()),
    )
    return sum(facts) >= 3 and (has_property or real_bank)


def recover_archived_assignment(case):
    if not case or not case.archived or not archived_case_has_assignment_identity(case):
        return False
    case.archived = False
    case.status = "Email Parsed - Review"
    case.updated_at = datetime.utcnow()
    db.session.commit()
    return True


def recover_structured_archived_cases(account, start_date, end_date):
    start_dt = datetime.combine(start_date, datetime.min.time())
    end_dt = datetime.combine(end_date + timedelta(days=1), datetime.min.time())
    candidates = ValuationCase.query.filter(
        ValuationCase.archived.is_(True),
        ValuationCase.source_email == account.email,
        db.func.coalesce(
            ValuationCase.email_received_at, ValuationCase.created_at
        ) >= start_dt,
        db.func.coalesce(
            ValuationCase.email_received_at, ValuationCase.created_at
        ) < end_dt,
    ).all()
    return sum(1 for case in candidates if recover_archived_assignment(case))


def _fetch_email_account_unlocked(
    account: EmailAccount, start_date=None, end_date=None,
    enrich_documents=True,
):
    created = 0
    updated = 0
    ignored = 0
    start_default, end_default = current_month_range()
    start_date = start_date or start_default
    end_date = end_date or end_default
    if start_date > end_date:
        raise ValueError("From date cannot be after To date.")
    account_id = account.id
    account_email = account.email
    host = detect_imap(account.email, account.provider, account.imap_host or "")
    password = decrypt_password(account.encrypted_password)
    # IMAP servers can silently stall while opening a TLS connection.  A
    # bounded timeout keeps one unavailable mailbox from blocking the MIS
    # refresh (and, consequently, other linked mailbox accounts).
    imap_timeout = max(5, int(os.getenv("IMAP_TIMEOUT_SECONDS", "25")))
    mail = None
    warning = ""
    try:
        mail = imaplib.IMAP4_SSL(host, 993, timeout=imap_timeout)
        mail.login(account.email, password)
        since = start_date.strftime("%d-%b-%Y")
        before = (end_date + timedelta(days=1)).strftime("%d-%b-%Y")
        message_refs = []
        scanned_folders = []
        is_gmail = (account.provider or "").casefold() == "gmail" or account.email.casefold().endswith("@gmail.com")
        for folder in imap_safe_assignment_folders(account, mail):
            status, _ = mail.select(folder)
            if status != "OK":
                continue
            status, payload = mail.search(None, "SINCE", since, "BEFORE", before)
            if status != "OK":
                continue
            scanned_folders.append(folder.strip('"'))
            for msg_id in payload[0].split():
                message_refs.append((folder, msg_id))
            # Gmail All Mail is authoritative and already includes Inbox and
            # custom-label messages. Stop after the first selectable folder so
            # the same full messages/attachments are not downloaded repeatedly.
            if is_gmail:
                break
        # Gmail's normal IMAP date search can miss a bank message when Gmail
        # has indexed it under a category/label.  Query the known LIFC sender
        # through Gmail's own search syntax as a narrow fallback.  This is
        # deliberately best-effort: a Gmail-specific search error must never
        # discard messages already found by the standard IMAP search.
        if is_gmail and mail.select('"[Gmail]/All Mail"')[0] == "OK":
            for sender_domain in ("lifl.in", "lifc.in"):
                gmail_query = (
                    f"from:{sender_domain} "
                    f"after:{(start_date - timedelta(days=1)):%Y/%m/%d} "
                    f"before:{(end_date + timedelta(days=1)):%Y/%m/%d}"
                )
                try:
                    status, payload = mail.search(
                        None, "X-GM-RAW", f'"{gmail_query}"'
                    )
                except imaplib.IMAP4.error as exc:
                    warning = (
                        "Gmail LIFC fallback search skipped; standard mailbox "
                        f"scan continued. {exc}"
                    )
                    continue
                if status == "OK":
                    for msg_id in payload[0].split():
                        message_refs.append(('"[Gmail]/All Mail"', msg_id))
        if not scanned_folders:
            return {"created": 0, "ignored": 0, "message": "Mailbox folder search failed"}

        # The IMAP search can take long enough for Render PostgreSQL to close
        # the connection that loaded ``account``. Release that idle session and
        # reacquire the account on a fresh, pre-pinged DB connection before any
        # MIS writes begin.
        db.session.remove()
        account = db.session.get(EmailAccount, account_id)
        if account is None:
            raise RuntimeError(f"Linked mailbox {account_email} is no longer available.")
        seen_message_ids = set()
        selected_folder = None
        for source_folder, msg_id in message_refs:
            # Fetch and process one message at a time. Holding a whole month of
            # MIME bodies (and Yahoo RFC822 attachment fallbacks) in a list can
            # exceed Render's 512 MB worker limit.
            if selected_folder != source_folder:
                if mail.select(source_folder)[0] != "OK":
                    continue
                selected_folder = source_folder
            raw = fetch_mis_message(mail, msg_id)
            if not raw:
                continue
            # A large Yahoo sweep can spend long enough fetching/parsing each
            # message for Render PostgreSQL to close the connection used by the
            # previous message. Drop that session before the next database
            # lookup so SQLAlchemy's pre-ping opens a healthy connection.
            db.session.remove()
            account = db.session.get(EmailAccount, account_id)
            if account is None:
                raise RuntimeError(
                    f"Linked mailbox {account_email} is no longer available."
                )
            message = email_lib.message_from_bytes(raw)
            subject = decode_header_value(message.get("Subject", ""))
            sender = decode_header_value(message.get("From", ""))
            body = latest_email_body(email_body(message))
            followup_mail = is_followup_email(subject, body)
            unique_id = message.get("Message-ID") or (
                f"{account.email}:{source_folder}:{msg_id.decode()}"
            )
            if unique_id in seen_message_ids:
                continue
            seen_message_ids.add(unique_id)
            received = None
            try:
                received = email_lib.utils.parsedate_to_datetime(message.get("Date")).replace(
                    tzinfo=None
                )
            except Exception:
                received = datetime.utcnow()
            existing_case = existing_case_for_message(
                account, unique_id, subject, received
            )
            if not is_valuation_email(subject, body, sender) and not followup_mail:
                # A repeat fetch must never make an already reviewed/accepted
                # MIS row disappear merely because a later parser version is
                # stricter.  New rejected messages are still ignored, and an
                # already archived rejection stays archived for review.
                if recover_archived_assignment(existing_case):
                    updated += 1
                elif existing_case and existing_case.source_email and not existing_case.archived:
                    existing_case.updated_at = datetime.utcnow()
                    db.session.commit()
                ignored += 1
                continue

            details = ai_extract_email(subject, body, sender)
            if not details.get("is_valuation", False) and not followup_mail:
                if recover_archived_assignment(existing_case):
                    updated += 1
                elif existing_case and existing_case.source_email and not existing_case.archived:
                    existing_case.updated_at = datetime.utcnow()
                    db.session.commit()
                ignored += 1
                continue

            non_billing_followup = bool(
                followup_mail
                or details.get("correction_mail")
                or details.get("correction_request_mail")
                or details.get("system_pending_mail")
            )

            # Stay fast for normal mail. Only when the address is absent, read
            # supported documents temporarily and discard their bytes after
            # extracting MIS text. Nothing is saved as a FileAsset.
            if (
                enrich_documents
                and details.get("is_valuation", False)
                and not non_billing_followup
            ):
                details = enrich_missing_address_from_email_document(
                    details, mail, msg_id
                )
            attachments = []
            details.pop("ai_error", None)
            if details.get("correction_mail") and apply_application_correction(
                details, existing_case, attachments
            ):
                updated += 1
                continue
            if non_billing_followup:
                if existing_case and _message_already_recorded(
                    existing_case, unique_id
                ):
                    continue
                target = existing_case_for_application(
                    details.get("application_number"), existing_case, subject
                )
                if not target and existing_case and not existing_case.archived:
                    target = existing_case
                if target:
                    if apply_followup_to_existing_case(
                        target, details, attachments, subject, received, unique_id
                    ):
                        updated += 1
                else:
                    ignored += 1
                continue
            if not details.get("is_valuation", False):
                ignored += 1
                continue
            duplicate_case = existing_case or existing_case_for_duplicate_assignment(
                details, account, subject, received
            )
            if duplicate_case:
                if existing_case and _message_already_recorded(existing_case, unique_id):
                    apply_email_details(
                        existing_case, details, account, subject, received, unique_id
                    )
                    db.session.commit()
                    updated += 1
                    continue
                if apply_followup_to_existing_case(
                    duplicate_case, details, attachments, subject, received, unique_id,
                    action="Duplicate assignment merged into existing MIS case",
                    mark_for_review=True,
                ):
                    updated += 1
                continue

            case = ValuationCase()
            apply_email_details(case, details, account, subject, received, unique_id)
            db.session.add(case)
            db.session.commit()
            store_email_attachments(case, attachments)
            if existing_case:
                updated += 1
            else:
                created += 1
        updated += recover_structured_archived_cases(account, start_date, end_date)
        updated += clean_non_billing_followups()
        account.last_fetch_at = datetime.utcnow()
        db.session.commit()
        deduplicated = merge_cross_mailbox_duplicate_cases()
        return {
            "created": created, "updated": updated, "ignored": ignored,
            "deduplicated": deduplicated,
            "message": (
                f"Fetched {start_date:%d-%m-%Y} to {end_date:%d-%m-%Y} "
                f"from {', '.join(scanned_folders)}"
            ),
            "warning": warning,
        }
    except (imaplib.IMAP4.abort, OSError) as exc:
        db.session.rollback()
        warning = f"Mailbox connection failed or ended early: {exc}"
        return {
            "created": created,
            "updated": updated,
            "ignored": ignored,
            "message": f"Partial fetch {start_date:%d-%m-%Y} to {end_date:%d-%m-%Y}",
            "warning": warning,
        }
    finally:
        if mail is not None:
            try:
                mail.logout()
            except Exception:
                pass


def fetch_email_account(
    account: EmailAccount, start_date=None, end_date=None,
    enrich_documents=True,
):
    """Run only one mailbox sweep per worker to prevent scheduler/UI overlap."""
    if not _email_fetch_lock.acquire(blocking=False):
        return {
            "created": 0,
            "updated": 0,
            "ignored": 0,
            "deduplicated": 0,
            "message": "Mailbox scan already running; the next scheduled retry will continue.",
            "warning": "Another MIS scan is already running.",
        }
    try:
        return _fetch_email_account_unlocked(
            account, start_date, end_date, enrich_documents=enrich_documents
        )
    finally:
        _email_fetch_lock.release()


def classify_photo(filename: str) -> str:
    name = Path(filename).stem.lower()
    for key, category in PHOTO_CATEGORIES.items():
        if key in name:
            return category
    return "Other Site Photo"


def extract_basic_text(filename: str, content: bytes, allow_ocr=True) -> str:
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf":
            reader = PdfReader(io.BytesIO(content))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)[:50000]
            if len(re.sub(r"\s+", "", text)) >= 80:
                return text
            return (offline_ocr_text(filename, content) or text) if allow_ocr else text
        if ext in PHOTO_EXTENSIONS:
            return offline_ocr_text(filename, content) if allow_ocr else ""
        if ext == ".docx":
            document = Document(io.BytesIO(content))
            output = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    line = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if line:
                        output.append(line)
            return "\n".join(output)[:50000]
        if ext in {".xlsx", ".xlsm"}:
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            output = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    line = " | ".join(str(v) for v in row if v not in (None, ""))
                    if line:
                        output.append(line)
            return "\n".join(output)[:50000]
    except Exception:
        return ""
    return ""


def ai_extract_document(filename: str, content: bytes, existing_text: str = "", source_kind="property_document"):
    return extract_property_asset(filename, content, existing_text, source_kind)


def store_asset(
    case_id, asset_type, filename, content, mime_type=None, category=None,
    source_kind="property_document", process_ai=True,
):
    filename = secure_filename(filename) or f"file_{secrets.token_hex(4)}"
    text = ""
    extraction = {}
    if asset_type in {"document", "visit_data"}:
        if process_ai:
            # Upload must stay lightweight on the 512 MB web worker. Even text
            # extraction can inflate compressed/scanned PDFs far beyond their
            # file size. Parse and OCR only during the explicit AI action.
            text = extract_basic_text(filename, content, allow_ocr=True)
            extraction = ai_extract_document(filename, content, text, source_kind)
    elif asset_type == "photo":
        category = category or classify_photo(filename)
    asset = FileAsset(
        case_id=case_id,
        asset_type=asset_type,
        category=category,
        filename=filename,
        mime_type=mime_type or "application/octet-stream",
        content=content,
        extracted_text=text,
        extraction_json=json.dumps(extraction, ensure_ascii=False) if extraction else None,
    )
    db.session.add(asset)
    db.session.commit()
    return asset


def safe_zip_members(zf):
    max_files = int(os.getenv("MAX_ZIP_FILES", "250"))
    max_uncompressed = int(os.getenv("MAX_ZIP_UNCOMPRESSED_MB", "40")) * 1024 * 1024
    max_single = int(os.getenv("MAX_SINGLE_UPLOAD_MB", "12")) * 1024 * 1024
    accepted = 0
    total_size = 0
    for member in zf.infolist():
        path = Path(member.filename)
        if member.is_dir() or ".." in path.parts or path.is_absolute():
            continue
        accepted += 1
        total_size += member.file_size
        if member.file_size > max_single:
            raise ValueError(
                f"{Path(member.filename).name} is too large; keep each file under "
                f"{max_single // (1024 * 1024)} MB."
            )
        if accepted > max_files or total_size > max_uncompressed:
            raise ValueError("ZIP is too large after extraction.")
        yield member


def read_upload_limited(item):
    max_bytes = int(os.getenv("MAX_SINGLE_UPLOAD_MB", "12")) * 1024 * 1024
    content = item.stream.read(max_bytes + 1)
    if len(content) > max_bytes:
        raise ValueError(
            f"{secure_filename(item.filename)} is too large; keep each file under "
            f"{max_bytes // (1024 * 1024)} MB."
        )
    return content


def valuation_calculation(form):
    def f(name, default=0):
        try:
            return float(form.get(name, default) or default)
        except Exception:
            return float(default)

    data = {
        "land_area": f("land_area"),
        "land_rate": f("land_rate"),
        "builtup_area": f("builtup_area"),
        "construction_rate": f("construction_rate"),
        "age_years": f("age_years"),
        "depreciation_percent": f("depreciation_percent"),
        "govt_land_rate": f("govt_land_rate"),
        "govt_construction_rate": f("govt_construction_rate"),
        "conservative_percent": f("conservative_percent", 100),
        "distress_percent": f("distress_percent", 80),
        "remarks": form.get("remarks", ""),
    }
    data["land_value"] = data["land_area"] * data["land_rate"]
    data["gross_building_value"] = data["builtup_area"] * data["construction_rate"]
    data["depreciation_amount"] = data["gross_building_value"] * data["depreciation_percent"] / 100
    data["net_building_value"] = data["gross_building_value"] - data["depreciation_amount"]
    data["market_value"] = data["land_value"] + data["net_building_value"]
    data["conservative_value"] = data["market_value"] * data["conservative_percent"] / 100
    data["distress_value"] = data["market_value"] * data["distress_percent"] / 100
    data["govt_value"] = (
        data["land_area"] * data["govt_land_rate"] +
        data["builtup_area"] * data["govt_construction_rate"]
    )
    return data


def numeric_from_value(value, default=0):
    if isinstance(value, (int, float)):
        return float(value)
    match = re.search(r"-?\d[\d,]*(?:\.\d+)?", str(value or ""))
    if not match:
        return float(default)
    try:
        return float(match.group(0).replace(",", ""))
    except ValueError:
        return float(default)


def valuation_defaults_from_profile(profile):
    """Apply the valuer's standing rules without inventing source facts."""
    profile = profile or {}
    stage_text = " ".join(str(profile.get(key) or "") for key in (
        "construction_stage", "construction_quality", "structure_type", "remarks",
    )).casefold()
    if any(token in stage_text for token in ("plinth", "foundation", "dpc")):
        construction_rate, conservative_percent = 300.0, 30.0
    elif any(token in stage_text for token in (
        "without plaster", "unplastered", "brick work", "brickwork", "bare brick",
    )):
        construction_rate, conservative_percent = 700.0, 70.0
    elif any(token in stage_text for token in (
        "complete", "completed", "plaster", "finished", "ready",
    )):
        construction_rate, conservative_percent = 1000.0, 100.0
    else:
        construction_rate, conservative_percent = 0.0, 100.0

    age_years = numeric_from_value(profile.get("property_age_years"))
    if not age_years:
        construction_year = numeric_from_value(profile.get("construction_year"))
        current_year = datetime.now(APP_TIMEZONE).year
        if 1900 <= construction_year <= current_year:
            age_years = float(current_year - int(construction_year))
    if age_years:
        profile["property_age_years"] = age_years
        profile["residual_age_years"] = max(0.0, 60.0 - age_years)

    govt_land_rate = numeric_from_value(profile.get("govt_land_rate"))
    market_land_rate = numeric_from_value(profile.get("land_rate"))
    if not market_land_rate and govt_land_rate:
        market_land_rate = govt_land_rate * 2

    return {
        "land_area": numeric_from_value(profile.get("land_area_as_per_docs")),
        "land_rate": market_land_rate,
        "builtup_area": numeric_from_value(profile.get("builtup_area_as_per_site")),
        "construction_rate": (
            numeric_from_value(profile.get("construction_rate"))
            or construction_rate
        ),
        "age_years": age_years,
        "depreciation_percent": 0.0,
        "govt_land_rate": govt_land_rate,
        "govt_construction_rate": numeric_from_value(
            profile.get("govt_construction_rate")
        ),
        "conservative_percent": conservative_percent,
        "distress_percent": 80.0,
    }


def report_mapping(case: ValuationCase, valuation: Valuation):
    extracted = safe_json(case.extracted_json)
    values = {
        "CASE_ID": case.id,
        "APPLICATION_NUMBER": case.application_number or "",
        "CUSTOMER_NAME": case.customer_name or "",
        "CONTACT_NUMBER": case.contact_number or "",
        "PROPERTY_ADDRESS": case.property_address or "",
        "BANK_NAME": case.bank_name or "",
        "BRANCH_NAME": case.branch_name or "",
        "CASE_TYPE": case.case_type or "",
        "STATUS": case.status or "",
        "VISIT_BY": case.visit_by or "",
        "REPORT_DATE": datetime.now().strftime("%d-%m-%Y"),
        "LAND_AREA": valuation.land_area,
        "LAND_RATE": valuation.land_rate,
        "LAND_VALUE": valuation.land_value,
        "BUILTUP_AREA": valuation.builtup_area,
        "CONSTRUCTION_RATE": valuation.construction_rate,
        "GROSS_BUILDING_VALUE": valuation.gross_building_value,
        "DEPRECIATION_PERCENT": valuation.depreciation_percent,
        "DEPRECIATION_AMOUNT": valuation.depreciation_amount,
        "NET_BUILDING_VALUE": valuation.net_building_value,
        "MARKET_VALUE": valuation.market_value,
        "CONSERVATIVE_VALUE": valuation.conservative_value,
        "DISTRESS_VALUE": valuation.distress_value,
        "GOVT_VALUE": valuation.govt_value,
        "REMARKS": valuation.remarks or "",
    }
    for key, value in extracted.items():
        values[f"EXTRACTED_{str(key).upper()}"] = value
    return values


def fill_template(content: bytes, mapping: dict) -> bytes:
    keep_vba = False
    wb = load_workbook(io.BytesIO(content), keep_vba=keep_vba)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    value = cell.value
                    for key, replacement in mapping.items():
                        value = value.replace(f"{{{{{key}}}}}", str(replacement if replacement is not None else ""))
                    cell.value = value
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def generic_report(case: ValuationCase, valuation: Valuation, assets):
    wb = Workbook()
    ws = wb.active
    ws.title = "Valuation Report"
    ws.merge_cells("A1:D1")
    ws["A1"] = "SVAI - SAKSHAM ASSOCIATE PROPERTY VALUATION REPORT"
    ws["A1"].font = Font(size=15, bold=True)
    ws["A1"].alignment = Alignment(horizontal="center")
    thin = Side(style="thin", color="999999")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    mapping = report_mapping(case, valuation)
    ordered = [
        ("Application Number", mapping["APPLICATION_NUMBER"], "Bank", mapping["BANK_NAME"]),
        ("Customer Name", mapping["CUSTOMER_NAME"], "Branch", mapping["BRANCH_NAME"]),
        ("Contact Number", mapping["CONTACT_NUMBER"], "Case Type", mapping["CASE_TYPE"]),
        ("Property Address", mapping["PROPERTY_ADDRESS"], "Status", mapping["STATUS"]),
        ("Land Area", mapping["LAND_AREA"], "Land Rate", mapping["LAND_RATE"]),
        ("Land Value", mapping["LAND_VALUE"], "Built-up Area", mapping["BUILTUP_AREA"]),
        ("Construction Rate", mapping["CONSTRUCTION_RATE"], "Gross Building Value", mapping["GROSS_BUILDING_VALUE"]),
        ("Depreciation %", mapping["DEPRECIATION_PERCENT"], "Depreciation Amount", mapping["DEPRECIATION_AMOUNT"]),
        ("Net Building Value", mapping["NET_BUILDING_VALUE"], "Market Value", mapping["MARKET_VALUE"]),
        ("Conservative Value", mapping["CONSERVATIVE_VALUE"], "Distress Value", mapping["DISTRESS_VALUE"]),
        ("Government Value", mapping["GOVT_VALUE"], "Report Date", mapping["REPORT_DATE"]),
        ("Remarks", mapping["REMARKS"], "", ""),
    ]
    for r, row in enumerate(ordered, 3):
        for c, value in enumerate(row, 1):
            cell = ws.cell(r, c, value)
            cell.border = border
            if c in (1, 3):
                cell.font = Font(bold=True)
                cell.fill = PatternFill("solid", fgColor="DDEBF7")
    start = 3 + len(ordered) + 2
    ws.merge_cells(start_row=start, start_column=1, end_row=start, end_column=4)
    ws.cell(start, 1, "UPLOADED DOCUMENTS / PHOTOS").font = Font(bold=True)
    for c, value in enumerate(["Type", "Category", "Filename", "Extraction"], 1):
        ws.cell(start + 1, c, value).font = Font(bold=True)
        ws.cell(start + 1, c).border = border
    for r, asset in enumerate(assets, start + 2):
        extraction = safe_json(asset.extraction_json)
        values = [asset.asset_type, asset.category or "", asset.filename, json.dumps(extraction, ensure_ascii=False)]
        for c, value in enumerate(values, 1):
            ws.cell(r, c, value).border = border
            ws.cell(r, c).alignment = Alignment(wrap_text=True, vertical="top")
    for col, width in zip("ABCD", [24, 34, 26, 55]):
        ws.column_dimensions[col].width = width
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


@app.before_request
def protect_requests():
    session.setdefault("_csrf_token", secrets.token_urlsafe(32))
    if request.method == "POST":
        submitted = request.form.get("_csrf_token") or request.headers.get("X-CSRF-Token", "")
        if not secrets.compare_digest(str(submitted), str(session["_csrf_token"])):
            abort(400, description="Invalid or missing CSRF token.")


@app.context_processor
def template_security():
    def case_km(case):
        stored = safe_json(case.extracted_json)
        profile = stored.get("case_profile") or stored.get("email") or stored
        return profile.get("distance_from_branch", profile.get("km", ""))
    return {
        "csrf_token": session.get("_csrf_token", ""),
        "case_km": case_km,
        "mailbox_source": mailbox_source,
        "concise_mis_address": concise_mis_address,
    }


@app.before_request
def ensure_setup():
    global _runtime_setup_done
    admin_email = os.getenv("ADMIN_EMAIL", "sakshamvaluer@yahoo.com").lower()
    if _runtime_setup_done == admin_email:
        return
    with _setup_lock:
        if _runtime_setup_done == admin_email:
            return
        admin_password = os.getenv("ADMIN_PASSWORD", "ChangeMe123!")
        if not User.query.filter_by(email=admin_email).first():
            db.session.add(User(
                email=admin_email,
                password_hash=generate_password_hash(admin_password),
                name="Saksham Associate Admin",
                role="admin",
            ))
            db.session.commit()
        if SEED_TEMPLATES_DIR.exists():
            seeds = {
                "DCB.xlsx": "DCB Bank",
                "SBFC.xlsx": "SBFC Finance",
                "Laxmi India.xlsx": "Laxmi India Finance",
                "Ummeed.docx": "Ummeed Housing Finance",
            }
            added = False
            for filename, bank_name in seeds.items():
                source = SEED_TEMPLATES_DIR / filename
                if not source.exists():
                    continue
                exists = FileAsset.query.filter_by(asset_type="template", filename=filename).first()
                if exists:
                    continue
                mime_type = (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    if source.suffix.lower() == ".docx"
                    else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
                db.session.add(FileAsset(
                    asset_type="template", category=bank_name, filename=filename,
                    mime_type=mime_type, content=source.read_bytes(),
                ))
                added = True
            if added:
                db.session.commit()
        _runtime_setup_done = admin_email


@app.route("/health")
def health():
    return jsonify({
        "status": "ok",
        "database": "connected",
        "ai_enabled": ai_enabled(),
        "document_processing_mode": (
            "Paid ChatGPT + Local Fallback"
            if document_ai_enabled() else "Free Local"
        ),
        "ai_provider": "OpenAI",
        "ai_model": OPENAI_MODEL,
        "time": datetime.utcnow().isoformat(),
    })


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        submitted_email = request.form.get("email", "").strip().lower()
        submitted_password = request.form.get("password", "")
        user = User.query.filter_by(email=submitted_email).first()
        authenticated = bool(
            user and check_password_hash(user.password_hash, submitted_password)
        )
        configured_admin_email = os.getenv(
            "ADMIN_EMAIL", "sakshamvaluer@yahoo.com"
        ).strip().lower()
        configured_admin_password = os.getenv("ADMIN_PASSWORD", "ChangeMe123!")
        if (
            user
            and not authenticated
            and submitted_email == configured_admin_email
            and configured_admin_password
            and secrets.compare_digest(submitted_password, configured_admin_password)
        ):
            # Persistent databases can retain an older password hash after the
            # Render ADMIN_PASSWORD secret is changed. Treat the configured
            # secret as the recovery authority and repair the stored hash.
            user.password_hash = generate_password_hash(configured_admin_password)
            db.session.commit()
            authenticated = True
        if authenticated:
            session.clear()
            session["_csrf_token"] = secrets.token_urlsafe(32)
            session["user_id"] = user.id
            session["role"] = user.role
            session["name"] = user.name
            return redirect(url_for("dashboard"))
        flash("Email or password is incorrect.", "error")
    return render_template("login.html")


@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        now = time.time()
        if now - float(session.get("password_reset_last_sent", 0) or 0) < 60:
            flash("Naya code mangne se pehle 60 seconds wait karein.", "error")
            return render_template("forgot_password.html")
        address = request.form.get("email", "").strip().lower()
        user = User.query.filter_by(email=address).first()
        if not user:
            flash("Registered SVAI email nahi mili.", "error")
            return render_template("forgot_password.html")
        active_accounts = EmailAccount.query.filter_by(active=True).order_by(
            EmailAccount.created_at.asc()
        ).all()
        # Prefer the recipient's own mailbox, but do not let one stale app
        # password prevent recovery when another linked mailbox can send.
        accounts = sorted(
            active_accounts,
            key=lambda candidate: candidate.email.strip().lower() != address,
        )
        recovery_code = os.getenv("ADMIN_RECOVERY_CODE", "").strip()
        recovery_available = (
            address == os.getenv(
                "ADMIN_EMAIL", "sakshamvaluer@yahoo.com"
            ).strip().lower()
            and len(recovery_code) >= 12
        )
        if not accounts:
            if recovery_available:
                session["password_reset"] = {
                    "user_id": user.id,
                    "email": user.email,
                    "code_hash": generate_password_hash(recovery_code),
                    "expires_at": now + 900,
                    "attempts": 0,
                    "recovery": True,
                }
                session["password_reset_last_sent"] = now
                flash(
                    "Linked mailbox available nahi hai. Administrator recovery "
                    "code se password reset karein.",
                    "success",
                )
                return redirect(url_for("reset_password"))
            flash(
                "Password reset code bhejne ke liye pehle se linked Gmail/Yahoo "
                "account chahiye. Local administrator se reset karwayein.",
                "error",
            )
            return render_template("forgot_password.html")
        code = f"{secrets.randbelow(1_000_000):06d}"
        delivery_error = None
        for account in accounts:
            try:
                send_password_reset_code(account, user.email, code)
                delivery_error = None
                break
            except Exception as exc:
                app.logger.warning(
                    "Password reset delivery failed via %s: %s",
                    account.email,
                    type(exc).__name__,
                )
                delivery_error = exc
        if delivery_error is not None:
            if recovery_available:
                session["password_reset"] = {
                    "user_id": user.id,
                    "email": user.email,
                    "code_hash": generate_password_hash(recovery_code),
                    "expires_at": now + 900,
                    "attempts": 0,
                    "recovery": True,
                }
                session["password_reset_last_sent"] = now
                flash(
                    "Email service unavailable hai. Administrator recovery "
                    "code se password reset karein.",
                    "success",
                )
                return redirect(url_for("reset_password"))
            flash(
                "Reset code email nahi ho paya. Internet aur linked Gmail/Yahoo "
                "app password check karke dobara try karein.",
                "error",
            )
            return render_template("forgot_password.html")
        session["password_reset"] = {
            "user_id": user.id,
            "email": user.email,
            "code_hash": generate_password_hash(code),
            "expires_at": now + 600,
            "attempts": 0,
        }
        session["password_reset_last_sent"] = now
        flash("6-digit reset code registered email par bhej diya gaya hai.", "success")
        return redirect(url_for("reset_password"))
    return render_template("forgot_password.html")


@app.route("/reset-password", methods=["GET", "POST"])
def reset_password():
    state = session.get("password_reset")
    if not state or float(state.get("expires_at", 0) or 0) < time.time():
        session.pop("password_reset", None)
        flash("Reset code expire ho gaya. Naya code mangayein.", "error")
        return redirect(url_for("forgot_password"))
    if request.method == "POST":
        code = request.form.get("code", "").strip()
        new_password = request.form.get("new_password", "")
        confirm_password = request.form.get("confirm_password", "")
        if int(state.get("attempts", 0)) >= 5:
            session.pop("password_reset", None)
            flash("Bahut zyada galat attempts. Naya reset code mangayein.", "error")
            return redirect(url_for("forgot_password"))
        if not check_password_hash(state.get("code_hash", ""), code):
            state["attempts"] = int(state.get("attempts", 0)) + 1
            session["password_reset"] = state
            flash("Reset code sahi nahi hai.", "error")
            return render_template(
                "reset_password.html", reset_email=state.get("email", "")
            )
        if len(new_password) < 8:
            flash("Naya password kam se kam 8 characters ka hona chahiye.", "error")
            return render_template(
                "reset_password.html", reset_email=state.get("email", "")
            )
        if new_password != confirm_password:
            flash("Dono naye passwords same nahi hain.", "error")
            return render_template(
                "reset_password.html", reset_email=state.get("email", "")
            )
        user = db.session.get(User, int(state["user_id"]))
        if not user:
            session.pop("password_reset", None)
            flash("Reset request valid nahi rahi. Dobara try karein.", "error")
            return redirect(url_for("forgot_password"))
        user.password_hash = generate_password_hash(new_password)
        db.session.commit()
        session.clear()
        flash("Password reset ho gaya. Ab naye password se login karein.", "success")
        return redirect(url_for("login"))
    return render_template(
        "reset_password.html", reset_email=state.get("email", "")
    )


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/settings", methods=["GET", "POST"])
@login_required
def settings_page():
    if request.method == "POST":
        api_key = request.form.get("openai_api_key", "").strip()
        new_password = request.form.get("new_password", "")
        paid_document_mode = (
            request.form.get("enable_paid_document_ai", "") == "true"
        )
        changed = False

        if api_key:
            if not (api_key.startswith("sk-") and len(api_key) > 20):
                flash("OpenAI API key sahi nahi lag rahi. Key sk- se shuru honi chahiye.", "error")
                return render_template(
                    "settings.html",
                    ai_enabled=ai_enabled(),
                    ai_model=OPENAI_MODEL,
                    document_ai_enabled=document_ai_enabled(),
                )
            update_env_file({"OPENAI_API_KEY": api_key})
            configure_openai(api_key)
            changed = True

        if paid_document_mode and not ai_enabled():
            flash(
                "Paid ChatGPT document reading on karne se pehle valid OpenAI "
                "API key save karein. Free Local Mode abhi bhi chalta rahega.",
                "error",
            )
            return render_template(
                "settings.html",
                ai_enabled=ai_enabled(),
                ai_model=OPENAI_MODEL,
                document_ai_enabled=document_ai_enabled(),
            )
        requested_mode = "true" if paid_document_mode else "false"
        current_mode = (
            "true"
            if os.getenv("OPENAI_DOCUMENT_EXTRACTION", "true").lower() == "true"
            else "false"
        )
        if requested_mode != current_mode:
            update_env_file({"OPENAI_DOCUMENT_EXTRACTION": requested_mode})
            os.environ["OPENAI_DOCUMENT_EXTRACTION"] = requested_mode
            changed = True

        if new_password:
            if len(new_password) < 8:
                flash("Naya password kam se kam 8 characters ka hona chahiye.", "error")
                return render_template(
                    "settings.html",
                    ai_enabled=ai_enabled(),
                    ai_model=OPENAI_MODEL,
                    document_ai_enabled=document_ai_enabled(),
                )
            user = db.session.get(User, session["user_id"])
            user.password_hash = generate_password_hash(new_password)
            db.session.commit()
            changed = True

        if changed:
            flash("Settings save ho gayi hain.", "success")
        else:
            flash("Koi nayi setting enter nahi ki gayi.", "error")
        return redirect(url_for("settings_page"))

    email_case_ids = db.session.query(ValuationCase.id).filter(
        ValuationCase.source_email.isnot(None)
    )
    email_document_count, email_document_bytes = db.session.query(
        db.func.count(FileAsset.id),
        db.func.coalesce(db.func.sum(db.func.length(FileAsset.content)), 0),
    ).filter(
        FileAsset.asset_type == "document",
        FileAsset.case_id.in_(email_case_ids),
    ).one()
    return render_template(
        "settings.html",
        ai_enabled=ai_enabled(),
        ai_model=OPENAI_MODEL,
        document_ai_enabled=document_ai_enabled(),
        email_document_count=email_document_count,
        email_document_mb=float(email_document_bytes or 0) / (1024 * 1024),
    )


@app.route("/settings/cleanup-email-documents", methods=["POST"])
@login_required
def cleanup_email_documents():
    email_case_ids = db.session.query(ValuationCase.id).filter(
        ValuationCase.source_email.isnot(None)
    )
    documents = FileAsset.query.filter(
        FileAsset.asset_type == "document",
        FileAsset.case_id.in_(email_case_ids),
    )
    count, total_bytes = db.session.query(
        db.func.count(FileAsset.id),
        db.func.coalesce(db.func.sum(db.func.length(FileAsset.content)), 0),
    ).filter(
        FileAsset.asset_type == "document",
        FileAsset.case_id.in_(email_case_ids),
    ).one()
    documents.delete(synchronize_session=False)
    db.session.commit()
    flash(
        f"{count} email-case document(s) delete hue; "
        f"lagbhag {total_bytes / (1024 * 1024):.1f} MB database space free hui. "
        "Gmail/Yahoo mails aur MIS cases safe hain.",
        "success",
    )
    return redirect(url_for("settings_page"))


def filter_cases_by_dates(query, start_date, end_date):
    start_dt = datetime.combine(start_date, datetime.min.time())
    end_dt = datetime.combine(end_date + timedelta(days=1), datetime.min.time())
    return query.filter(db.or_(
        db.and_(
            ValuationCase.email_received_at.isnot(None),
            ValuationCase.email_received_at >= start_dt,
            ValuationCase.email_received_at < end_dt,
        ),
        db.and_(
            ValuationCase.email_received_at.is_(None),
            ValuationCase.created_at >= start_dt,
            ValuationCase.created_at < end_dt,
        ),
    ))


def filter_billing_ready_cases(query):
    """Keep incomplete assignments visible in Review Queue, not billing MIS."""
    return query.filter(db.or_(
        ValuationCase.status.is_(None),
        ValuationCase.status != "Email Parsed - Review",
    ))


@app.route("/")
@login_required
def dashboard():
    clean_non_billing_followups()
    search = request.args.get("q", "").strip()
    include_archived = request.args.get("archived") == "1"
    show_review = request.args.get("review") == "1"
    default_from, default_to = current_month_range()
    date_from = parse_iso_date(request.args.get("from"), default_from)
    date_to = parse_iso_date(request.args.get("to"), default_to)
    if date_from > date_to:
        date_from, date_to = date_to, date_from
    query = ValuationCase.query
    if not include_archived:
        query = query.filter_by(archived=False)
        if show_review:
            query = query.filter(ValuationCase.status == "Email Parsed - Review")
        else:
            query = filter_billing_ready_cases(query)
    query = filter_cases_by_dates(query, date_from, date_to)
    if search:
        pattern = f"%{search}%"
        query = query.filter(db.or_(
            ValuationCase.application_number.ilike(pattern),
            ValuationCase.customer_name.ilike(pattern),
            ValuationCase.property_address.ilike(pattern),
            ValuationCase.bank_name.ilike(pattern),
        ))
    cases = query.order_by(
        db.func.coalesce(ValuationCase.email_received_at, ValuationCase.created_at).desc()
    ).limit(2000).all()
    stats = {
        "cases": filter_billing_ready_cases(
            ValuationCase.query.filter_by(archived=False)
        ).count(),
        "review_cases": ValuationCase.query.filter_by(
            archived=False, status="Email Parsed - Review"
        ).count(),
        "reports": FileAsset.query.filter_by(asset_type="report").count(),
    }
    return render_template(
        "dashboard.html", cases=cases, stats=stats, search=search,
        include_archived=include_archived, show_review=show_review,
        date_from=date_from, date_to=date_to,
        engineers=SiteEngineer.query.filter_by(active=True).order_by(SiteEngineer.name).all(),
        whatsapp_groups=WhatsAppGroup.query.filter_by(active=True).order_by(WhatsAppGroup.name).all(),
        ai_enabled=ai_enabled(), ai_model=OPENAI_MODEL,
    )


@app.route("/reports")
@login_required
def reports_page():
    reports = FileAsset.query.options(defer(FileAsset.content)).filter_by(asset_type="report").order_by(
        FileAsset.created_at.desc()
    ).all()
    case_ids = {item.case_id for item in reports if item.case_id}
    cases = {
        item.id: item
        for item in ValuationCase.query.filter(ValuationCase.id.in_(case_ids)).all()
    } if case_ids else {}
    return render_template("reports.html", reports=reports, cases=cases)


@app.route("/reports/<int:report_id>/delete", methods=["POST"])
@login_required
def delete_report(report_id):
    report = FileAsset.query.filter_by(id=report_id, asset_type="report").first_or_404()
    filename = report.filename
    db.session.delete(report)
    db.session.commit()
    flash(f"Generated report {filename} delete ho gayi.", "success")
    return redirect(url_for("reports_page"))


@app.route("/reports/cleanup-inputs", methods=["POST"])
@login_required
def cleanup_report_inputs():
    report_case_ids = {
        case_id for (case_id,) in db.session.query(FileAsset.case_id).filter(
            FileAsset.asset_type == "report",
            FileAsset.case_id.isnot(None),
        ).distinct().all()
    }
    deleted = 0
    if report_case_ids:
        deleted = FileAsset.query.filter(
            FileAsset.case_id.in_(report_case_ids),
            FileAsset.asset_type.in_({"document", "photo", "visit_data", "case_template"}),
        ).delete(synchronize_session=False)
        db.session.commit()
    flash(f"{deleted} completed-report upload(s) storage se remove kiye gaye.", "success")
    return redirect(url_for("reports_page"))


@app.route("/site-engineers", methods=["GET", "POST"])
@login_required
def site_engineers():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        mobile = re.sub(r"\D", "", request.form.get("mobile_number", ""))
        area = request.form.get("area", "").strip()
        if len(mobile) == 10:
            mobile = "91" + mobile
        if not name or not 11 <= len(mobile) <= 15:
            flash("Engineer name aur valid WhatsApp number enter karein.", "error")
        else:
            engineer = SiteEngineer.query.filter_by(mobile_number=mobile).first()
            if engineer:
                engineer.name, engineer.area, engineer.active = name, area, True
            else:
                db.session.add(SiteEngineer(name=name, mobile_number=mobile, area=area))
            db.session.commit()
            flash(f"{name} site engineer saved.", "success")
        return redirect(url_for("site_engineers"))
    return render_template(
        "site_engineers.html",
        engineers=SiteEngineer.query.order_by(SiteEngineer.name).all(),
        whatsapp_groups=WhatsAppGroup.query.order_by(WhatsAppGroup.name).all(),
    )


@app.route("/whatsapp-groups", methods=["POST"])
@login_required
def save_whatsapp_group():
    name = request.form.get("name", "").strip()
    area = request.form.get("area", "").strip()
    invite_url = normalize_whatsapp_group_link(request.form.get("invite_url", ""))
    if not invite_url:
        flash("Valid WhatsApp group invite link enter karein.", "error")
    elif not name:
        flash("WhatsApp group name enter karein.", "error")
    else:
        group = WhatsAppGroup.query.filter_by(invite_url=invite_url).first()
        if group:
            group.name, group.area, group.active = name, area, True
        else:
            db.session.add(WhatsAppGroup(name=name, area=area, invite_url=invite_url))
        db.session.commit()
        flash(f"{name} WhatsApp group saved.", "success")
    return redirect(url_for("site_engineers"))


@app.route("/whatsapp-groups/<int:group_id>/delete", methods=["POST"])
@login_required
def delete_whatsapp_group(group_id):
    group = WhatsAppGroup.query.get_or_404(group_id)
    group.active = False
    db.session.commit()
    flash(f"{group.name} group list se removed.", "success")
    return redirect(url_for("site_engineers"))


@app.route("/site-engineers/<int:engineer_id>/delete", methods=["POST"])
@login_required
def delete_site_engineer(engineer_id):
    engineer = SiteEngineer.query.get_or_404(engineer_id)
    engineer.active = False
    db.session.commit()
    flash(f"{engineer.name} engineer list se removed.", "success")
    return redirect(url_for("site_engineers"))


@app.route("/cases/<int:case_id>/initiate-visit", methods=["POST"])
@login_required
def initiate_visit(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    recipient = request.form.get("recipient", "").strip()
    kind, separator, raw_id = recipient.partition(":")
    if not separator or not raw_id.isdigit():
        flash("Pehle saved engineer ya area group select karein.", "error")
        return redirect(request.referrer or url_for("dashboard"))
    message = "\n".join([
        "New Site Visit - SVAI",
        f"Customer: {case.customer_name or 'Not available'}",
        f"Customer Mobile: {case.contact_number or 'Not available'}",
        f"Application No: {case.application_number or 'Not available'}",
        f"Bank: {case.bank_name or 'Not available'}",
        f"Case Type: {case.case_type or 'Not available'}",
        f"Address: {concise_mis_address(case.property_address) or 'Not available'}",
        f"Branch: {case.branch_name or 'Not available'}",
    ])
    if kind == "engineer":
        engineer = SiteEngineer.query.filter_by(id=int(raw_id), active=True).first()
        if not engineer:
            flash("Selected site engineer available nahi hai.", "error")
            return redirect(url_for("dashboard"))
        case.visit_by = engineer.name
        db.session.commit()
        return redirect(f"https://wa.me/{engineer.mobile_number}?text={quote(message)}")
    if kind == "group":
        group = WhatsAppGroup.query.filter_by(id=int(raw_id), active=True).first()
        if not group:
            flash("Selected WhatsApp group available nahi hai.", "error")
            return redirect(url_for("dashboard"))
        case.visit_by = group.name
        db.session.commit()
        return redirect(f"https://wa.me/?text={quote(message)}")
    flash("Valid engineer ya WhatsApp group select karein.", "error")
    return redirect(url_for("dashboard"))


@app.route("/mis/import", methods=["POST"])
@login_required
def import_mis():
    upload = request.files.get("mis_file")
    if not upload or not upload.filename:
        flash("MIS .xlsx file select karein.", "error")
        return redirect(url_for("dashboard"))
    if Path(upload.filename).suffix.casefold() != ".xlsx":
        flash("MIS import ke liye .xlsx file hi upload karein.", "error")
        return redirect(url_for("dashboard"))
    try:
        rows = mis_import_rows(upload)
    except ValueError as exc:
        flash(str(exc), "error")
        return redirect(url_for("dashboard"))

    created = 0
    matched = 0
    active_cases = ValuationCase.query.filter_by(archived=False).all()
    for row in rows:
        app_key = normalized_application_number(row.get("application_number"))
        received_at = row.get("received_at")
        target = next((
            case for case in active_cases
            if app_key and normalized_application_number(case.application_number) == app_key
        ), None)
        if target is None and not app_key:
            customer_key = re.sub(r"\s+", " ", str(row.get("customer_name") or "")).strip().casefold()
            bank_key = re.sub(r"\s+", " ", str(row.get("bank_name") or "")).strip().casefold()
            target = next((
                case for case in active_cases
                if customer_key
                and re.sub(r"\s+", " ", str(case.customer_name or "")).strip().casefold() == customer_key
                and (not bank_key or re.sub(r"\s+", " ", str(case.bank_name or "")).strip().casefold() == bank_key)
                and received_at and (case.email_received_at or case.created_at).date() == received_at.date()
            ), None)
        if target is None:
            target = ValuationCase(
                email_received_at=received_at,
                status=str(row.get("status") or "Imported MIS").strip(),
                extracted_json=json.dumps({
                    "mis_import": {"distance_from_branch": row.get("distance", "")}
                }),
            )
            db.session.add(target)
            active_cases.append(target)
            created += 1
        else:
            matched += 1
        for field in (
            "application_number", "customer_name", "contact_number", "bank_name",
            "case_type", "property_address", "visit_by", "branch_name",
        ):
            value = str(row.get(field) or "").strip()
            if value and not str(getattr(target, field) or "").strip():
                setattr(target, field, value)
    db.session.commit()
    deduplicated = merge_cross_mailbox_duplicate_cases()
    flash(
        f"MIS import complete: {created} case(s) added, {matched} existing matched. "
        f"{deduplicated} repeated case(s) merged. Duplicate application rows nahi banayi gayi.",
        "success",
    )
    return redirect(url_for("dashboard"))


@app.route("/cases/new", methods=["GET", "POST"])
@login_required
def new_case():
    if request.method == "POST":
        application_number = request.form.get("application_number", "").strip()
        if not application_number:
            flash("Application number enter karein.", "error")
            return render_template("new_case.html")
        existing = ValuationCase.query.filter(
            db.func.lower(ValuationCase.application_number) == application_number.lower(),
            ValuationCase.archived.is_(False),
        ).first()
        if existing:
            flash("Ye application number pehle se hai; existing file khol di gayi.", "success")
            return redirect(url_for("case_detail", case_id=existing.id))
        portal_case = request.form.get("portal_case") == "1"
        case = ValuationCase(
            application_number=application_number,
            customer_name=request.form.get("customer_name", "").strip(),
            contact_number=request.form.get("contact_number", "").strip(),
            property_address=request.form.get("property_address", "").strip(),
            bank_name=request.form.get("bank_name", "").strip(),
            branch_name=request.form.get("branch_name", "").strip(),
            case_type=request.form.get("case_type", "").strip(),
            visit_by=request.form.get("visit_by", "").strip(),
            status="Portal Pending" if portal_case else "Files Pending",
        )
        db.session.add(case)
        db.session.commit()
        return redirect(url_for("case_detail", case_id=case.id))
    return render_template("new_case.html")


@app.route("/make-report", methods=["GET", "POST"])
@login_required
def make_report():
    master_templates = FileAsset.query.filter_by(asset_type="template").order_by(
        FileAsset.category, FileAsset.created_at.desc()
    ).all()
    bank_names = []
    for item in master_templates:
        name = (item.category or "").strip()
        if name and name not in bank_names:
            bank_names.append(name)
    if request.method == "POST":
        application_number = request.form.get("application_number", "").strip()
        bank_name = request.form.get("bank_name", "").strip()
        if not application_number:
            flash("Application number enter karein.", "error")
            return render_template(
                "make_report.html", bank_names=bank_names,
                application_number=application_number, bank_name=bank_name,
            )
        application_key = normalized_application_number(application_number)
        case = next((
            item for item in ValuationCase.query.filter_by(archived=False).all()
            if normalized_application_number(item.application_number) == application_key
        ), None)
        if case is None:
            case = ValuationCase(
                application_number=application_number,
                bank_name=bank_name,
                status="Files Pending",
            )
            db.session.add(case)
        elif bank_name:
            case.bank_name = bank_name
        db.session.commit()
        if not matching_master_template(case.bank_name):
            flash(
                "Selected bank ka master report format nahi mila. Report Formats me ek baar upload karein.",
                "error",
            )
        return redirect(url_for("case_detail", case_id=case.id, report=1))
    return render_template(
        "make_report.html", bank_names=bank_names,
        application_number="", bank_name="",
    )


@app.route("/cases/<int:case_id>")
@login_required
def case_detail(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    valuation = Valuation.query.filter_by(case_id=case_id).first()
    assets = FileAsset.query.options(defer(FileAsset.content)).filter_by(
        case_id=case_id
    ).order_by(FileAsset.created_at.desc()).all()
    extraction = {}
    for asset in assets:
        data = safe_json(asset.extraction_json)
        if data:
            extraction[asset.id] = data
    templates = FileAsset.query.options(defer(FileAsset.content)).filter(
        db.or_(
            FileAsset.asset_type == "template",
            db.and_(FileAsset.asset_type == "case_template", FileAsset.case_id == case_id),
        )
    ).order_by(FileAsset.created_at.desc()).all()
    recommended_template = matching_master_template(case.bank_name)
    return render_template(
        "case_detail.html", case=case, valuation=valuation, assets=assets,
        extraction=extraction, templates=templates, ai_enabled=ai_enabled(),
        document_ai_enabled=document_ai_enabled(),
        ai_model=OPENAI_MODEL, case_profile=safe_json(case.extracted_json),
        recommended_template=recommended_template,
        # Report controls stay visible on every case page. Hiding them behind
        # ?report=1 caused operators to press the heavier Process AI action.
        report_mode=True,
    )


@app.route("/cases/<int:case_id>/update", methods=["POST"])
@login_required
def update_case(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    for field in [
        "application_number", "customer_name", "contact_number", "property_address",
        "bank_name", "branch_name", "case_type", "visit_by", "status"
    ]:
        setattr(case, field, request.form.get(field, getattr(case, field)) or "")
    db.session.commit()
    flash("Case details saved.", "success")
    return redirect(url_for("case_detail", case_id=case_id))


@app.route("/cases/<int:case_id>/km", methods=["POST"])
@login_required
def update_case_km(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    km = numeric_km(request.form.get("distance_from_branch", ""))
    if km is None:
        flash("Billing ke liye valid K.M. enter karein.", "error")
    else:
        stored = safe_json(case.extracted_json)
        profile = dict(stored.get("case_profile") or stored.get("email") or stored)
        profile["distance_from_branch"] = km
        if "case_profile" in stored or "email" in stored:
            stored["case_profile"] = profile
        else:
            stored = {"case_profile": profile}
        case.extracted_json = json.dumps(stored, ensure_ascii=False, default=str)
        db.session.commit()
        flash(f"{km:g} K.M. MIS aur billing ke liye saved.", "success")
    return redirect(request.referrer or url_for("dashboard"))


@app.route("/cases/<int:case_id>/upload", methods=["POST"])
@login_required
def upload_case_files(case_id):
    return handle_case_upload(case_id, request.form.get("upload_kind", "visit"))


def handle_case_upload(case_id, upload_kind):
    case = ValuationCase.query.get_or_404(case_id)
    uploaded = request.files.getlist("files")
    count = 0
    page_number = 0
    if upload_kind == "documents":
        default_type = "document"
        source_kind = "property_document"
    elif upload_kind == "visit_form":
        default_type = "visit_data"
        source_kind = "visit_data"
    elif upload_kind == "visit":
        default_type = "photo"
        source_kind = "visit_data"
    else:
        default_type = "photo"
        source_kind = "visit_data"
    for item in uploaded:
        if not item or not item.filename:
            continue
        filename = secure_filename(item.filename)
        ext = Path(filename).suffix.lower()
        if ext == ".zip":
            try:
                # Werkzeug already spools larger requests to a temporary file.
                # Read ZIP members directly from that stream instead of keeping
                # a second full compressed copy in the 512 MB web worker.
                item.stream.seek(0)
                with zipfile.ZipFile(item.stream) as zf:
                    for member in safe_zip_members(zf):
                        inner_name = secure_filename(Path(member.filename).name)
                        inner_ext = Path(inner_name).suffix.lower()
                        if inner_ext not in DOCUMENT_EXTENSIONS | PHOTO_EXTENSIONS:
                            continue
                        inner = zf.read(member)
                        if upload_kind == "documents":
                            asset_type = "document"
                        elif upload_kind == "visit_form":
                            asset_type = "visit_data"
                        elif inner_ext in PHOTO_EXTENSIONS:
                            asset_type = "photo"
                        else:
                            asset_type = default_type
                        page_number += 1
                        store_asset(
                            case_id, asset_type, inner_name, inner,
                            category=(
                                f"Visit Form Page {page_number}"
                                if upload_kind == "visit_form" else None
                            ),
                            source_kind=source_kind, process_ai=False,
                        )
                        count += 1
            except (zipfile.BadZipFile, ValueError) as exc:
                message = str(exc) if isinstance(exc, ValueError) else "File is not a valid ZIP."
                flash(f"{filename}: {message}", "error")
        elif ext in PHOTO_EXTENSIONS:
            try:
                content = read_upload_limited(item)
            except ValueError as exc:
                flash(str(exc), "error")
                continue
            if upload_kind == "documents":
                asset_type = "document"
            elif upload_kind == "visit_form":
                asset_type = "visit_data"
            else:
                asset_type = "photo"
            page_number += 1
            store_asset(
                case_id, asset_type, filename, content, item.mimetype,
                category=(
                    f"Visit Form Page {page_number}"
                    if upload_kind == "visit_form" else None
                ),
                source_kind=source_kind, process_ai=False,
            )
            count += 1
        elif ext in DOCUMENT_EXTENSIONS:
            try:
                content = read_upload_limited(item)
            except ValueError as exc:
                flash(str(exc), "error")
                continue
            page_number += 1
            store_asset(
                case_id, default_type, filename, content, item.mimetype,
                category=(
                    f"Visit Form Page {page_number}"
                    if upload_kind == "visit_form" else None
                ),
                source_kind=source_kind, process_ai=False,
            )
            count += 1
        db.session.expire_all()
        gc.collect()
    if count:
        case.status = "Files Uploaded - AI Pending"
        db.session.commit()
    flash(f"{count} file(s) processed.", "success")
    return redirect(url_for("case_detail", case_id=case_id))


def quick_asset_type(filename):
    ext = Path(filename).suffix.lower()
    name = Path(filename).stem.casefold()
    if ext in PHOTO_EXTENSIONS:
        return "photo", "visit_data"
    if any(token in name for token in (
        "visit", "engineer", "site data", "site_data", "inspection", "sketch",
    )):
        return "visit_data", "visit_data"
    return "document", "property_document"


@app.route("/cases/<int:case_id>/upload/all", methods=["POST"])
@login_required
def upload_all_case_files(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    count = 0
    for item in request.files.getlist("files"):
        if not item or not item.filename:
            continue
        filename = secure_filename(item.filename)
        content = item.read()
        ext = Path(filename).suffix.lower()
        if ext == ".zip":
            try:
                with zipfile.ZipFile(io.BytesIO(content)) as bundle:
                    for member in safe_zip_members(bundle):
                        inner_name = secure_filename(Path(member.filename).name)
                        if Path(inner_name).suffix.lower() not in DOCUMENT_EXTENSIONS:
                            continue
                        asset_type, source_kind = quick_asset_type(inner_name)
                        store_asset(
                            case.id, asset_type, inner_name, bundle.read(member),
                            source_kind=source_kind, process_ai=False,
                        )
                        count += 1
            except (zipfile.BadZipFile, ValueError) as exc:
                message = str(exc) if isinstance(exc, ValueError) else "File is not a valid ZIP."
                flash(f"{filename}: {message}", "error")
        elif ext in DOCUMENT_EXTENSIONS:
            asset_type, source_kind = quick_asset_type(filename)
            store_asset(
                case.id, asset_type, filename, content, item.mimetype,
                source_kind=source_kind, process_ai=False,
            )
            count += 1

    template = request.files.get("template")
    if template and template.filename:
        filename = secure_filename(template.filename)
        if Path(filename).suffix.lower() in TEMPLATE_EXTENSIONS:
            db.session.add(FileAsset(
                case_id=case.id,
                asset_type="case_template",
                category=case.bank_name,
                filename=filename,
                mime_type=template.mimetype,
                content=template.read(),
            ))
            count += 1
        else:
            flash("Report format sirf XLSX, XLSM ya DOCX hona chahiye.", "error")

    if count:
        case.status = (
            "Portal Pending"
            if case.status == "Portal Pending"
            else "Files Uploaded - Review Pending"
        )
        db.session.commit()
    flash(f"{count} file(s) upload ho gayi. Ab details check karke report banayein.", "success")
    return redirect(url_for("case_detail", case_id=case.id))


@app.route("/cases/<int:case_id>/upload/visit", methods=["POST"])
@login_required
def upload_visit_files(case_id):
    return handle_case_upload(case_id, "visit")


@app.route("/cases/<int:case_id>/upload/visit-form", methods=["POST"])
@login_required
def upload_visit_form_files(case_id):
    return handle_case_upload(case_id, "visit_form")


@app.route("/cases/<int:case_id>/upload/documents", methods=["POST"])
@login_required
def upload_property_documents(case_id):
    return handle_case_upload(case_id, "documents")


@app.route("/cases/<int:case_id>/upload/template", methods=["POST"])
@login_required
def upload_case_template(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    item = request.files.get("template")
    if not item or not item.filename:
        flash("Select a bank valuation template.", "error")
        return redirect(url_for("case_detail", case_id=case_id))
    filename = secure_filename(item.filename)
    if Path(filename).suffix.lower() not in TEMPLATE_EXTENSIONS:
        flash("Only .xlsx, .xlsm or .docx report formats are supported.", "error")
        return redirect(url_for("case_detail", case_id=case_id))
    db.session.add(FileAsset(
        case_id=case.id, asset_type="case_template",
        category=request.form.get("bank_name", "").strip() or case.bank_name,
        filename=filename, mime_type=item.mimetype, content=item.read(),
    ))
    case.status = "Report Format Uploaded"
    db.session.commit()
    flash("Case-specific valuation format uploaded.", "success")
    return redirect(url_for("case_detail", case_id=case_id))


def valuation_as_dict(valuation):
    if not valuation:
        return {}
    fields = [
        "land_area", "land_rate", "builtup_area", "construction_rate", "age_years",
        "depreciation_percent", "govt_land_rate", "govt_construction_rate",
        "conservative_percent", "distress_percent", "land_value",
        "gross_building_value", "depreciation_amount", "net_building_value",
        "market_value", "conservative_value", "distress_value", "govt_value", "remarks",
    ]
    return {field: getattr(valuation, field) for field in fields}


def embedded_pdf_photos(asset):
    """Return the largest real image from each uploaded PDF page.

    Visit PDFs commonly contain one full-page site photograph plus a tiny
    scanner watermark. Property-document PDFs use the same structure. Keeping
    only the largest image prevents logos/watermarks from entering reports.
    """
    if Path(asset.filename).suffix.lower() != ".pdf":
        return []
    if asset.asset_type not in {"document", "visit_data"}:
        return []
    name = Path(asset.filename).stem.casefold()
    is_property_document = any(token in name for token in (
        "property_paper", "property paper", "registry", "sale_deed", "sale deed",
    ))
    is_visit_source = asset.asset_type == "visit_data" or any(
        token in name for token in ("visit", "inspection", "site_data", "site data")
    )
    if not (is_property_document or is_visit_source):
        return []
    output = []
    try:
        reader = PdfReader(io.BytesIO(asset.content))
        for page_index, page in enumerate(reader.pages):
            images = list(page.images)
            if not images:
                continue
            largest = max(images, key=lambda item: len(item.data or b""))
            if len(largest.data or b"") < 10_000:
                continue
            if is_property_document:
                category = "Property Document"
            elif page_index == 0:
                category = "Front Elevation"
            else:
                category = "Other Site Photo"
            output.append({
                "filename": f"{Path(asset.filename).stem}_page_{page_index + 1}_{largest.name}",
                "category": category,
                "content": largest.data,
            })
    except Exception:
        return []
    return output


@app.route("/cases/<int:case_id>/process-ai", methods=["POST"])
@login_required
def process_case_ai(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    paid_mode = document_ai_enabled()
    assets = FileAsset.query.filter_by(case_id=case_id).all()
    document_extractions = []
    visit_extractions = []
    processed = 0
    for asset in assets:
        inferred_type, inferred_source = quick_asset_type(asset.filename)
        if asset.asset_type == "document" and inferred_type == "visit_data":
            asset.asset_type = "visit_data"
        if asset.asset_type == "photo":
            result = classify_property_photo(asset.filename, asset.content)
            category = result.get("category", "Other Site Photo")
            asset.category = category
            asset.extraction_json = json.dumps(result, ensure_ascii=False)
            processed += 1
        elif asset.asset_type in {"document", "visit_data"}:
            source_kind = (
                "visit_data" if asset.asset_type == "visit_data" else inferred_source
            )
            asset.extracted_text = asset.extracted_text or extract_basic_text(asset.filename, asset.content)
            extraction = ai_extract_document(
                asset.filename, asset.content, asset.extracted_text or "", source_kind
            )
            asset.extraction_json = json.dumps(extraction, ensure_ascii=False)
            processed += 1
        db.session.add(asset)
    db.session.commit()

    for asset in assets:
        extraction = safe_json(asset.extraction_json)
        if not extraction:
            continue
        if asset.asset_type == "document":
            document_extractions.append(extraction)
            if any(token in Path(asset.filename).stem.casefold() for token in (
                "technical_report", "technical report", "valuation_report", "valuation report",
            )):
                visit_extractions.append(extract_property_asset(
                    asset.filename, asset.content, asset.extracted_text or "", "visit_data"
                ))
        elif asset.asset_type == "visit_data":
            visit_extractions.append(extraction)
    current = safe_json(case.extracted_json)
    email_data = current.get("email", current if "application_number" in current else {})
    valuation = Valuation.query.filter_by(case_id=case_id).first()
    case_profile = build_case_profile(
        email_data, document_extractions, visit_extractions, valuation_as_dict(valuation)
    )
    previous_profile = current.get("case_profile", {})
    if previous_profile.get("source_reviewed"):
        for field in SOURCE_REVIEW_FIELDS:
            if field in previous_profile:
                case_profile[field] = previous_profile[field]
        case_profile["source_reviewed"] = True
        case_profile["source_reviewed_at"] = previous_profile.get(
            "source_reviewed_at", ""
        )
        case_profile["survey_khasra_plot_no"] = (
            case_profile.get("survey_khasra_plot_no_as_per_docs")
            or case_profile.get("survey_khasra_plot_no_as_per_site")
            or ""
        )
    case.extracted_json = json.dumps(
        {"email": email_data, "case_profile": case_profile}, ensure_ascii=False, default=str
    )
    valuation = valuation or Valuation(case_id=case_id)
    draft_sources = valuation_defaults_from_profile(case_profile)
    for field, source_value in draft_sources.items():
        current_value = getattr(valuation, field, 0)
        standing_rule_field = field in {
            "depreciation_percent", "conservative_percent", "distress_percent",
        }
        old_generic_default = (
            (field == "conservative_percent" and current_value == 80)
            or (field == "distress_percent" and current_value == 70)
        )
        if (
            source_value not in ("", None)
            and (not current_value or standing_rule_field or old_generic_default)
        ):
            setattr(valuation, field, numeric_from_value(source_value))
    recalculated = valuation_calculation({
        "land_area": valuation.land_area,
        "land_rate": valuation.land_rate,
        "builtup_area": valuation.builtup_area,
        "construction_rate": valuation.construction_rate,
        "age_years": valuation.age_years,
        "depreciation_percent": valuation.depreciation_percent,
        "govt_land_rate": valuation.govt_land_rate,
        "govt_construction_rate": valuation.govt_construction_rate,
        "conservative_percent": valuation.conservative_percent,
        "distress_percent": valuation.distress_percent,
        "remarks": valuation.remarks or case_profile.get("remarks", ""),
    })
    for field, value in recalculated.items():
        setattr(valuation, field, value)
    db.session.add(valuation)
    if not case.application_number:
        case.application_number = str(case_profile.get("application_number", ""))
    if not case.customer_name or case.customer_name == "To be reviewed":
        case.customer_name = (
            case_profile.get("customer_name") or case_profile.get("applicant_name")
            or ""
        )
    if not case.property_address:
        case.property_address = (
            case_profile.get("property_address_as_per_site")
            or case_profile.get("property_address_as_per_docs") or ""
        )
    if not case.contact_number:
        case.contact_number = str(case_profile.get("contact_number", ""))
    case.status = (
        "AI Processed - Review Data"
        if paid_mode else "Local Files Processed - Review Data"
    )
    db.session.commit()
    flash(
        (
            f"ChatGPT + local fallback processed {processed} file(s)."
            if paid_mode else
            f"Free local mode processed {processed} file(s)."
        )
        + " Available data se valuation fields auto-filled.",
        "success",
    )
    return redirect(url_for("case_detail", case_id=case_id))


SOURCE_REVIEW_FIELDS = (
    "owner_name",
    "title_document_number",
    "registration_number",
    "registration_date",
    "land_tenure",
    "approving_authority",
    "plan_details",
    "construction_permission",
    "property_address_as_per_docs",
    "survey_khasra_plot_no_as_per_docs",
    "land_area_as_per_docs",
    "builtup_area_as_per_docs",
    "north_boundary_as_per_docs",
    "south_boundary_as_per_docs",
    "east_boundary_as_per_docs",
    "west_boundary_as_per_docs",
    "property_usage_as_per_docs",
    "property_address_as_per_site",
    "survey_khasra_plot_no_as_per_site",
    "land_area_as_per_site",
    "builtup_area_as_per_site",
    "north_boundary_as_per_site",
    "south_boundary_as_per_site",
    "east_boundary_as_per_site",
    "west_boundary_as_per_site",
    "property_usage_as_per_site",
    "road_width",
    "road_type",
    "latitude",
    "longitude",
    "occupancy",
    "person_met",
    "visit_engineer",
    "visit_date",
    "structure_type",
    "construction_quality",
    "construction_stage",
    "construction_year",
    "property_age_years",
    "residual_age_years",
    "number_of_floors",
    "floor_wise_usage",
    "room_configuration",
)


@app.route("/cases/<int:case_id>/source-review", methods=["POST"])
@login_required
def save_source_review(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    stored = safe_json(case.extracted_json)
    profile = dict(stored.get("case_profile") or stored.get("email") or stored)
    for field in SOURCE_REVIEW_FIELDS:
        if field in request.form:
            profile[field] = request.form.get(field, "").strip()
    profile["survey_khasra_plot_no"] = (
        profile.get("survey_khasra_plot_no_as_per_docs")
        or profile.get("survey_khasra_plot_no_as_per_site")
        or ""
    )
    profile["source_reviewed"] = True
    profile["source_reviewed_at"] = datetime.now(APP_TIMEZONE).isoformat()
    email_data = stored.get("email", {})
    case.extracted_json = json.dumps(
        {"email": email_data, "case_profile": profile},
        ensure_ascii=False,
        default=str,
    )
    case.property_address = (
        profile.get("property_address_as_per_site")
        or profile.get("property_address_as_per_docs")
        or case.property_address
        or ""
    )
    case.status = "Source Data Reviewed"
    db.session.commit()
    flash(
        "Documents aur Actual Site ki reviewed details save ho gayi. "
        "Ab report isi separation se banegi.",
        "success",
    )
    return redirect(url_for("case_detail", case_id=case_id))


@app.route("/assets/<int:asset_id>")
@login_required
def download_asset(asset_id):
    asset = FileAsset.query.get_or_404(asset_id)
    return send_file(
        io.BytesIO(asset.content), mimetype=asset.mime_type,
        as_attachment=True, download_name=asset.filename
    )


@app.route("/assets/<int:asset_id>/delete", methods=["POST"])
@login_required
def delete_asset(asset_id):
    asset = FileAsset.query.get_or_404(asset_id)
    case_id = asset.case_id
    db.session.delete(asset)
    db.session.commit()
    return redirect(url_for("case_detail", case_id=case_id))


@app.route("/cases/<int:case_id>/valuation", methods=["POST"])
@login_required
def save_valuation(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    data = valuation_calculation(request.form)
    valuation = Valuation.query.filter_by(case_id=case_id).first() or Valuation(case_id=case_id)
    for key, value in data.items():
        setattr(valuation, key, value)
    db.session.add(valuation)
    case.status = "Valuation Completed"
    db.session.commit()
    flash("Valuation calculated and saved.", "success")
    return redirect(url_for("case_detail", case_id=case_id))


@app.route("/templates", methods=["GET", "POST"])
@login_required
def templates_page():
    if request.method == "POST":
        item = request.files.get("template")
        bank_name = request.form.get("bank_name", "").strip()
        if item and Path(item.filename).suffix.lower() in TEMPLATE_EXTENSIONS:
            content = item.read()
            asset = FileAsset(
                asset_type="template", category=bank_name,
                filename=secure_filename(item.filename), mime_type=item.mimetype,
                content=content,
            )
            db.session.add(asset)
            db.session.commit()
            flash("Bank template uploaded.", "success")
    templates = FileAsset.query.filter_by(asset_type="template").order_by(FileAsset.created_at.desc()).all()
    return render_template("templates.html", templates=templates)


@app.route("/templates/<int:asset_id>/delete", methods=["POST"])
@login_required
def delete_template(asset_id):
    asset = FileAsset.query.get_or_404(asset_id)
    db.session.delete(asset)
    db.session.commit()
    return redirect(url_for("templates_page"))


@app.route("/conveyance")
@login_required
def conveyance_page():
    try:
        month = normalized_payment_month(request.args.get("month"))
    except ValueError:
        month = datetime.now(APP_TIMEZONE).strftime("%Y-%m")
    profiles = StaffPaymentProfile.query.order_by(
        StaffPaymentProfile.active.desc(), StaffPaymentProfile.name
    ).all()
    entries = StaffMonthlyPayment.query.filter_by(month=month).join(
        StaffPaymentProfile
    ).order_by(StaffPaymentProfile.category, StaffPaymentProfile.name).all()
    totals = {
        "base": sum(item.base_amount for item in entries),
        "conveyance": sum(item.conveyance_amount for item in entries),
        "advance": sum(item.advance_amount for item in entries),
        "adjustment": sum(item.adjustment_amount for item in entries),
        "final": sum(item.final_amount for item in entries),
    }
    return render_template(
        "conveyance.html", month=month, profiles=profiles, entries=entries, totals=totals
    )


@app.route("/conveyance/profiles/save", methods=["POST"])
@login_required
def save_staff_payment_profile():
    profile_id = request.form.get("profile_id", "").strip()
    name = re.sub(r"\s+", " ", request.form.get("name", "")).strip()
    try:
        if not name:
            raise ValueError("Staff/office name zaroori hai.")
        category = request.form.get("category", "Staff")
        if category not in {"Staff", "Office"}:
            raise ValueError("Category valid nahi hai.")
        payment_mode = request.form.get("payment_mode", "Per Visit")
        if payment_mode not in {"Salary", "Per Visit", "Fixed / Other"}:
            raise ValueError("Payment basis valid nahi hai.")
        profile = (
            db.session.get(StaffPaymentProfile, int(profile_id))
            if profile_id.isdigit() else None
        )
        duplicate = StaffPaymentProfile.query.filter(
            db.func.lower(StaffPaymentProfile.name) == name.casefold()
        ).first()
        if duplicate and (not profile or duplicate.id != profile.id):
            profile = duplicate
        profile = profile or StaffPaymentProfile(name=name)
        profile.name = name
        profile.category = category
        profile.payment_mode = payment_mode
        profile.fixed_salary = nonnegative_number(
            request.form.get("fixed_salary"), "Monthly salary/fixed amount"
        )
        profile.visit_rate = nonnegative_number(request.form.get("visit_rate"), "Visit rate")
        profile.petrol_rate = nonnegative_number(request.form.get("petrol_rate"), "Petrol rate")
        profile.active = request.form.get("active", "1") == "1"
        db.session.add(profile)
        db.session.commit()
        flash(f"{profile.name} ki payment settings saved.", "success")
    except (ValueError, TypeError) as exc:
        db.session.rollback()
        flash(str(exc), "error")
    return redirect(url_for("conveyance_page", month=request.form.get("month", "")))


@app.route("/conveyance/profiles/import", methods=["POST"])
@login_required
def import_staff_payment_profiles():
    upload = request.files.get("staff_file")
    try:
        if not upload or not upload.filename or Path(upload.filename).suffix.lower() != ".xlsx":
            raise ValueError("Employee names import ke liye .xlsx file upload karein.")
        names = staff_names_from_workbook(upload)
        added = 0
        for name in names:
            existing = StaffPaymentProfile.query.filter(
                db.func.lower(StaffPaymentProfile.name) == name.casefold()
            ).first()
            if existing:
                continue
            is_office = normalized_header(name) == "office"
            db.session.add(StaffPaymentProfile(
                name=name,
                category="Office" if is_office else "Staff",
                payment_mode="Fixed / Other" if is_office else "Per Visit",
            ))
            added += 1
        db.session.commit()
        flash(f"{added} naye staff name imported; existing names duplicate nahi hue.", "success")
    except (ValueError, TypeError) as exc:
        db.session.rollback()
        flash(str(exc), "error")
    return redirect(url_for("conveyance_page", month=request.form.get("month", "")))


@app.route("/conveyance/monthly/save", methods=["POST"])
@login_required
def save_staff_monthly_payment():
    month = request.form.get("month", "")
    try:
        month = normalized_payment_month(month)
        profile = db.session.get(
            StaffPaymentProfile, int(request.form.get("profile_id", "0"))
        )
        if not profile or not profile.active:
            raise ValueError("Active staff/office name select karein.")
        amounts = staff_payment_amounts(
            profile,
            visits=request.form.get("visits"),
            km=request.form.get("km"),
            base_override=request.form.get("base_override"),
            advance=request.form.get("advance_amount"),
            adjustment=request.form.get("adjustment_amount"),
        )
        entry = StaffMonthlyPayment.query.filter_by(
            profile_id=profile.id, month=month
        ).first() or StaffMonthlyPayment(profile_id=profile.id, month=month)
        for field, value in amounts.items():
            setattr(entry, field, value)
        entry.status = request.form.get("status", "Pending")
        if entry.status not in {"Pending", "Paid"}:
            raise ValueError("Payment status valid nahi hai.")
        entry.notes = request.form.get("notes", "").strip()
        db.session.add(entry)
        db.session.commit()
        flash(f"{profile.name} ka {month} final payable ₹{entry.final_amount:,.2f} saved.", "success")
    except (ValueError, TypeError) as exc:
        db.session.rollback()
        flash(str(exc), "error")
    return redirect(url_for("conveyance_page", month=month))


@app.route("/conveyance/export")
@login_required
def export_staff_payments():
    try:
        month = normalized_payment_month(request.args.get("month"))
    except ValueError as exc:
        flash(str(exc), "error")
        return redirect(url_for("conveyance_page"))
    entries = StaffMonthlyPayment.query.filter_by(month=month).join(
        StaffPaymentProfile
    ).order_by(StaffPaymentProfile.category, StaffPaymentProfile.name).all()
    if not entries:
        flash("Selected month me export ke liye koi payment row nahi hai.", "error")
        return redirect(url_for("conveyance_page", month=month))
    content = staff_payments_workbook(month, entries)
    return send_file(
        io.BytesIO(content), as_attachment=True,
        download_name=f"SVAI_Conveyance_MIS_{month}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route("/billing", methods=["GET", "POST"])
@login_required
def billing_page():
    default_from, default_to = current_month_range()
    saved_templates = BillingTemplate.query.order_by(
        BillingTemplate.bank_name, BillingTemplate.created_at.desc()
    ).all()
    if request.method == "POST":
        bank_name = request.form.get("bank_name", "").strip()
        branch_name = request.form.get("branch_name", "").strip()
        new_template = request.files.get("billing_template")
        selected_id = request.form.get("billing_template_id", "").strip()
        try:
            if new_template and new_template.filename:
                suffix = Path(new_template.filename).suffix.lower()
                if suffix not in {".xlsx"}:
                    raise ValueError("Billing format ke liye .xlsx Excel file upload karein.")
                if not bank_name:
                    raise ValueError("Naya format save karne ke liye Bank Name zaroori hai.")
                template = BillingTemplate(
                    bank_name=bank_name,
                    branch_name=branch_name,
                    filename=secure_filename(new_template.filename),
                    mime_type=new_template.mimetype,
                    content=new_template.read(),
                )
                db.session.add(template)
                db.session.commit()
                flash("Original bank billing format safely saved. Ab isi format se bill banega.", "success")
                return redirect(url_for("billing_page", template_id=template.id))
            if not selected_id:
                raise ValueError("Bank ka saved invoice format select karein, ya naya format upload karein.")
            template = BillingTemplate.query.get_or_404(int(selected_id))
            bank_name = bank_name or template.bank_name
            branch_name = branch_name or template.branch_name or ""
            submitted_slab_amounts = [value for value in request.form.getlist("slab_amount[]") if str(value).strip()]
            slabs = parse_billing_slabs(request.form) if submitted_slab_amounts else saved_billing_slabs(bank_name, branch_name)
            if not slabs:
                raise ValueError("Is bank ke KM rates pehle Save Rates me set karein.")
            save_billing_slabs(bank_name, branch_name, slabs)
            db.session.commit()
            source = request.form.get("source", "live")
            if source == "upload":
                mis_file = request.files.get("mis_file")
                if not mis_file or not mis_file.filename:
                    raise ValueError("Upload MIS option ke liye MIS .xlsx file select karein.")
                rows = billing_upload_rows(mis_file)
            else:
                from_date = parse_iso_date(request.form.get("from"), default_from)
                to_date = parse_iso_date(request.form.get("to"), default_to)
                cases = filter_cases_by_dates(ValuationCase.query, from_date, to_date).order_by(
                    db.func.coalesce(ValuationCase.email_received_at, ValuationCase.created_at)
                ).all()
                rows = billing_case_rows(cases)
            selected_bank = normalized_header(bank_name)
            if selected_bank:
                rows = [row for row in rows if normalized_header(row.get("bank_name")) == selected_bank]
            selected_branch = normalized_header(branch_name or template.branch_name)
            if selected_branch:
                rows = [row for row in rows if normalized_header(row.get("branch_name")) == selected_branch]
            if not rows:
                raise ValueError("Is bank/range ke liye koi MIS case nahi mila. Bank name ya date range check karein.")
            output, pending_rows = generate_billing_workbook(template.content, rows, slabs)
            if pending_rows:
                flash(
                    f"{len(pending_rows)} case(s) me KM/rate slab match nahi hua; unki Fee blank rakhi gayi hai. "
                    "MIS me K.M fill karke phir bill generate karein.", "error"
                )
            filename = secure_filename(
                f"{template.bank_name}_Invoice_{datetime.now(APP_TIMEZONE):%Y%m%d_%H%M}.xlsx"
            )
            return send_file(
                io.BytesIO(output), as_attachment=True, download_name=filename,
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except (ValueError, TypeError) as exc:
            flash(str(exc), "error")
        except Exception as exc:
            flash(f"Original invoice format fill nahi ho saka; koi generic bill nahi banaya gaya. {exc}", "error")
    selected_template_id = request.args.get("template_id", "")
    selected_template = BillingTemplate.query.get(selected_template_id) if str(selected_template_id).isdigit() else None
    saved_slabs = saved_billing_slabs(
        selected_template.bank_name, selected_template.branch_name or ""
    ) if selected_template else []
    return render_template(
        "billing.html", templates=saved_templates, default_from=default_from,
        default_to=default_to, selected_template_id=str(selected_template_id),
        saved_slabs=saved_slabs,
    )


@app.route("/billing/templates/<int:template_id>/delete", methods=["POST"])
@login_required
def delete_billing_template(template_id):
    template = BillingTemplate.query.get_or_404(template_id)
    db.session.delete(template)
    db.session.commit()
    flash("Billing format removed.", "success")
    return redirect(url_for("billing_page"))


def resolve_case_report_template(case, template_id=""):
    template = None
    if str(template_id).isdigit():
        template = FileAsset.query.options(defer(FileAsset.content)).filter(
            FileAsset.id == int(template_id),
            db.or_(
                FileAsset.asset_type == "template",
                db.and_(FileAsset.asset_type == "case_template", FileAsset.case_id == case.id),
            ),
        ).first()
    if template is None:
        template = FileAsset.query.options(defer(FileAsset.content)).filter_by(
            case_id=case.id, asset_type="case_template"
        ).order_by(FileAsset.created_at.desc()).first()
    if template is None:
        template = matching_master_template(case.bank_name)
    return template


@app.route("/cases/<int:case_id>/local-report-data")
@login_required
def local_report_data(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    valuation = Valuation.query.filter_by(case_id=case_id).first()
    stored = safe_json(case.extracted_json)
    profile = dict(stored.get("case_profile") or stored.get("email") or stored)
    profile.update({
        "application_number": case.application_number,
        "customer_name": case.customer_name,
        "contact_number": case.contact_number,
        "property_address": case.property_address,
        "bank_name": case.bank_name,
        "branch_name": case.branch_name,
        "case_type": case.case_type,
        "visit_by": case.visit_by,
        "report_date": datetime.now(APP_TIMEZONE).strftime("%d-%m-%Y"),
        **valuation_as_dict(valuation),
    })
    assets = FileAsset.query.options(defer(FileAsset.content)).filter(
        FileAsset.case_id == case_id,
        FileAsset.asset_type.in_(("document", "visit_data", "photo")),
    ).order_by(FileAsset.id).all()
    template = resolve_case_report_template(case, request.args.get("template_id", ""))
    if template is None:
        return jsonify({"error": "Bank report format missing."}), 400
    extension = Path(template.filename).suffix.lower()
    if extension not in TEMPLATE_EXTENSIONS:
        return jsonify({"error": "Unsupported bank report format."}), 400
    report_name = secure_filename(
        f"SVAI_DRAFT_{case.application_number or case.id}_"
        f"{case.customer_name or 'Pending_Name'}{extension}"
    )
    return jsonify({
        "profile": profile,
        "report_name": report_name,
        "template": {
            "url": url_for("download_asset", asset_id=template.id),
            "filename": template.filename,
        },
        "assets": [{
            "url": url_for("download_asset", asset_id=asset.id),
            "filename": asset.filename,
            "asset_type": asset.asset_type,
            "category": asset.category or "",
        } for asset in assets],
    })


@app.route("/cases/<int:case_id>/local-report-save", methods=["POST"])
@login_required
def save_local_report(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    item = request.files.get("report")
    if not item or not item.filename:
        return jsonify({"error": "Generated report missing."}), 400
    filename = secure_filename(item.filename)
    extension = Path(filename).suffix.lower()
    if extension not in TEMPLATE_EXTENSIONS:
        return jsonify({"error": "Generated report type is not supported."}), 400
    content = item.read()
    mime_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if extension == ".docx" else
        "application/vnd.ms-excel.sheet.macroEnabled.12"
        if extension == ".xlsm" else
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    report = FileAsset(
        case_id=case.id, asset_type="report", category=case.bank_name,
        filename=filename, mime_type=mime_type, content=content,
    )
    db.session.add(report)
    for asset in FileAsset.query.filter(
        FileAsset.case_id == case.id,
        FileAsset.asset_type.in_(("document", "photo", "visit_data", "case_template")),
    ).all():
        db.session.delete(asset)
    case.status = "Draft Report Generated"
    db.session.commit()
    return jsonify({
        "status": "ok",
        "report_id": report.id,
        "download_url": url_for("download_asset", asset_id=report.id),
    })


@app.route("/cases/<int:case_id>/report", methods=["POST"])
@login_required
def generate_report(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    if case.status == "Portal Pending":
        flash("Ye online portal case hai; local report pending rakhi gayi hai.", "error")
        return redirect(url_for("case_detail", case_id=case_id))
    valuation = Valuation.query.filter_by(case_id=case_id).first()
    if not valuation:
        valuation = Valuation(case_id=case_id, remarks="DRAFT: valuation figures pending review")
        db.session.add(valuation)
        db.session.commit()
    assets = FileAsset.query.filter_by(case_id=case_id).all()
    missing_inputs = []
    if not any(asset.asset_type == "document" for asset in assets):
        missing_inputs.append("Property Documents")
    if not any(asset.asset_type == "visit_data" for asset in assets):
        missing_inputs.append("Visit Form / MP Kisan")
    if not any(asset.asset_type == "photo" for asset in assets):
        missing_inputs.append("Site Photos")
    if missing_inputs:
        flash(
            "Report generate nahi hui. Mandatory upload missing: "
            + ", ".join(missing_inputs),
            "error",
        )
        return redirect(url_for("case_detail", case_id=case_id, report=1))
    # Generate must never silently create a blank report when the operator
    # forgot to press "Process All Files" first. Process current inputs here
    # and rebuild the source-separated profile before filling the template.
    stored_before = safe_json(case.extracted_json)
    reviewed_profile = stored_before.get("case_profile", {})
    source_reviewed = bool(reviewed_profile.get("source_reviewed"))
    document_extractions = []
    visit_extractions = []
    for asset in assets:
        inferred_type, inferred_source = quick_asset_type(asset.filename)
        if asset.asset_type == "document" and inferred_type == "visit_data":
            asset.asset_type = "visit_data"
        if asset.asset_type == "photo" and not asset.extraction_json and not source_reviewed:
            result = classify_property_photo(asset.filename, asset.content)
            asset.category = result.get("category", "Other Site Photo")
            asset.extraction_json = json.dumps(result, ensure_ascii=False)
        elif (
            asset.asset_type in {"document", "visit_data"}
            and not asset.extraction_json
            and not source_reviewed
        ):
            source_kind = "visit_data" if asset.asset_type == "visit_data" else inferred_source
            asset.extracted_text = asset.extracted_text or extract_basic_text(
                asset.filename, asset.content
            )
            extraction = ai_extract_document(
                asset.filename, asset.content, asset.extracted_text or "", source_kind
            )
            asset.extraction_json = json.dumps(extraction, ensure_ascii=False)
        extraction = safe_json(asset.extraction_json)
        if asset.asset_type == "document" and extraction:
            document_extractions.append(extraction)
        elif asset.asset_type == "visit_data" and extraction:
            visit_extractions.append(extraction)
        db.session.add(asset)
    email_data = stored_before.get(
        "email", stored_before if "application_number" in stored_before else {}
    )
    refreshed_profile = build_case_profile(
        email_data, document_extractions, visit_extractions, valuation_as_dict(valuation)
    )
    if reviewed_profile.get("source_reviewed"):
        for field in SOURCE_REVIEW_FIELDS:
            if field in reviewed_profile:
                refreshed_profile[field] = reviewed_profile[field]
        refreshed_profile["source_reviewed"] = True
        refreshed_profile["source_reviewed_at"] = reviewed_profile.get(
            "source_reviewed_at", ""
        )
    case.extracted_json = json.dumps(
        {"email": email_data, "case_profile": refreshed_profile},
        ensure_ascii=False,
        default=str,
    )
    db.session.commit()
    template_id = request.form.get("template_id")
    stored = safe_json(case.extracted_json)
    profile = dict(stored.get("case_profile") or stored.get("email") or stored)
    profile.update({
        "application_number": case.application_number,
        "customer_name": case.customer_name,
        "contact_number": case.contact_number,
        "property_address": case.property_address,
        "bank_name": case.bank_name,
        "branch_name": case.branch_name,
        "case_type": case.case_type,
        "visit_by": case.visit_by,
        "report_date": datetime.now(APP_TIMEZONE).strftime("%d-%m-%Y"),
        **valuation_as_dict(valuation),
    })
    photo_assets = [
        {
            "filename": asset.filename,
            "category": asset.category,
            "content": asset.content,
        }
        for asset in assets if asset.asset_type == "photo"
    ]
    for asset in assets:
        photo_assets.extend(embedded_pdf_photos(asset))
    visit_images = [
        asset for asset in assets
        if (
            asset.asset_type == "visit_data"
            and Path(asset.filename).suffix.lower() in PHOTO_EXTENSIONS
            and (asset.category or "").startswith("Visit Form Page ")
        )
    ]
    visit_images.sort(
        key=lambda asset: numeric_from_value(asset.category or "", asset.id)
    )
    used_visit_ids = set()
    for asset in visit_images:
        filename_key = Path(asset.filename).stem.casefold()
        category = None
        if any(token in filename_key for token in ("google", "location", "map")):
            category = "Google Map"
        elif any(token in filename_key for token in ("kisan", "khasra", "bhulekh", "bhu")):
            category = "MP Kisan"
        if category:
            photo_assets.append({
                "filename": asset.filename,
                "category": category,
                "content": asset.content,
            })
            used_visit_ids.add(asset.id)
    if len(visit_images) >= 4:
        for asset, category in (
            (visit_images[-2], "Site Sketch"),
            (visit_images[-1], "Location Map"),
        ):
            if asset.id in used_visit_ids:
                continue
            photo_assets.append({
                "filename": asset.filename,
                "category": category,
                "content": asset.content,
            })
    output = None
    extension = ""
    mime_type = ""
    template = None
    if template_id:
        try:
            numeric_template_id = int(template_id)
        except (TypeError, ValueError):
            numeric_template_id = 0
        template = FileAsset.query.filter(
            FileAsset.id == numeric_template_id,
            db.or_(
                FileAsset.asset_type == "template",
                db.and_(FileAsset.asset_type == "case_template", FileAsset.case_id == case_id),
            ),
        ).first()
    if template is None:
        template = FileAsset.query.filter_by(
            case_id=case_id, asset_type="case_template"
        ).order_by(FileAsset.created_at.desc()).first()
    if template is None:
        template = matching_master_template(case.bank_name)
    if template is None:
        flash(
            "Pehle bank ka original Excel/Word valuation format upload aur select karein. "
            "SVAI generic format nahi banayega.",
            "error",
        )
        return redirect(url_for("case_detail", case_id=case_id, report=1))
    try:
        extension = Path(template.filename).suffix.lower()
        if extension == ".docx":
            output = fill_docx_template(
                template.content,
                profile,
                photo_assets,
                template_name=template.filename,
                bank_name=template.category or case.bank_name,
            )
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            output = fill_excel_template(
                template.content, profile, photo_assets,
                template_name=template.filename,
                bank_name=template.category or case.bank_name,
            )
            if extension == ".xlsm":
                mime_type = "application/vnd.ms-excel.sheet.macroEnabled.12"
            else:
                extension = ".xlsx"
                mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    except Exception as exc:
        flash(
            f"Original bank format ko safe tarike se fill nahi kiya ja saka; "
            f"generic/ टूटा report नहीं बनाया गया. {exc}",
            "error",
        )
        return redirect(url_for("case_detail", case_id=case_id))
    report_name = (
        f"SVAI_DRAFT_{case.application_number or case.id}_"
        f"{case.customer_name or 'Pending_Name'}{extension}"
    )
    report = FileAsset(
        case_id=case_id, asset_type="report", category=case.bank_name,
        filename=secure_filename(report_name), mime_type=mime_type,
        content=output,
    )
    db.session.add(report)
    # Keep the uploaded registry, visit sheets and site photographs. Operators
    # need those original sources for corrections and safe re-generation; a
    # successful draft must not silently destroy the case working set.
    case.status = "Draft Report Generated"
    db.session.commit()
    return send_file(
        io.BytesIO(output), as_attachment=True, download_name=report.filename,
        mimetype=report.mime_type
    )


@app.route("/cases/<int:case_id>/archive", methods=["POST"])
@login_required
def archive_case(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    case.archived = not case.archived
    case.status = "Archived" if case.archived else "Reopened"
    db.session.commit()
    return redirect(url_for("dashboard", archived="1" if case.archived else "0"))


@app.route("/cases/<int:case_id>/delete", methods=["POST"])
@login_required
def delete_case(case_id):
    case = ValuationCase.query.get_or_404(case_id)
    label = case.application_number or case.customer_name or f"Case #{case.id}"
    FileAsset.query.filter_by(case_id=case.id).delete(synchronize_session=False)
    Valuation.query.filter_by(case_id=case.id).delete(synchronize_session=False)
    db.session.delete(case)
    db.session.commit()
    flash(f"{label} MIS se permanently delete ho gaya.", "success")
    return redirect(url_for("dashboard"))


@app.route("/email-accounts", methods=["GET", "POST"])
@login_required
def email_accounts():
    if request.method == "POST":
        address = request.form.get("email", "").strip().lower()
        password = re.sub(r"\s+", "", request.form.get("password", ""))
        provider = request.form.get("provider", "").strip().lower()
        domain = address.rsplit("@", 1)[-1] if "@" in address else ""
        gmail_address = domain in {"gmail.com", "googlemail.com"}
        yahoo_address = domain == "yahoo.com" or domain.startswith("yahoo.")
        if provider not in {"gmail", "yahoo"} or not (gmail_address or yahoo_address):
            flash("Only Gmail and Yahoo email IDs are supported.", "error")
            return redirect(url_for("email_accounts"))
        if (provider == "gmail" and not gmail_address) or (provider == "yahoo" and not yahoo_address):
            flash("Selected provider does not match the email ID.", "error")
            return redirect(url_for("email_accounts"))
        if len(password) != 16:
            flash("Enter the 16-character Gmail/Yahoo app password (spaces are ignored).", "error")
            return redirect(url_for("email_accounts"))
        custom_host = ""
        bank_name = request.form.get("bank_name", "").strip()
        host = detect_imap(address, provider, custom_host)
        try:
            mail = imaplib.IMAP4_SSL(host, 993)
            mail.login(address, password)
            mail.logout()
            account = EmailAccount.query.filter_by(email=address).first() or EmailAccount(email=address)
            account.encrypted_password = encrypt_password(password)
            account.provider = provider
            account.imap_host = custom_host
            account.bank_name = bank_name
            account.active = True
            db.session.add(account)
            db.session.commit()
            flash("Email verified and linked.", "success")
        except Exception as exc:
            flash(f"Email verification failed: {exc}", "error")
    accounts = EmailAccount.query.order_by(EmailAccount.created_at.desc()).all()
    date_from, date_to = current_month_range()
    return render_template(
        "email_accounts.html", accounts=accounts, date_from=date_from, date_to=date_to,
        ai_enabled=ai_enabled(), ai_model=OPENAI_MODEL,
    )


@app.route("/email-accounts/<int:account_id>/fetch", methods=["POST"])
@login_required
def fetch_one_email(account_id):
    account = EmailAccount.query.get_or_404(account_id)
    default_from, default_to = current_month_range()
    date_from = parse_iso_date(request.form.get("from"), default_from)
    date_to = parse_iso_date(request.form.get("to"), default_to)
    try:
        result = fetch_email_account(
            account, date_from, date_to, enrich_documents=False
        )
        flash(
            f"{result['created']} new, {result.get('updated', 0)} corrected valuation "
            f"case(s); {result.get('deduplicated', 0)} duplicate MIS row(s) merged; "
            f"{result['ignored']} unrelated email(s) ignored."
            + (f" {result.get('warning')}" if result.get("warning") else ""),
            "error" if result.get("warning") else "success",
        )
    except Exception as exc:
        flash(f"Email fetch failed: {exc}", "error")
    return redirect(url_for("email_accounts"))


@app.route("/email-accounts/fetch-all", methods=["POST"])
@login_required
def fetch_all_emails():
    total = 0
    updated = 0
    deduplicated = 0
    errors = []
    warnings = []
    ignored = 0
    default_from, default_to = current_month_range()
    date_from = parse_iso_date(request.form.get("from"), default_from)
    date_to = parse_iso_date(request.form.get("to"), default_to)
    for account in EmailAccount.query.filter_by(active=True).all():
        try:
            result = fetch_email_account(
                account, date_from, date_to, enrich_documents=False
            )
            total += result["created"]
            updated += result.get("updated", 0)
            deduplicated += result.get("deduplicated", 0)
            ignored += result["ignored"]
            if result.get("warning"):
                warnings.append(f"{account.email}: {result['warning']}")
        except Exception as exc:
            db.session.rollback()
            errors.append(f"{account.email}: {exc}")
    flash(
        f"{total} valuation case(s) added; {updated} existing case(s) corrected; "
        f"{deduplicated} duplicate MIS row(s) merged; "
        f"{ignored} unrelated email(s) ignored. "
        f"Range: {date_from:%d-%m-%Y} to {date_to:%d-%m-%Y}."
        + (f" Warnings: {'; '.join(warnings)}" if warnings else "")
        + (f" Errors: {'; '.join(errors)}" if errors else ""),
        "success" if not errors and not warnings else "error",
    )
    return redirect(url_for(
        "dashboard", **{"from": date_from.isoformat(), "to": date_to.isoformat()}
    ))


@app.route("/email-accounts/<int:account_id>/delete", methods=["POST"])
@login_required
def delete_email_account(account_id):
    account = EmailAccount.query.get_or_404(account_id)
    db.session.delete(account)
    db.session.commit()
    return redirect(url_for("email_accounts"))


@app.route("/mis/export")
@login_required
def export_mis():
    default_from, default_to = current_month_range()
    date_from = parse_iso_date(request.args.get("from"), default_from)
    date_to = parse_iso_date(request.args.get("to"), default_to)
    query = filter_cases_by_dates(
        filter_billing_ready_cases(ValuationCase.query.filter_by(archived=False)),
        date_from, date_to
    )
    cases = query.order_by(
        db.func.coalesce(ValuationCase.email_received_at, ValuationCase.created_at)
    ).all()
    mis_template = SEED_TEMPLATES_DIR / "MIS format.xlsx"
    if mis_template.exists():
        wb = load_workbook(mis_template)
        ws = wb["ALL BANK"] if "ALL BANK" in wb.sheetnames else wb.active
        if ws.max_row > 1:
            ws.delete_rows(2, ws.max_row - 1)
        # The bank MIS format predates the received-time column. Insert it in
        # place so its layout remains the source template rather than rebuilt.
        if str(ws.cell(1, 3).value or "").strip().casefold() != "time":
            ws.insert_cols(3)
        headers = [
            "SR NO", "Date", "Time", "CUSTOMER NAME", "APPLICATION NO",
            "CONTACT NUMBER", "CASE TYPE", "BANK", "STATUS", "ADDRESS",
            "VISIT BY", "BRANCH", "Pending", "K.M",
        ]
        for column, header in enumerate(headers, 1):
            ws.cell(1, column).value = header
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "ALL BANK"
        ws.append([
            "SR NO", "Date", "Time", "CUSTOMER NAME", "APPLICATION NO", "CONTACT NUMBER",
            "CASE TYPE", "BANK", "STATUS", "ADDRESS", "VISIT BY", "BRANCH",
            "Pending", "K.M",
        ])
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="FFE699")
    for index, case in enumerate(cases, 1):
        stored = safe_json(case.extracted_json)
        profile = stored.get("case_profile") or stored.get("email") or stored
        case_date = case.email_received_at or case.created_at
        ws.append([
            index, case_date if case_date else "",
            case_date.strftime("%I:%M %p") if case_date else "",
            case.customer_name, case.application_number, case.contact_number,
            case.case_type, case.bank_name, case.status, case.property_address,
            case.visit_by, case.branch_name, profile.get("pending", ""),
            profile.get("distance_from_branch", profile.get("km", "")),
        ])
        ws.cell(index + 1, 2).number_format = "d-mmm-yy"
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return send_file(
        out, as_attachment=True,
        download_name=f"SVAI_MIS_{date_from:%Y%m%d}_{date_to:%Y%m%d}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route("/profile/password", methods=["POST"])
@login_required
def change_password():
    user = User.query.get(session["user_id"])
    current = request.form.get("current_password", "")
    new = request.form.get("new_password", "")
    if not user or not check_password_hash(user.password_hash, current):
        flash("Current password is incorrect.", "error")
    elif len(new) < 8:
        flash("New password must contain at least 8 characters.", "error")
    else:
        user.password_hash = generate_password_hash(new)
        db.session.commit()
        flash("Password changed.", "success")
    return redirect(url_for("dashboard"))


def scheduled_email_fetch():
    with app.app_context():
        # Automatic capture scans only today's incoming assignments. Older
        # dates are intentionally recovered through the manual From/To fetch.
        # A short minute job cannot monopolize the fetch lock with a month scan.
        today = datetime.now(APP_TIMEZONE).date()
        for account in EmailAccount.query.filter_by(active=True).all():
            try:
                # The minute scheduler is for reliable case capture, not heavy
                # attachment processing. Full documents remain available to a
                # user-triggered fetch/report flow when missing fields matter.
                result = fetch_email_account(
                    account, today, today, enrich_documents=False
                )
                app.logger.info(
                    "Scheduled MIS fetch completed for %s (%s to %s): %s new, %s updated",
                    account.email,
                    today,
                    today,
                    result.get("created", 0),
                    result.get("updated", 0),
                )
            except Exception as exc:
                app.logger.warning("Scheduled email fetch failed for %s: %s", account.email, exc)
            finally:
                # Repeated MIME/IMAP parsing can leave large temporary arenas
                # resident in a 512 MB Render worker. Release ORM references and
                # collect between mailboxes so the web UI remains responsive.
                db.session.remove()
                gc.collect()
                try:
                    import ctypes
                    ctypes.CDLL("libc.so.6").malloc_trim(0)
                except (OSError, AttributeError):
                    pass


def start_scheduler():
    if not BackgroundScheduler:
        return
    if os.getenv("ENABLE_EMAIL_SCHEDULER", "false").lower() != "true":
        return
    scheduler = BackgroundScheduler(daemon=True)
    minutes = max(1, int(os.getenv("EMAIL_FETCH_MINUTES", "1")))
    scheduler.add_job(
        scheduled_email_fetch,
        "interval",
        minutes=minutes,
        id="email_fetch",
        replace_existing=True,
        max_instances=1,
        coalesce=True,
    )
    scheduler.start()


with app.app_context():
    db.create_all()
start_scheduler()


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", "8000")), debug=os.getenv("FLASK_DEBUG") == "1")
